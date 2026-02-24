"""
创建测试用的Word文档
"""
from docx import Document
from pathlib import Path

# 创建文档
doc = Document()

# 添加标题
doc.add_heading('基于深度学习的图像识别算法研究', level=0)

# 添加摘要
doc.add_heading('摘要', level=1)
doc.add_paragraph(
    '本文提出了一种基于深度学习的图像识别算法。'
    '通过构建卷积神经网络模型，实现对图像的高精度分类。'
    '实验结果表明，该算法在ImageNet数据集上达到了95%的准确率，'
    '相比传统方法提升了20个百分点。'
)

# 添加关键词
doc.add_paragraph('关键词：深度学习；图像识别；卷积神经网络；计算机视觉')

# 添加引言
doc.add_heading('1. 引言', level=1)
doc.add_paragraph(
    '随着人工智能技术的快速发展，图像识别已经成为计算机视觉领域的研究热点。'
    '传统的图像识别方法主要依赖人工设计的特征，存在特征表达能力不足的问题。'
    '深度学习技术的出现为图像识别带来了新的解决方案。'
)

# 添加方法部分
doc.add_heading('2. 研究方法', level=1)
doc.add_paragraph(
    '本文采用卷积神经网络(CNN)作为基础模型。'
    '网络结构包括5层卷积层和3层全连接层。'
    '使用ReLU作为激活函数，Dropout防止过拟合。'
    '优化器选用Adam，学习率设置为0.001。'
)

# 添加实验结果
doc.add_heading('3. 实验结果', level=1)
doc.add_paragraph(
    '在ImageNet数据集上进行了充分的实验验证。'
    '训练集包含120万张图像，测试集包含5万张图像。'
    '经过100个epoch的训练，模型收敛良好。'
    '最终在测试集上达到95.2%的Top-1准确率。'
)

# 添加结论
doc.add_heading('4. 结论', level=1)
doc.add_paragraph(
    '本文提出的深度学习图像识别算法取得了优异的性能。'
    '未来工作将关注模型的轻量化和实时性优化。'
)

# 添加参考文献
doc.add_heading('参考文献', level=1)
doc.add_paragraph('[1] Krizhevsky A, Sutskever I, Hinton G E. ImageNet classification with deep convolutional neural networks[J]. Communications of the ACM, 2017.')
doc.add_paragraph('[2] He K, Zhang X, Ren S, et al. Deep residual learning for image recognition[C]//CVPR, 2016.')

# 保存文档
output_path = Path('test_paper.docx')
doc.save(output_path)

print(f'测试文档已创建: {output_path.absolute()}')
