"""
论文格式模板生成脚本
生成3个标准学术论文 .docx 模板文件及对应 .json 元数据文件
目标目录: backend/templates/builtin/
"""

import json
import os
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
from docx.shared import Length


# ──────────────────────────────────────────────────────────────
# 工具函数
# ──────────────────────────────────────────────────────────────

def set_east_asia_font(run, font_name: str):
    """设置段落 run 的中文（东亚）字体，使用 XML 操作。"""
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), run.font.name if run.font.name else 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), run.font.name if run.font.name else 'Times New Roman')


def set_style_east_asia_font(style, font_name: str):
    """设置样式级别的中文字体（直接操作 style 的 rPr XML）。"""
    rPr = style.font._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def set_paragraph_spacing(para, space_before_pt=None, space_after_pt=None,
                           line_spacing_multiple=None, line_spacing_pt=None):
    """统一设置段落间距和行距。"""
    fmt = para.paragraph_format
    if space_before_pt is not None:
        fmt.space_before = Pt(space_before_pt)
    if space_after_pt is not None:
        fmt.space_after = Pt(space_after_pt)
    if line_spacing_multiple is not None:
        from docx.oxml import OxmlElement as _Oxml
        from docx.enum.text import WD_LINE_SPACING
        fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        fmt.line_spacing = line_spacing_multiple
    if line_spacing_pt is not None:
        from docx.enum.text import WD_LINE_SPACING
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = Pt(line_spacing_pt)


def add_heading_paragraph(doc, text: str, level: int,
                           font_name: str, east_asia_font: str,
                           font_size_pt: float, bold: bool = True,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT,
                           space_before_pt=0.0, space_after_pt=0.0,
                           line_spacing_multiple=1.5):
    """添加标题段落，手动控制所有格式（避免依赖内置样式继承）。"""
    para = doc.add_paragraph()
    fmt = para.paragraph_format
    fmt.alignment = alignment
    fmt.space_before = Pt(space_before_pt)
    fmt.space_after = Pt(space_after_pt)
    from docx.enum.text import WD_LINE_SPACING
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line_spacing_multiple

    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.font.bold = bold
    set_east_asia_font(run, east_asia_font)
    return para


def add_body_paragraph(doc, text: str,
                       font_name: str, east_asia_font: str,
                       font_size_pt: float, bold: bool = False,
                       alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       space_before_pt=0.0, space_after_pt=0.0,
                       line_spacing_multiple=1.5,
                       first_line_indent_chars=2):
    """添加正文段落。"""
    para = doc.add_paragraph()
    fmt = para.paragraph_format
    fmt.alignment = alignment
    fmt.space_before = Pt(space_before_pt)
    fmt.space_after = Pt(space_after_pt)
    from docx.enum.text import WD_LINE_SPACING
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line_spacing_multiple
    # 首行缩进（字符数 × 字号）
    if first_line_indent_chars > 0:
        fmt.first_line_indent = Pt(font_size_pt * first_line_indent_chars)

    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.font.bold = bold
    set_east_asia_font(run, east_asia_font)
    return para


def write_json_metadata(path: Path, data: dict):
    """将元数据写入 JSON 文件（UTF-8 + ensure_ascii=False）。"""
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


# ──────────────────────────────────────────────────────────────
# 模板 1：清华大学毕业论文模板
# ──────────────────────────────────────────────────────────────

def create_tsinghua_template(output_dir: Path) -> Path:
    """
    清华大学毕业论文模板
    页面：A4，上2.5 下2.5 左3 右2 (cm)
    正文：宋体 12pt，1.5倍行距
    一级标题：黑体 16pt，居中，段前24pt段后6pt
    二级标题：黑体 14pt，左对齐，段前12pt段后6pt
    三级标题：黑体 12pt，左对齐，段前6pt段后0pt
    英文/数字：Times New Roman
    """
    doc = Document()

    # ---------- 页面设置 ----------
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)

    # ---------- 修改内置 Normal 样式（正文基础）----------
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    set_style_east_asia_font(normal_style, 'SimSun')

    # ---------- 封面 ----------
    # 校名
    p_school = doc.add_paragraph()
    p_school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_school.add_run('清华大学')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(22)
    r.font.bold = True
    set_east_asia_font(r, 'SimHei')

    # 论文类型
    p_type = doc.add_paragraph()
    p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_type.add_run('本科毕业论文')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(18)
    r.font.bold = True
    set_east_asia_font(r, 'SimHei')

    # 空行
    for _ in range(3):
        doc.add_paragraph()

    # 论文标题
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run('（论文题目）')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(16)
    r.font.bold = True
    set_east_asia_font(r, 'SimHei')

    # 空行
    for _ in range(5):
        doc.add_paragraph()

    # 作者信息表格
    info_items = [
        ('院　　系：', '（学院/系名称）'),
        ('专　　业：', '（专业名称）'),
        ('姓　　名：', '（学生姓名）'),
        ('学　　号：', '（学号）'),
        ('指导教师：', '（教师姓名）'),
        ('完成日期：', '2026年2月'),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_label = p.add_run(label)
        r_label.font.name = 'Times New Roman'
        r_label.font.size = Pt(12)
        set_east_asia_font(r_label, 'SimSun')
        r_value = p.add_run(value)
        r_value.font.name = 'Times New Roman'
        r_value.font.size = Pt(12)
        set_east_asia_font(r_value, 'SimSun')

    # ---------- 分页 ----------
    doc.add_page_break()

    # ---------- 摘要 ----------
    add_heading_paragraph(
        doc, '摘  要', level=0,
        font_name='Times New Roman', east_asia_font='SimHei',
        font_size_pt=16, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before_pt=24, space_after_pt=6,
        line_spacing_multiple=1.5
    )
    add_body_paragraph(
        doc,
        '本文以（研究主题）为研究对象，综合运用（研究方法），对（研究问题）进行了系统深入的研究。'
        '首先，（第一部分概述）；其次，（第二部分概述）；最后，（结论概述）。'
        '研究结果表明，（主要结论）。本文的研究成果对于（应用领域）具有重要的理论和实践意义。',
        font_name='Times New Roman', east_asia_font='SimSun',
        font_size_pt=12, line_spacing_multiple=1.5,
        space_before_pt=0, space_after_pt=6
    )
    p_kw = doc.add_paragraph()
    r_kw_label = p_kw.add_run('关键词：')
    r_kw_label.font.name = 'Times New Roman'
    r_kw_label.font.size = Pt(12)
    r_kw_label.font.bold = True
    set_east_asia_font(r_kw_label, 'SimSun')
    r_kw = p_kw.add_run('关键词1；关键词2；关键词3；关键词4；关键词5')
    r_kw.font.name = 'Times New Roman'
    r_kw.font.size = Pt(12)
    set_east_asia_font(r_kw, 'SimSun')

    # ---------- ABSTRACT ----------
    doc.add_paragraph()
    add_heading_paragraph(
        doc, 'Abstract', level=0,
        font_name='Times New Roman', east_asia_font='SimHei',
        font_size_pt=16, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before_pt=24, space_after_pt=6,
        line_spacing_multiple=1.5
    )
    add_body_paragraph(
        doc,
        'This paper takes (research topic) as the object of study, '
        'comprehensively uses (research methods), and conducts a systematic '
        'and in-depth study on (research questions). The results show that '
        '(main conclusions). The findings of this study have important '
        'theoretical and practical significance for (application field).',
        font_name='Times New Roman', east_asia_font='SimSun',
        font_size_pt=12, line_spacing_multiple=1.5,
        space_before_pt=0, space_after_pt=6,
        first_line_indent_chars=0
    )
    p_kwen = doc.add_paragraph()
    r_kwen_label = p_kwen.add_run('Key words: ')
    r_kwen_label.font.name = 'Times New Roman'
    r_kwen_label.font.size = Pt(12)
    r_kwen_label.font.bold = True
    r_kwen = p_kwen.add_run('keyword1; keyword2; keyword3; keyword4; keyword5')
    r_kwen.font.name = 'Times New Roman'
    r_kwen.font.size = Pt(12)

    # ---------- 分页 -> 正文 ----------
    doc.add_page_break()

    # ---------- 第一章 ----------
    add_heading_paragraph(
        doc, '第一章  绪论', level=1,
        font_name='Times New Roman', east_asia_font='SimHei',
        font_size_pt=16, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before_pt=24, space_after_pt=6,
        line_spacing_multiple=1.5
    )

    add_heading_paragraph(
        doc, '1.1  研究背景与意义', level=2,
        font_name='Times New Roman', east_asia_font='SimHei',
        font_size_pt=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before_pt=12, space_after_pt=6,
        line_spacing_multiple=1.5
    )
    add_body_paragraph(
        doc,
        '随着（领域/背景描述），（研究问题）逐渐成为学术界和工业界关注的热点。'
        '（背景描述继续）……本课题的研究具有重要的理论价值和现实意义。',
        font_name='Times New Roman', east_asia_font='SimSun',
        font_size_pt=12, line_spacing_multiple=1.5
    )

    add_heading_paragraph(
        doc, '1.2  国内外研究现状', level=2,
        font_name='Times New Roman', east_asia_font='SimHei',
        font_size_pt=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before_pt=12, space_after_pt=6,
        line_spacing_multiple=1.5
    )
    add_body_paragraph(
        doc,
        '国内外学者对（研究主题）已开展了大量研究工作。（国外研究现状概述）[1]。'
        '（国内研究现状概述）[2,3]。然而，现有研究在（不足之处）方面仍存在一定局限。',
        font_name='Times New Roman', east_asia_font='SimSun',
        font_size_pt=12, line_spacing_multiple=1.5
    )

    add_heading_paragraph(
        doc, '1.2.1  国外研究现状', level=3,
        font_name='Times New Roman', east_asia_font='SimHei',
        font_size_pt=12, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before_pt=6, space_after_pt=0,
        line_spacing_multiple=1.5
    )
    add_body_paragraph(
        doc,
        '（国外研究现状详细描述。此处为示例占位文字，实际撰写时需替换为真实文献综述内容。）',
        font_name='Times New Roman', east_asia_font='SimSun',
        font_size_pt=12, line_spacing_multiple=1.5
    )

    add_heading_paragraph(
        doc, '1.3  研究内容与论文组织结构', level=2,
        font_name='Times New Roman', east_asia_font='SimHei',
        font_size_pt=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before_pt=12, space_after_pt=6,
        line_spacing_multiple=1.5
    )
    add_body_paragraph(
        doc,
        '本文共分为X章，各章节安排如下：\n'
        '第一章为绪论，介绍研究背景、研究意义、国内外研究现状及论文组织结构。\n'
        '第二章为（章节内容摘要）。\n'
        '第X章为总结与展望，对全文进行总结并展望未来研究方向。',
        font_name='Times New Roman', east_asia_font='SimSun',
        font_size_pt=12, line_spacing_multiple=1.5
    )

    # ---------- 第二章 ----------
    doc.add_page_break()
    add_heading_paragraph(
        doc, '第二章  （章节标题）', level=1,
        font_name='Times New Roman', east_asia_font='SimHei',
        font_size_pt=16, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before_pt=24, space_after_pt=6,
        line_spacing_multiple=1.5
    )
    add_heading_paragraph(
        doc, '2.1  （节标题）', level=2,
        font_name='Times New Roman', east_asia_font='SimHei',
        font_size_pt=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before_pt=12, space_after_pt=6,
        line_spacing_multiple=1.5
    )
    add_body_paragraph(
        doc,
        '（正文示例。此处为占位文字，正文内容宋体小四（12pt），Times New Roman，1.5倍行距，'
        '首行缩进2字符，两端对齐。）',
        font_name='Times New Roman', east_asia_font='SimSun',
        font_size_pt=12, line_spacing_multiple=1.5
    )

    # ---------- 参考文献 ----------
    doc.add_page_break()
    add_heading_paragraph(
        doc, '参考文献', level=0,
        font_name='Times New Roman', east_asia_font='SimHei',
        font_size_pt=16, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before_pt=24, space_after_pt=6,
        line_spacing_multiple=1.5
    )
    refs = [
        '[1] 作者1, 作者2. 文献标题[J]. 期刊名称, 2023, 卷(期): 起止页码.',
        '[2] Author A, Author B. Title of the article[J]. Journal Name, 2022, Vol(No): pages.',
        '[3] 作者. 专著名[M]. 出版地: 出版社, 2021.',
    ]
    for ref in refs:
        p_ref = doc.add_paragraph()
        fmt = p_ref.paragraph_format
        fmt.first_line_indent = Pt(-24)
        fmt.left_indent = Pt(24)
        from docx.enum.text import WD_LINE_SPACING
        fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        fmt.line_spacing = 1.5
        r_ref = p_ref.add_run(ref)
        r_ref.font.name = 'Times New Roman'
        r_ref.font.size = Pt(12)
        set_east_asia_font(r_ref, 'SimSun')

    # ---------- 致谢 ----------
    doc.add_page_break()
    add_heading_paragraph(
        doc, '致  谢', level=0,
        font_name='Times New Roman', east_asia_font='SimHei',
        font_size_pt=16, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before_pt=24, space_after_pt=6,
        line_spacing_multiple=1.5
    )
    add_body_paragraph(
        doc,
        '感谢导师（姓名）教授在本文写作过程中给予的悉心指导和大力支持。'
        '感谢（同学/同事）在研究过程中提供的帮助与支持。'
        '感谢家人的理解与支持，使我能够顺利完成学业。',
        font_name='Times New Roman', east_asia_font='SimSun',
        font_size_pt=12, line_spacing_multiple=1.5
    )

    out_path = output_dir / '清华大学毕业论文模板.docx'
    doc.save(str(out_path))
    return out_path


# ──────────────────────────────────────────────────────────────
# 模板 2：浙江大学毕业论文模板
# ──────────────────────────────────────────────────────────────

def create_zju_template(output_dir: Path) -> Path:
    """
    浙江大学毕业论文模板
    页面：A4，上3 下2.5 左3 右2.5 (cm)
    正文：宋体 12pt，1.5倍行距
    一级标题：黑体 14pt，居中，段前18pt段后6pt
    二级标题：黑体 12pt，左对齐，段前12pt段后0pt
    三级标题：宋体 12pt 加粗，左对齐
    """
    doc = Document()

    # ---------- 页面设置 ----------
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(3.0)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

    # ---------- Normal 样式 ----------
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    set_style_east_asia_font(normal_style, 'SimSun')

    # ---------- 封面 ----------
    # 校徽占位
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_logo.add_run('ZHEJIANG UNIVERSITY')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.font.bold = True

    doc.add_paragraph()

    p_school = doc.add_paragraph()
    p_school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_school.add_run('浙江大学')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(22)
    r.font.bold = True
    set_east_asia_font(r, 'SimHei')

    p_type = doc.add_paragraph()
    p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_type.add_run('毕业论文（设计）')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(16)
    r.font.bold = True
    set_east_asia_font(r, 'SimHei')

    for _ in range(4):
        doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run('（论文题目）')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(18)
    r.font.bold = True
    set_east_asia_font(r, 'SimHei')

    for _ in range(6):
        doc.add_paragraph()

    info_items = [
        ('学院（系）：', '（学院名称）'),
        ('专    业：', '（专业名称）'),
        ('学    号：', '（学号）'),
        ('学生姓名：', '（姓名）'),
        ('指导教师：', '（姓名，职称）'),
        ('合作导师：', '（姓名，职称）（如有）'),
        ('完成日期：', '2026 年 2 月'),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_label = p.add_run(label)
        r_label.font.name = 'Times New Roman'
        r_label.font.size = Pt(12)
        set_east_asia_font(r_label, 'SimSun')
        r_value = p.add_run(value)
        r_value.font.name = 'Times New Roman'
        r_value.font.size = Pt(12)
        set_east_asia_font(r_value, 'SimSun')

    # ---------- 分页 ----------
    doc.add_page_break()

    # ---------- 摘要 ----------
    add_heading_paragraph(
        doc, '摘  要', level=0,
        font_name='Times New Roman', east_asia_font='SimHei',
        font_size_pt=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before_pt=18, space_after_pt=6,
        line_spacing_multiple=1.5
    )
    add_body_paragraph(
        doc,
        '（摘要正文。阐述研究目的、研究方法、主要研究结果和结论。摘要应简明扼要，'
        '一般为300字左右。本示例为格式占位文字，请替换为实际论文摘要内容。）',
        font_name='Times New Roman', east_asia_font='SimSun',
        font_size_pt=12, line_spacing_multiple=1.5,
        space_before_pt=0, space_after_pt=6
    )
    p_kw = doc.add_paragraph()
    r_kw_label = p_kw.add_run('关键词：')
    r_kw_label.font.name = 'Times New Roman'
    r_kw_label.font.size = Pt(12)
    r_kw_label.font.bold = True
    set_east_asia_font(r_kw_label, 'SimSun')
    r_kw = p_kw.add_run('关键词1；关键词2；关键词3；关键词4；关键词5')
    r_kw.font.name = 'Times New Roman'
    r_kw.font.size = Pt(12)
    set_east_asia_font(r_kw, 'SimSun')

    doc.add_paragraph()

    add_heading_paragraph(
        doc, 'Abstract', level=0,
        font_name='Times New Roman', east_asia_font='SimHei',
        font_size_pt=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before_pt=18, space_after_pt=6,
        line_spacing_multiple=1.5
    )
    add_body_paragraph(
        doc,
        '(Abstract content. Describe the purpose, methods, main results and '
        'conclusions of the research. The abstract should be concise and '
        'generally around 300 words. This is a placeholder for formatting purposes.)',
        font_name='Times New Roman', east_asia_font='SimSun',
        font_size_pt=12, line_spacing_multiple=1.5,
        space_before_pt=0, space_after_pt=6,
        first_line_indent_chars=0
    )
    p_kwen = doc.add_paragraph()
    r_kwen_label = p_kwen.add_run('Key words: ')
    r_kwen_label.font.name = 'Times New Roman'
    r_kwen_label.font.size = Pt(12)
    r_kwen_label.font.bold = True
    r_kwen = p_kwen.add_run('keyword1; keyword2; keyword3; keyword4; keyword5')
    r_kwen.font.name = 'Times New Roman'
    r_kwen.font.size = Pt(12)

    # ---------- 分页 -> 正文 ----------
    doc.add_page_break()

    # ---------- 第1章 ----------
    add_heading_paragraph(
        doc, '1  绪论', level=1,
        font_name='Times New Roman', east_asia_font='SimHei',
        font_size_pt=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before_pt=18, space_after_pt=6,
        line_spacing_multiple=1.5
    )

    add_heading_paragraph(
        doc, '1.1  研究背景', level=2,
        font_name='Times New Roman', east_asia_font='SimHei',
        font_size_pt=12, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before_pt=12, space_after_pt=0,
        line_spacing_multiple=1.5
    )
    add_body_paragraph(
        doc,
        '（研究背景描述。介绍本研究领域的发展现状，说明开展本研究的重要性和必要性。'
        '此处为示例占位文字，请替换为实际内容。）',
        font_name='Times New Roman', east_asia_font='SimSun',
        font_size_pt=12, line_spacing_multiple=1.5
    )

    add_heading_paragraph(
        doc, '1.2  研究目的与意义', level=2,
        font_name='Times New Roman', east_asia_font='SimHei',
        font_size_pt=12, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before_pt=12, space_after_pt=0,
        line_spacing_multiple=1.5
    )
    add_body_paragraph(
        doc,
        '（研究目的与意义说明。阐述本研究拟解决的科学问题或工程问题，'
        '以及研究成果的理论价值和实践价值。）',
        font_name='Times New Roman', east_asia_font='SimSun',
        font_size_pt=12, line_spacing_multiple=1.5
    )

    add_heading_paragraph(
        doc, '1.2.1  理论意义', level=3,
        font_name='Times New Roman', east_asia_font='SimSun',
        font_size_pt=12, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before_pt=6, space_after_pt=0,
        line_spacing_multiple=1.5
    )
    add_body_paragraph(
        doc,
        '（理论意义详细说明。三级标题采用宋体12pt加粗，此为示例。）',
        font_name='Times New Roman', east_asia_font='SimSun',
        font_size_pt=12, line_spacing_multiple=1.5
    )

    add_heading_paragraph(
        doc, '1.3  论文结构', level=2,
        font_name='Times New Roman', east_asia_font='SimHei',
        font_size_pt=12, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before_pt=12, space_after_pt=0,
        line_spacing_multiple=1.5
    )
    add_body_paragraph(
        doc,
        '本文共分X章。第一章绪论，介绍研究背景、目的与论文结构。'
        '第二章（内容概述）。第X章总结全文，展望未来研究方向。',
        font_name='Times New Roman', east_asia_font='SimSun',
        font_size_pt=12, line_spacing_multiple=1.5
    )

    # ---------- 第2章 ----------
    doc.add_page_break()
    add_heading_paragraph(
        doc, '2  （第二章标题）', level=1,
        font_name='Times New Roman', east_asia_font='SimHei',
        font_size_pt=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before_pt=18, space_after_pt=6,
        line_spacing_multiple=1.5
    )
    add_heading_paragraph(
        doc, '2.1  （节标题）', level=2,
        font_name='Times New Roman', east_asia_font='SimHei',
        font_size_pt=12, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before_pt=12, space_after_pt=0,
        line_spacing_multiple=1.5
    )
    add_body_paragraph(
        doc,
        '（正文示例内容。浙江大学毕业论文正文采用宋体12pt，Times New Roman，'
        '1.5倍行距，首行缩进2字符。此处为格式示例占位文字。）',
        font_name='Times New Roman', east_asia_font='SimSun',
        font_size_pt=12, line_spacing_multiple=1.5
    )

    # ---------- 参考文献 ----------
    doc.add_page_break()
    add_heading_paragraph(
        doc, '参考文献', level=0,
        font_name='Times New Roman', east_asia_font='SimHei',
        font_size_pt=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before_pt=18, space_after_pt=6,
        line_spacing_multiple=1.5
    )
    refs = [
        '[1] 作者1, 作者2. 参考文献标题[J]. 期刊名, 2024, 卷(期): 页码-页码.',
        '[2] Author A, Author B. Article title[J]. Journal Name, 2023, Vol(Issue): pp-pp.',
        '[3] 作者. 著作名称[M]. 出版地: 出版社, 2022.',
    ]
    for ref in refs:
        p_ref = doc.add_paragraph()
        fmt = p_ref.paragraph_format
        fmt.first_line_indent = Pt(-24)
        fmt.left_indent = Pt(24)
        from docx.enum.text import WD_LINE_SPACING
        fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        fmt.line_spacing = 1.5
        r_ref = p_ref.add_run(ref)
        r_ref.font.name = 'Times New Roman'
        r_ref.font.size = Pt(12)
        set_east_asia_font(r_ref, 'SimSun')

    out_path = output_dir / '浙江大学毕业论文模板.docx'
    doc.save(str(out_path))
    return out_path


# ──────────────────────────────────────────────────────────────
# 模板 3：中文核心期刊投稿模板
# ──────────────────────────────────────────────────────────────

def create_core_journal_template(output_dir: Path) -> Path:
    """
    中文核心期刊投稿模板
    页面：A4，上2.5 下2.5 左2.5 右2.5 (cm)
    正文：宋体 10.5pt（五号），1.3倍行距
    文章标题：黑体 16pt，居中
    作者：宋体 10.5pt，居中
    摘要关键词：楷体 10.5pt（用 SimSun 加注解，楷体用 KaiTi 但系统可能缺失 -> 用斜体模拟）
    一级标题：黑体 10.5pt，左对齐，末尾加全角句号
    二级标题：宋体 10.5pt 加粗，左对齐
    """
    # 注意：KaiTi（楷体）在部分系统中可用（Windows自带），尝试使用
    # 如不可用则降级为 SimSun（宋体）
    KAITI_FONT = 'KaiTi'  # Windows 楷体

    doc = Document()

    # ---------- 页面设置 ----------
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # ---------- Normal 样式 ----------
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(10.5)
    set_style_east_asia_font(normal_style, 'SimSun')

    # ---------- 文章标题 ----------
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p_title.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after  = Pt(6)
    from docx.enum.text import WD_LINE_SPACING
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.3
    r = p_title.add_run('（文章题目）')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(16)
    r.font.bold = True
    set_east_asia_font(r, 'SimHei')

    # ---------- 作者行 ----------
    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p_author.paragraph_format
    fmt.space_before = Pt(6)
    fmt.space_after  = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.3
    r = p_author.add_run('作者姓名¹，合作者姓名²')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10.5)
    set_east_asia_font(r, 'SimSun')

    # ---------- 作者单位 ----------
    p_aff = doc.add_paragraph()
    p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p_aff.paragraph_format
    fmt.space_before = Pt(2)
    fmt.space_after  = Pt(6)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.3
    r = p_aff.add_run('（1. 单位全称，省市 邮编；2. 单位全称，省市 邮编）')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9)
    set_east_asia_font(r, 'SimSun')

    # ---------- 中文摘要 ----------
    p_abs_label = doc.add_paragraph()
    fmt = p_abs_label.paragraph_format
    fmt.space_before = Pt(6)
    fmt.space_after  = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.3
    r_label = p_abs_label.add_run('摘要：')
    r_label.font.name = 'Times New Roman'
    r_label.font.size = Pt(10.5)
    r_label.font.bold = True
    set_east_asia_font(r_label, KAITI_FONT)
    r_content = p_abs_label.add_run(
        '本文研究了（研究主题），采用（研究方法）对（研究对象）进行了系统分析。'
        '研究结果表明，（主要结论）。本研究对于（应用领域）具有重要参考价值。'
        '（摘要一般200~300字，此处为示例占位文字。）'
    )
    r_content.font.name = 'Times New Roman'
    r_content.font.size = Pt(10.5)
    set_east_asia_font(r_content, KAITI_FONT)

    # ---------- 关键词 ----------
    p_kw = doc.add_paragraph()
    fmt = p_kw.paragraph_format
    fmt.space_before = Pt(3)
    fmt.space_after  = Pt(6)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.3
    r_kw_label = p_kw.add_run('关键词：')
    r_kw_label.font.name = 'Times New Roman'
    r_kw_label.font.size = Pt(10.5)
    r_kw_label.font.bold = True
    set_east_asia_font(r_kw_label, KAITI_FONT)
    r_kw = p_kw.add_run('关键词1；关键词2；关键词3；关键词4；关键词5')
    r_kw.font.name = 'Times New Roman'
    r_kw.font.size = Pt(10.5)
    set_east_asia_font(r_kw, KAITI_FONT)

    # ---------- 中图分类号 / 文献标识码 ----------
    p_class = doc.add_paragraph()
    fmt = p_class.paragraph_format
    fmt.space_before = Pt(3)
    fmt.space_after  = Pt(6)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.3
    r_class = p_class.add_run('中图分类号：XXXXX    文献标识码：A    文章编号：XXXX-XXXX(2026)XX-XXXX-XX')
    r_class.font.name = 'Times New Roman'
    r_class.font.size = Pt(10.5)
    set_east_asia_font(r_class, 'SimSun')

    # ---------- 分割线 ----------
    p_line = doc.add_paragraph()
    fmt = p_line.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after  = Pt(0)
    pPr = p_line._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ---------- 正文：一级标题 ----------
    def add_journal_h1(doc, text):
        """核心期刊一级标题：黑体 10.5pt，左对齐，编号+全角句号结尾"""
        para = doc.add_paragraph()
        fmt = para.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.space_before = Pt(6)
        fmt.space_after  = Pt(3)
        fmt.first_line_indent = Pt(0)
        from docx.enum.text import WD_LINE_SPACING
        fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        fmt.line_spacing = 1.3
        r = para.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10.5)
        r.font.bold = True
        set_east_asia_font(r, 'SimHei')
        return para

    def add_journal_h2(doc, text):
        """核心期刊二级标题：宋体 10.5pt 加粗，左对齐"""
        para = doc.add_paragraph()
        fmt = para.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.space_before = Pt(3)
        fmt.space_after  = Pt(0)
        fmt.first_line_indent = Pt(0)
        from docx.enum.text import WD_LINE_SPACING
        fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        fmt.line_spacing = 1.3
        r = para.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10.5)
        r.font.bold = True
        set_east_asia_font(r, 'SimSun')
        return para

    def add_journal_body(doc, text):
        """核心期刊正文：宋体 10.5pt，1.3倍行距，首行缩进2字符"""
        para = doc.add_paragraph()
        fmt = para.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt.space_before = Pt(0)
        fmt.space_after  = Pt(0)
        fmt.first_line_indent = Pt(10.5 * 2)
        from docx.enum.text import WD_LINE_SPACING
        fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        fmt.line_spacing = 1.3
        r = para.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10.5)
        set_east_asia_font(r, 'SimSun')
        return para

    # 注意：一级标题末尾用全角句号"。"
    add_journal_h1(doc, '1  引言。')
    add_journal_body(
        doc,
        '（引言内容。介绍研究背景、国内外研究现状及本文主要贡献。'
        '核心期刊正文采用宋体五号（10.5pt），Times New Roman，1.3倍行距，首行缩进2字符。'
        '此处为示例占位文字，请替换为实际内容[1]。）'
    )

    add_journal_h1(doc, '2  研究方法。')
    add_journal_h2(doc, '2.1  数据来源')
    add_journal_body(
        doc,
        '（数据来源说明。描述本研究所使用的数据集或实验材料，包括数据获取方式、'
        '样本规模、时间范围等基本信息[2]。）'
    )
    add_journal_h2(doc, '2.2  分析方法')
    add_journal_body(
        doc,
        '（分析方法说明。详细描述本研究采用的研究方法和技术路线，'
        '包括理论框架、数学模型或实验设计等[3]。）'
    )

    add_journal_h1(doc, '3  结果与分析。')
    add_journal_body(
        doc,
        '（结果与分析。呈现主要研究结果，结合图表进行详细分析和讨论，'
        '与已有研究进行比较，解释结果的意义。）'
    )

    add_journal_h1(doc, '4  结论。')
    add_journal_body(
        doc,
        '（结论。总结本文主要研究发现，指出研究的创新点和局限性，'
        '展望未来研究方向。）'
    )

    # ---------- 英文摘要 ----------
    doc.add_paragraph()
    p_en_title = doc.add_paragraph()
    p_en_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_en_title.add_run('（English Title of the Article）')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.font.bold = True

    p_en_author = doc.add_paragraph()
    p_en_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_en_author.add_run('Author Name¹, Co-author Name²')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10.5)

    p_en_aff = doc.add_paragraph()
    p_en_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_en_aff.add_run('(1. Institution, City, Postcode; 2. Institution, City, Postcode)')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9)

    p_en_abs = doc.add_paragraph()
    fmt = p_en_abs.paragraph_format
    fmt.space_before = Pt(6)
    from docx.enum.text import WD_LINE_SPACING
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.3
    r_en_abs_label = p_en_abs.add_run('Abstract: ')
    r_en_abs_label.font.name = 'Times New Roman'
    r_en_abs_label.font.size = Pt(10.5)
    r_en_abs_label.font.bold = True
    r_en_abs = p_en_abs.add_run(
        'This paper investigates (research topic) using (research methods). '
        'The results indicate that (main conclusions). '
        'This study provides valuable insights for (application field).'
    )
    r_en_abs.font.name = 'Times New Roman'
    r_en_abs.font.size = Pt(10.5)

    p_en_kw = doc.add_paragraph()
    fmt = p_en_kw.paragraph_format
    fmt.space_before = Pt(3)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.3
    r_en_kw_label = p_en_kw.add_run('Key words: ')
    r_en_kw_label.font.name = 'Times New Roman'
    r_en_kw_label.font.size = Pt(10.5)
    r_en_kw_label.font.bold = True
    r_en_kw = p_en_kw.add_run('keyword1; keyword2; keyword3; keyword4; keyword5')
    r_en_kw.font.name = 'Times New Roman'
    r_en_kw.font.size = Pt(10.5)

    # ---------- 参考文献 ----------
    doc.add_paragraph()
    p_ref_title = doc.add_paragraph()
    fmt = p_ref_title.paragraph_format
    fmt.space_before = Pt(6)
    fmt.space_after  = Pt(3)
    from docx.enum.text import WD_LINE_SPACING
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.3
    r = p_ref_title.add_run('参考文献')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10.5)
    r.font.bold = True
    set_east_asia_font(r, 'SimHei')

    refs = [
        '[1] 作者1, 作者2. 期刊文章标题[J]. 期刊名称, 2024, 卷(期): 起止页.',
        '[2] 作者. 著作名[M]. 出版地: 出版社, 2023.',
        '[3] Author A, Author B. Article title[J]. Journal Name, 2022, Vol(No): pp-pp.',
    ]
    for ref in refs:
        p_ref = doc.add_paragraph()
        fmt = p_ref.paragraph_format
        fmt.first_line_indent = Pt(-21)
        fmt.left_indent = Pt(21)
        from docx.enum.text import WD_LINE_SPACING
        fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        fmt.line_spacing = 1.3
        r_ref = p_ref.add_run(ref)
        r_ref.font.name = 'Times New Roman'
        r_ref.font.size = Pt(9)
        set_east_asia_font(r_ref, 'SimSun')

    out_path = output_dir / '中文核心期刊投稿模板.docx'
    doc.save(str(out_path))
    return out_path


# ──────────────────────────────────────────────────────────────
# 元数据定义
# ──────────────────────────────────────────────────────────────

METADATA_LIST = [
    {
        "template_id": "tsinghua_v1",
        "name": "清华大学毕业论文模板",
        "category": "universities",
        "school_or_journal": "清华大学",
        "description": "清华大学研究生/本科生毕业论文标准格式模板，参照清华大学学位论文撰写规范。"
                        "页面A4，上2.5cm下2.5cm左3cm右2cm；正文宋体12pt（小四），1.5倍行距；"
                        "一级标题黑体16pt居中，二级标题黑体14pt，三级标题黑体12pt。",
        "version": "1.0",
        "author": "system",
        "created_at": "2026-02-24T00:00:00",
        "updated_at": "2026-02-24T00:00:00",
        "usage_count": 0,
        "is_public": True,
        "tags": ["清华大学", "毕业论文", "学位论文", "本科", "硕士"],
        "file_path": None,
        "preview_image": None,
        "format_spec": {
            "page_size": "A4",
            "margins_cm": {"top": 2.5, "bottom": 2.5, "left": 3.0, "right": 2.0},
            "body_font": "宋体 12pt（小四）",
            "body_line_spacing": "1.5倍",
            "heading1": "黑体 16pt，居中，段前24pt，段后6pt",
            "heading2": "黑体 14pt，左对齐，段前12pt，段后6pt",
            "heading3": "黑体 12pt，左对齐，段前6pt，段后0pt",
            "english_font": "Times New Roman"
        }
    },
    {
        "template_id": "zju_v1",
        "name": "浙江大学毕业论文模板",
        "category": "universities",
        "school_or_journal": "浙江大学",
        "description": "浙江大学本科/研究生毕业论文（设计）标准格式模板，参照浙江大学本科生毕业论文撰写规范。"
                        "页面A4，上3cm下2.5cm左3cm右2.5cm；正文宋体12pt，1.5倍行距；"
                        "一级标题黑体14pt居中，二级标题黑体12pt，三级标题宋体12pt加粗。",
        "version": "1.0",
        "author": "system",
        "created_at": "2026-02-24T00:00:00",
        "updated_at": "2026-02-24T00:00:00",
        "usage_count": 0,
        "is_public": True,
        "tags": ["浙江大学", "毕业论文", "学位论文", "本科", "设计"],
        "file_path": None,
        "preview_image": None,
        "format_spec": {
            "page_size": "A4",
            "margins_cm": {"top": 3.0, "bottom": 2.5, "left": 3.0, "right": 2.5},
            "body_font": "宋体 12pt",
            "body_line_spacing": "1.5倍",
            "heading1": "黑体 14pt，居中，段前18pt，段后6pt",
            "heading2": "黑体 12pt，左对齐，段前12pt，段后0pt",
            "heading3": "宋体 12pt 加粗，左对齐",
            "english_font": "Times New Roman"
        }
    },
    {
        "template_id": "core_journal_v1",
        "name": "中文核心期刊投稿模板",
        "category": "journals",
        "school_or_journal": "中文核心期刊（通用）",
        "description": "适用于中文核心期刊（北大核心、CSSCI、CSCD等）的通用投稿格式模板。"
                        "页面A4，四边距均为2.5cm；正文宋体10.5pt（五号），1.3倍行距；"
                        "文章标题黑体16pt居中，作者宋体10.5pt居中，摘要关键词楷体10.5pt，"
                        "一级标题黑体10.5pt末尾加全角句号，二级标题宋体10.5pt加粗。",
        "version": "1.0",
        "author": "system",
        "created_at": "2026-02-24T00:00:00",
        "updated_at": "2026-02-24T00:00:00",
        "usage_count": 0,
        "is_public": True,
        "tags": ["核心期刊", "学术论文", "投稿", "北大核心", "CSSCI", "CSCD"],
        "file_path": None,
        "preview_image": None,
        "format_spec": {
            "page_size": "A4",
            "margins_cm": {"top": 2.5, "bottom": 2.5, "left": 2.5, "right": 2.5},
            "body_font": "宋体 10.5pt（五号）",
            "body_line_spacing": "1.3倍",
            "article_title": "黑体 16pt，居中",
            "author": "宋体 10.5pt，居中",
            "abstract_keywords": "楷体 10.5pt",
            "heading1": "黑体 10.5pt，左对齐，末尾加全角句号",
            "heading2": "宋体 10.5pt 加粗，左对齐",
            "english_font": "Times New Roman"
        }
    },
]


# ──────────────────────────────────────────────────────────────
# 主程序
# ──────────────────────────────────────────────────────────────

def main():
    # 确定输出目录（相对于本脚本所在目录 backend/）
    script_dir = Path(__file__).parent.resolve()
    output_dir = script_dir / 'templates' / 'builtin'
    output_dir.mkdir(parents=True, exist_ok=True)

    print('=' * 60)
    print('论文格式模板生成脚本')
    print(f'输出目录: {output_dir}')
    print('=' * 60)

    # 模板生成函数映射（按 template_id 对应）
    generators = {
        'tsinghua_v1':       create_tsinghua_template,
        'zju_v1':            create_zju_template,
        'core_journal_v1':   create_core_journal_template,
    }

    results = []

    for meta in METADATA_LIST:
        tid = meta['template_id']
        name = meta['name']
        gen_func = generators[tid]

        print(f'\n正在生成: {name}')

        # 生成 .docx
        try:
            docx_path = gen_func(output_dir)
            size_kb = docx_path.stat().st_size / 1024
            print(f'  .docx 生成成功')
            print(f'  路径: {docx_path}')
            print(f'  大小: {size_kb:.1f} KB')

            # 更新元数据中的 file_path
            meta_copy = dict(meta)
            meta_copy['file_path'] = str(docx_path)

            # 生成对应 .json 元数据
            json_path = output_dir / f'{docx_path.stem}.json'
            write_json_metadata(json_path, meta_copy)
            print(f'  .json 生成成功')
            print(f'  路径: {json_path}')

            results.append({
                'name': name,
                'docx': str(docx_path),
                'json': str(json_path),
                'size_kb': size_kb,
                'status': 'success'
            })

        except Exception as e:
            print(f'  [ERROR] 生成失败: {e}')
            import traceback
            traceback.print_exc()
            results.append({
                'name': name,
                'status': 'failed',
                'error': str(e)
            })

    # ---------- 汇总报告 ----------
    print('\n' + '=' * 60)
    print('生成结果汇总')
    print('=' * 60)
    success_count = sum(1 for r in results if r['status'] == 'success')
    print(f'成功: {success_count}/{len(results)}')
    for r in results:
        if r['status'] == 'success':
            print(f'  [OK]  {r["name"]}')
            print(f'        docx: {r["docx"]}')
            print(f'        json: {r["json"]}')
            print(f'        大小: {r["size_kb"]:.1f} KB')
        else:
            print(f'  [FAIL] {r["name"]} -> {r.get("error", "未知错误")}')

    print('\n所有模板生成完毕。')
    print(f'文件保存位置: {output_dir}')


if __name__ == '__main__':
    main()
