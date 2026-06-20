"""
POC: 样式注入方案验证脚本

目的: 验证"样式定义注入 + pStyle 打标签"方案相比当前 direct formatting 方案
     在正确性和性能上的提升。

用法:
  cd backend
  python scripts/poc_style_injection.py \
      --input storage/uploads/xxx.docx \
      --template-id tsinghua_v1 \
      --output /tmp/result.docx

  # 或列出可用模板
  python scripts/poc_style_injection.py --list-templates

核心技术:
  1. 从模板 .docx 的 ZIP 里抽出 word/styles.xml / numbering.xml / theme1.xml
  2. 注入到用户文档 ZIP:
       - styles.xml / theme1.xml 直接替换
       - numbering.xml 合并（保留用户 numId 定义 + 追加模板独有部分）
  3. 用 StructureAnalyzer 识别章节结构
  4. 对识别到的段落: 清除 run 的 direct formatting + 设置 pStyle
  5. 遍历所有 doc.sections 应用页面设置
"""

import argparse
import copy
import shutil
import sys
import time
import zipfile
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm
from lxml import etree

# 把 backend 目录加到 path，以便 import app 模块
_BACKEND_DIR = Path(__file__).parent.parent
sys.path.insert(0, str(_BACKEND_DIR))

from app.core.formatter.structure_analyzer import StructureAnalyzer  # noqa: E402
from app.core.formatter.template_manager import TemplateManager  # noqa: E402


# 要从模板抽取的 XML 部件
TEMPLATE_XML_PARTS = (
    "word/styles.xml",
    "word/numbering.xml",
    "word/theme/theme1.xml",
)

# 章节类型 → Word 样式名候选列表（按优先级）
SECTION_TO_STYLE_CANDIDATES = {
    "title":       ["Title", "标题"],
    "chapter":     ["Heading 1", "标题 1", "标题1", "一级标题"],
    "section_1":   ["Heading 2", "标题 2", "标题2", "二级标题"],
    "section_2":   ["Heading 3", "标题 3", "标题3", "三级标题"],
    "section_3":   ["Heading 4", "标题 4", "标题4", "四级标题"],
    "abstract_cn": ["Abstract", "摘要"],
    "abstract_en": ["Abstract", "摘要"],
    "keywords_cn": ["Normal", "正文"],
    "keywords_en": ["Normal", "正文"],
    "references":  ["References", "参考文献"],
    "toc":         ["TOC 1", "目录 1", "目录"],
    "body":        ["Normal", "正文"],
}

# direct formatting 里要清掉的子节点（打 pStyle 前清一遍，让样式定义生效）
DIRECT_FORMAT_TAGS_TO_CLEAR = {
    "sz", "szCs", "rFonts", "b", "bCs", "i", "iCs", "u", "color",
}


# ─────────────────────────────────────────────────────────────
# ZIP 层操作：抽取模板部件 / 注入到用户文档
# ─────────────────────────────────────────────────────────────

def extract_template_parts(template_docx: Path) -> Dict[str, bytes]:
    """从模板 docx ZIP 里抽出 styles.xml / numbering.xml / theme1.xml 的字节"""
    parts: Dict[str, bytes] = {}
    with zipfile.ZipFile(template_docx, "r") as z:
        names = set(z.namelist())
        for t in TEMPLATE_XML_PARTS:
            if t in names:
                parts[t] = z.read(t)
    return parts


def merge_numbering(user_xml: bytes, template_xml: bytes) -> bytes:
    """
    合并 numbering.xml：
      - 以用户 numbering 为基（保留用户已用的 numId 和 abstractNumId 定义）
      - 追加模板里用户没有的 abstractNum / num 节点
    避免直接替换导致用户原有列表错位。
    """
    user_root = etree.fromstring(user_xml)
    template_root = etree.fromstring(template_xml)

    w_num = qn("w:num")
    w_abs_num = qn("w:abstractNum")
    w_num_id_attr = qn("w:numId")
    w_abs_num_id_attr = qn("w:abstractNumId")

    user_num_ids = {el.get(w_num_id_attr) for el in user_root.findall(w_num)}
    user_abs_ids = {el.get(w_abs_num_id_attr) for el in user_root.findall(w_abs_num)}

    # abstractNum 要在 num 之前，保持顺序
    for abs_num in template_root.findall(w_abs_num):
        if abs_num.get(w_abs_num_id_attr) not in user_abs_ids:
            user_root.append(copy.deepcopy(abs_num))
    for num in template_root.findall(w_num):
        if num.get(w_num_id_attr) not in user_num_ids:
            user_root.append(copy.deepcopy(num))

    return etree.tostring(
        user_root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone=True,
    )


def inject_template_parts(
    user_docx: Path, template_parts: Dict[str, bytes], output_docx: Path
) -> None:
    """注入模板 XML 到用户文档（生成新文件，不改原件）"""
    tmp = output_docx.with_suffix(output_docx.suffix + ".tmp")
    with zipfile.ZipFile(user_docx, "r") as zin:
        existing_names = set(zin.namelist())
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.namelist():
                if item == "word/numbering.xml" and item in template_parts:
                    merged = merge_numbering(zin.read(item), template_parts[item])
                    zout.writestr(item, merged)
                elif item in template_parts:
                    zout.writestr(item, template_parts[item])
                else:
                    zout.writestr(item, zin.read(item))
            # 用户文档没有但模板有的部件（如用户文档没 numbering.xml），追加进去
            for name, data in template_parts.items():
                if name not in existing_names:
                    zout.writestr(name, data)
    shutil.move(str(tmp), str(output_docx))


# ─────────────────────────────────────────────────────────────
# 段落层：清 direct formatting + 打 pStyle 标签
# ─────────────────────────────────────────────────────────────

def clear_direct_formatting(paragraph) -> None:
    """清掉段落所有 run 的 direct formatting（让 pStyle 样式定义生效）"""
    for run in paragraph.runs:
        rPr = run._element.find(qn("w:rPr"))
        if rPr is None:
            continue
        for child in list(rPr):
            if etree.QName(child).localname in DIRECT_FORMAT_TAGS_TO_CLEAR:
                rPr.remove(child)


def resolve_style_name(doc: Document, section_type: str) -> Optional[str]:
    """从候选里找第一个文档真实存在的样式名"""
    candidates = SECTION_TO_STYLE_CANDIDATES.get(section_type, ["Normal"])
    available = {s.name for s in doc.styles}
    for name in candidates:
        if name in available:
            return name
    return None


def tag_sections_with_pstyle(doc: Document, sections: List[Dict]) -> Dict[str, int]:
    """对识别到的章节打 pStyle 标签"""
    stats: Dict[str, int] = {}
    skipped = 0
    for sec in sections:
        idx = sec.get("index")
        stype = sec.get("type")
        if idx is None or idx >= len(doc.paragraphs):
            skipped += 1
            continue
        para = doc.paragraphs[idx]
        style_name = resolve_style_name(doc, stype)
        if style_name is None:
            skipped += 1
            continue
        try:
            clear_direct_formatting(para)
            para.style = doc.styles[style_name]
            stats[stype] = stats.get(stype, 0) + 1
        except Exception as e:
            print(f"  [warn] 段落 {idx}({stype}) 打标签失败: {e}")
            skipped += 1
    if skipped:
        stats["__skipped__"] = skipped
    return stats


# ─────────────────────────────────────────────────────────────
# 页面设置：遍历所有 section
# ─────────────────────────────────────────────────────────────

def apply_page_setup_all(doc: Document, page_cfg: Dict) -> int:
    """所有 doc.sections 都应用页面设置"""
    count = 0
    for section in doc.sections:
        section.page_width = Cm(page_cfg["page_width"])
        section.page_height = Cm(page_cfg["page_height"])
        section.top_margin = Cm(page_cfg["margin_top"])
        section.bottom_margin = Cm(page_cfg["margin_bottom"])
        section.left_margin = Cm(page_cfg["margin_left"])
        section.right_margin = Cm(page_cfg["margin_right"])
        count += 1
    return count


# ─────────────────────────────────────────────────────────────
# 主流程
# ─────────────────────────────────────────────────────────────

def list_templates(template_dir: Path) -> None:
    tm = TemplateManager(
        builtin_dir=template_dir / "builtin",
        user_dir=template_dir / "user",
    )
    templates = tm.list_templates()
    if not templates:
        print("(无模板)")
        return
    print(f"{'ID':<20} {'分类':<15} {'名称'}")
    print("-" * 70)
    for t in templates:
        print(f"{t.template_id:<20} {t.category:<15} {t.name}")


def run_poc(
    input_path: Path,
    template_id: str,
    output_path: Path,
    template_dir: Path,
) -> None:
    print("=" * 70)
    print(f"POC: 样式注入方案")
    print(f"输入: {input_path}")
    print(f"模板: {template_id}")
    print(f"输出: {output_path}")
    print("=" * 70)

    timings: List[tuple] = []
    t0 = time.time()

    # Step 1: 解析模板
    tm = TemplateManager(
        builtin_dir=template_dir / "builtin",
        user_dir=template_dir / "user",
    )
    template_meta = tm.get_template(template_id)
    if template_meta is None:
        print(f"[error] 模板不存在: {template_id}")
        print(f"        可用模板请运行: python {Path(__file__).name} --list-templates")
        sys.exit(1)
    template_docx = Path(template_meta.file_path)
    template_cfg = tm.extract_config(template_id)
    page_cfg = {
        "page_width":    template_cfg.page_width,
        "page_height":   template_cfg.page_height,
        "margin_top":    template_cfg.margin_top,
        "margin_bottom": template_cfg.margin_bottom,
        "margin_left":   template_cfg.margin_left,
        "margin_right":  template_cfg.margin_right,
    }
    t1 = time.time()
    timings.append(("加载模板配置", t1 - t0))
    print(f"[{t1 - t0:.2f}s] 模板加载: {template_meta.name}")
    print(f"         页面 {page_cfg['page_width']:.1f}×{page_cfg['page_height']:.1f}cm, "
          f"边距 T{page_cfg['margin_top']:.1f}/B{page_cfg['margin_bottom']:.1f}/"
          f"L{page_cfg['margin_left']:.1f}/R{page_cfg['margin_right']:.1f}cm")

    # Step 2: 抽取模板 XML 部件
    template_parts = extract_template_parts(template_docx)
    t2 = time.time()
    timings.append(("抽取模板 XML", t2 - t1))
    print(f"[{t2 - t1:.2f}s] 模板 XML 部件: {list(template_parts.keys())}")

    # Step 3: 注入到用户文档（ZIP 层操作）
    output_path.parent.mkdir(parents=True, exist_ok=True)
    inject_template_parts(input_path, template_parts, output_path)
    t3 = time.time()
    timings.append(("注入样式定义", t3 - t2))
    print(f"[{t3 - t2:.2f}s] 注入样式定义完成")

    # Step 4: 加载注入后文档
    doc = Document(str(output_path))
    para_count = len(doc.paragraphs)
    t4 = time.time()
    timings.append(("加载注入后文档", t4 - t3))
    print(f"[{t4 - t3:.2f}s] 加载文档: {para_count} 段落, "
          f"{len(doc.sections)} sections, {len(doc.styles)} 样式")

    # Step 5: 结构识别
    analyzer = StructureAnalyzer(use_ai=False)
    analysis = analyzer.analyze(doc)
    sections = analysis["sections"]
    t5 = time.time()
    timings.append(("结构识别", t5 - t4))
    print(f"[{t5 - t4:.2f}s] 结构识别: {len(sections)} 章节, 质量 {analysis['quality']:.2%}")

    # Step 6: 段落打 pStyle 标签 + 清 direct formatting
    tag_stats = tag_sections_with_pstyle(doc, sections)
    t6 = time.time()
    timings.append(("pStyle 打标签", t6 - t5))
    print(f"[{t6 - t5:.2f}s] pStyle 标签: {tag_stats}")

    # Step 7: 所有 section 页面设置
    n_sec = apply_page_setup_all(doc, page_cfg)
    t7 = time.time()
    timings.append(("页面设置", t7 - t6))
    print(f"[{t7 - t6:.2f}s] 页面设置: {n_sec} 个 section")

    # Step 8: 保存
    doc.save(str(output_path))
    t8 = time.time()
    timings.append(("保存文档", t8 - t7))
    print(f"[{t8 - t7:.2f}s] 保存完成")

    # 总结
    total = t8 - t0
    size_in = input_path.stat().st_size / 1024
    size_out = output_path.stat().st_size / 1024
    print("=" * 70)
    print(f"✅ 总耗时: {total:.2f}s   |   输入 {size_in:.1f} KB → 输出 {size_out:.1f} KB")
    print("   阶段耗时明细:")
    for name, dur in timings:
        pct = dur / total * 100 if total > 0 else 0
        print(f"     {name:<20} {dur:>6.2f}s  ({pct:>4.1f}%)")
    print("=" * 70)
    print("\n[人工验证清单]")
    print("  [ ] 用 Microsoft Word 打开输出文件，确认无'需要修复'提示")
    print("  [ ] 抽查 5 个章节标题：字体、字号、加粗、居中是否符合模板")
    print("  [ ] 抽查 5 个正文段落：字体、字号、首行缩进、行距是否符合模板")
    print("  [ ] 检查所有页的页边距（含封面、目录、正文）是否一致")
    print("  [ ] 检查用户原文档里的加粗/斜体/颜色等内联格式是否保留")
    print("  [ ] 检查图片、表格、公式是否完整保留")


def main():
    parser = argparse.ArgumentParser(
        description="POC: 样式注入方案验证",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("--input", help="用户论文 .docx 路径")
    parser.add_argument("--template-id", help="模板 ID（如 tsinghua_v1）")
    parser.add_argument("--output", help="输出 .docx 路径")
    parser.add_argument(
        "--template-dir",
        default=str(_BACKEND_DIR / "templates"),
        help="模板根目录（默认 backend/templates）",
    )
    parser.add_argument(
        "--list-templates",
        action="store_true",
        help="列出可用模板并退出",
    )
    args = parser.parse_args()

    template_dir = Path(args.template_dir)

    if args.list_templates:
        list_templates(template_dir)
        return

    if not (args.input and args.template_id and args.output):
        parser.error("--input / --template-id / --output 三者都必填")

    input_path = Path(args.input)
    if not input_path.exists():
        parser.error(f"输入文件不存在: {input_path}")

    run_poc(
        input_path=input_path,
        template_id=args.template_id,
        output_path=Path(args.output),
        template_dir=template_dir,
    )


if __name__ == "__main__":
    main()
