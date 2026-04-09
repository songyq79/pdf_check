"""
智能评价 API
POST /api/v1/evaluation/upload        → 上传论文，提交评价任务
GET  /api/v1/evaluation/status/{id}   → 查询进度
GET  /api/v1/evaluation/result/{id}   → 获取完整结果
GET  /api/v1/evaluation/download/{id} → 下载 Word 报告
"""

import asyncio
import re
import uuid
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime

from docx import Document
from fastapi import APIRouter, File, HTTPException, UploadFile
from fastapi.responses import FileResponse
from loguru import logger

from app.config import settings
from app.services.file_service import save_upload_file, validate_file
from app.workers.celery_app import celery_app
from app.workers.evaluation_tasks import run_evaluation

router = APIRouter()

# ─────────────────────────────────────────────────────────────
# 正则常量（模块级编译）
# ─────────────────────────────────────────────────────────────

# 参考文献行：支持三种格式：
# 1. [1] 编号格式
# 2. 1. 中文/英文 格式
# 3. 无编号但含 [J][D][M][C][R] 文献类型标识符（参考文献区节标题之后）
_REF_LINE_PAT = re.compile(r'^\[\d+\]|^\d+\.\s+[\u4e00-\u9fff A-Za-z]')
# 无编号参考文献行（需在参考文献区才生效）
_REF_LINE_NOID_PAT = re.compile(r'\[J\]|\[D\]|\[M\]|\[C\]|\[R\]|\[P\]|\[S\]')

# 参考文献区节标题
_REF_SECTION_PAT = re.compile(r'^参考文献\s*$', re.IGNORECASE)

# 正文引用标注：[1] [1,2] [1-3]
_CITATION_PAT = re.compile(r'\[\d+(?:[,，\-–]\d+)*\]')

# 摘要标题行
_ABSTRACT_PAT = re.compile(r'^(摘\s*要|abstract|摘　要)$', re.IGNORECASE)

# 关键词行
_KEYWORD_PAT = re.compile(r'^(关键词|keywords)[：:\s]*(.*)', re.IGNORECASE)

# 章节标题识别：只匹配真正的章节编号，避免把文献年份等误识别为标题
# 规则：必须是行首的编号格式，后面紧跟中文或空格+中文
_HEADING_PAT = re.compile(
    r'^('
    r'第[一二三四五六七八九十百\d]+[章节部分]\s*\S'     # 第一章xxx
    r'|\d+(\.\d+)*\s+[\u4e00-\u9fff]'                   # 1 绪论 / 1.1 背景（数字+空格+中文）
    r'|[一二三四五六七八九十]+[、]\s*[\u4e00-\u9fff]'    # 一、研究背景
    r')',
)

# 封面/扉页特征行（不应被识别为标题）
_COVER_PAT = re.compile(
    r'^(中国|北京|上海|南京|武汉|西安|成都|广州|哈尔滨|大连|天津)'
    r'|大学|学院|高等学历|继续教育|毕业设计|毕业论文|学位论文'
    r'|姓\s*名[：:]|学\s*号[：:]|专\s*业[：:]|指导教师[：:]|院\s*系[：:]'
    r'|题\s*目[：:]|日\s*期[：:]|\d{4}\s*年\s*\d{1,2}\s*月'
)

# 论文题目行（封面上的题目，通常以"题目："开头）
_TITLE_LINE_PAT = re.compile(r'^题\s*目[：:]\s*(.*)')


def _is_heading(text: str) -> bool:
    """判断是否为章节标题，同时排除封面行和段落介绍行"""
    stripped = text.strip()
    if len(stripped) > 40:           # 标题一般不超过40字
        return False
    if '。' in stripped or '，' in stripped:  # 含句号/逗号说明是正文段落不是标题
        return False
    if _COVER_PAT.search(stripped):  # 封面特征行排除
        return False
    return bool(_HEADING_PAT.match(stripped))


def _extract_text_fallback(file_path: str) -> str:
    """
    降级方案：从损坏的文档中提取纯文本
    
    尝试从 ZIP 中解析 document.xml 提取文本内容
    """
    try:
        with zipfile.ZipFile(file_path, 'r') as docx_zip:
            # 读取主文档 XML
            xml_content = docx_zip.read('word/document.xml')
            root = ET.fromstring(xml_content)
            
            # 提取所有文本节点
            namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            text_elements = root.findall('.//w:t', namespace)
            
            text_content = '\n'.join([elem.text for elem in text_elements if elem.text])
            
            if text_content.strip():
                logger.info(f"[evaluation] 降级提取成功: {len(text_content)} 字符")
                return text_content
                
    except Exception as e:
        logger.warning(f"[evaluation] 降级提取失败: {e}")
    
    return ""


def _extract_structure_from_text(text_content: str) -> dict:
    """
    降级方案：从纯文本中提取结构化内容
    
    只能提取基本的文本信息，无法识别复杂的文档结构
    """
    lines = [line.strip() for line in text_content.split('\n') if line.strip()]
    
    if not lines:
        return {
            "title": "未命名论文",
            "abstract": "（未提取到摘要）",
            "keywords": "（未提取到关键词）",
            "outline": "（未识别到章节结构）",
            "conclusion": "（未提取到结论）",
            "references": "（未提取到参考文献）",
            "chapter_summaries": "（未提取到章节内容）",
            "language_sample": "（未提取到正文内容）",
            "citation_samples": "（正文中未检测到 [数字] 格式的引用标注）",
        }
    
    # 提取标题（第一个足够长的行）
    title = "未命名论文"
    for line in lines[:20]:
        if len(line) >= 8 and not _COVER_PAT.search(line):
            title = line[:120]
            break
    
    # 提取摘要
    abstract_lines = []
    in_abstract = False
    for i, line in enumerate(lines):
        if _ABSTRACT_PAT.match(line):
            in_abstract = True
            continue
        if in_abstract:
            if _KEYWORD_PAT.match(line) or _is_heading(line):
                break
            abstract_lines.append(line)
            if len('\n'.join(abstract_lines)) > 800:
                break
    abstract = '\n'.join(abstract_lines)[:800] or "（未提取到摘要）"
    
    # 提取关键词
    keywords = ""
    for line in lines:
        kw = _KEYWORD_PAT.match(line)
        if kw:
            keywords = kw.group(2).strip()
            break
    if not keywords:
        keywords = "（未提取到关键词）"
    
    # 提取参考文献
    ref_lines = []
    in_ref = False
    for line in lines:
        if _REF_SECTION_PAT.match(line):
            in_ref = True
            continue
        if in_ref:
            if _REF_LINE_PAT.match(line) or _REF_LINE_NOID_PAT.search(line):
                ref_lines.append(line)
                if len(ref_lines) >= 60:
                    break
    references = '\n'.join(ref_lines) or "（未提取到参考文献）"
    
    # 提取章节标题（大纲）
    outline_lines = []
    for line in lines:
        if _is_heading(line):
            outline_lines.append(line)
            if len(outline_lines) >= 20:
                break
    outline = '\n'.join(outline_lines) or "（未识别到章节结构）"
    
    # 提取正文样本（跳过前面的封面、摘要等，取中间部分）
    body_lines = []
    skip_count = 0
    for line in lines:
        # 跳过封面、摘要、关键词
        if _COVER_PAT.search(line) or _ABSTRACT_PAT.match(line) or _KEYWORD_PAT.match(line):
            skip_count += 1
            continue
        # 跳过参考文献
        if _REF_SECTION_PAT.match(line):
            break
        # 跳过标题行
        if _is_heading(line):
            continue
        # 收集正文
        if skip_count > 0 and len(line) > 20:
            body_lines.append(line)
            if len('\n'.join(body_lines)) > 3000:
                break
    
    language_sample = '\n'.join(body_lines[:30]) or "（未提取到正文内容）"
    chapter_summaries = '\n'.join(body_lines[:50]) or "（未提取到章节内容）"
    
    # 提取结论（最后几段正文）
    conclusion_lines = []
    for line in reversed(lines):
        if _REF_SECTION_PAT.match(line) or _REF_LINE_PAT.match(line):
            continue
        if len(line) > 20 and not _is_heading(line):
            conclusion_lines.insert(0, line)
            if len('\n'.join(conclusion_lines)) > 1500:
                break
    conclusion = '\n'.join(conclusion_lines[-10:]) or "（未提取到结论）"
    
    # 提取引用标注样本
    citation_lines = [line for line in body_lines if _CITATION_PAT.search(line)]
    citation_samples = '\n'.join(citation_lines[:20]) or "（正文中未检测到 [数字] 格式的引用标注）"
    
    return {
        "title": title,
        "abstract": abstract,
        "keywords": keywords,
        "outline": outline,
        "conclusion": conclusion,
        "references": references,
        "chapter_summaries": chapter_summaries,
        "language_sample": language_sample,
        "citation_samples": citation_samples,
    }


def _extract_title(doc: Document) -> str:
    """
    按优先级提取论文标题：
    1. Word 文档属性的 title（作者填写过才有）
    2. 封面"题目：xxx"行
    3. 样式为 Heading 1 / Title 的第一个段落
    4. 第一个较长的非封面非空行（≥8字）
    """
    # 优先级1：文档属性
    core_props = doc.core_properties
    if core_props.title and len(core_props.title.strip()) > 4:
        return core_props.title.strip()[:120]

    # 优先级2：封面题目行
    for p in doc.paragraphs[:30]:
        text = p.text.strip()
        m = _TITLE_LINE_PAT.match(text)
        if m and m.group(1).strip():
            return m.group(1).strip()[:120]

    # 优先级3：Heading 1 / Title 样式
    for p in doc.paragraphs[:50]:
        style_name = (p.style.name or "").lower()
        if any(s in style_name for s in ("heading 1", "title", "标题 1", "标题1")):
            text = p.text.strip()
            if text and not _COVER_PAT.search(text):
                return text[:120]

    # 优先级4：第一个足够长的正文行（跳过封面）
    for p in doc.paragraphs[:40]:
        text = p.text.strip()
        if len(text) >= 8 and not _COVER_PAT.search(text) and not _ABSTRACT_PAT.match(text):
            return text[:120]

    return "未命名论文"


def _extract_structure(doc: Document) -> dict:
    """
    从 docx 中提取结构化内容，供各评价维度按需取用。

    返回字段：
      title            论文标题
      abstract         摘要（≤800字）
      keywords         关键词
      outline          章节大纲文本（仅一级/二级标题）
      conclusion       结论章节全文（≤1500字）
      references       参考文献全文（最多60条）
      chapter_summaries 各章节前300字摘要拼接（≤10章）
      language_sample  从核心正文章节抽取的语言样本（≤3000字）
      citation_samples 含引用标注的正文句子抽样（≤2000字）

    注意：chapters 列表不返回（数据量大，不传给 Celery Worker）
    """
    # ── 提取标题 ──────────────────────────────────────────────
    title = _extract_title(doc)

    # ── 一次遍历，按状态机分区 ───────────────────────────────
    abstract_lines = []
    keywords = ""
    ref_lines = []
    chapters = []        # [{"heading": str, "lines": [str], "level": int}]
    cur_chapter = None
    state = "preamble"   # preamble → abstract → body → references

    for p in doc.paragraphs:
        line = p.text.strip()
        if not line:
            continue

        # ── 进入参考文献区 ──────────────────────────────────
        if state != "references":
            if _REF_SECTION_PAT.match(line):
                state = "references"
                continue
            # 没有"参考文献"标题、直接以编号开头的文献行也切换
            if _REF_LINE_PAT.match(line) and state == "body":
                state = "references"
                ref_lines.append(line)
                continue

        if state == "references":
            # 支持有编号格式 [1]/1. 和无编号格式（含[J][D]等标识符）
            if _REF_LINE_PAT.match(line) or _REF_LINE_NOID_PAT.search(line):
                ref_lines.append(line)
            # 参考文献区的非文献行（如致谢、空行）忽略
            continue

        # ── 摘要区 ──────────────────────────────────────────
        if _ABSTRACT_PAT.match(line):
            state = "abstract"
            continue

        if state == "abstract":
            kw = _KEYWORD_PAT.match(line)
            if kw:
                keywords = kw.group(2).strip()
                state = "body"
                continue
            if _is_heading(line):
                state = "body"
                # 不 continue，让下面的 body 逻辑处理这个标题
            else:
                abstract_lines.append(line)
                continue

        # ── 关键词单独行（摘要之后、正文之前）──────────────
        kw = _KEYWORD_PAT.match(line)
        if kw and not keywords:
            keywords = kw.group(2).strip()
            state = "body"
            continue

        # ── 正文区 ──────────────────────────────────────────
        if state in ("preamble", "body"):
            if _is_heading(line):
                state = "body"
                # 判断标题级别：含小数点且有两段以上为二级，否则一级
                # 级别判断：三段数字(1.1.1)=3级，两段(1.1)=2级，其余=1级
                if re.match(r'^\d+\.\d+\.\d+', line):
                    level = 3
                elif re.match(r'^\d+\.\d+', line) or re.match(r'^第[^章节部分]*[节]', line):
                    level = 2
                else:
                    level = 1
                cur_chapter = {"heading": line, "lines": [], "level": level}
                chapters.append(cur_chapter)
            else:
                if cur_chapter is not None and state == "body":
                    cur_chapter["lines"].append(line)

    # ── 整理各输出字段 ───────────────────────────────────────

    abstract = "\n".join(abstract_lines)[:800]

    # 大纲：一级标题直接列出，二级缩进，三级及以下忽略（避免太细）
    outline_parts = []
    for ch in chapters:
        if ch["level"] == 1:
            outline_parts.append(ch["heading"])
        elif ch["level"] == 2:
            outline_parts.append(f"  └ {ch['heading']}")
        # 三级及以下不纳入大纲
    outline = "\n".join(outline_parts) if outline_parts else "（未识别到章节结构）"

    # ── 给每个一级章节聚合其所有子章节内容 ─────────────────────
    # 背景：一级标题（如"第3章"）自身 lines 通常为空，
    # 内容全在子章节（3.1、3.2...）的 lines 里。
    # 构建 parent_bodies: {一级章节index: 合并后的全部正文}
    lv1_indices = [i for i, ch in enumerate(chapters) if ch["level"] == 1]
    def _get_chapter_full_body(ch_idx: int) -> str:
        """获取一个一级章节及其所有子章节的合并正文"""
        # 找下一个一级章节的位置
        next_lv1 = len(chapters)
        for idx in lv1_indices:
            if idx > ch_idx:
                next_lv1 = idx
                break
        all_lines = []
        for i in range(ch_idx, next_lv1):
            all_lines.extend(chapters[i]["lines"])
        return "\n".join(all_lines)

    # 结论：找含"结论/总结/结语"关键词的最后一个一级章节，包含其子章节内容
    conclusion = ""
    for lv1_idx in reversed(lv1_indices):
        ch = chapters[lv1_idx]
        if re.search(r'结论|总结|结语|conclusions?', ch["heading"], re.IGNORECASE):
            full_body = _get_chapter_full_body(lv1_idx)
            conclusion = (ch["heading"] + "\n" + full_body)[:1500]
            break
    if not conclusion and lv1_indices:
        last_idx = lv1_indices[-1]
        ch = chapters[last_idx]
        full_body = _get_chapter_full_body(last_idx)
        conclusion = (ch["heading"] + "\n" + full_body)[:1500]

    # 参考文献（最多60条）
    references = "\n".join(ref_lines[:60])

    # 章节摘要：每个一级章节取聚合正文前400字，跳过文献综述
    chapter_summary_parts = []
    for lv1_idx in lv1_indices:
        ch = chapters[lv1_idx]
        if re.search(r'文献综述|研究综述|文献回顾|literature', ch["heading"], re.IGNORECASE):
            continue
        full_body = _get_chapter_full_body(lv1_idx)
        if full_body.strip():
            chapter_summary_parts.append(f"【{ch['heading']}】\n{full_body[:400]}")
        if len(chapter_summary_parts) >= 10:
            break
    chapter_summaries = "\n\n".join(chapter_summary_parts)

    # 语言样本：从核心正文一级章节均匀抽取（排除绪论、文献综述、结论、致谢）
    # 每章取500字，最多6章共3000字
    skip_pat = re.compile(r'绪论|前言|引言|文献综述|文献回顾|致谢|附录|结论|总结|结语|literature|references', re.IGNORECASE)
    core_lv1_indices = [i for i in lv1_indices if not skip_pat.search(chapters[i]["heading"])]
    
    # 如果核心章节为空，则使用所有一级章节（避免返回空内容）
    if not core_lv1_indices:
        core_lv1_indices = lv1_indices
    
    step = max(1, len(core_lv1_indices) // 6)
    lang_parts = []
    for lv1_idx in core_lv1_indices[::step][:6]:
        ch = chapters[lv1_idx]
        full_body = _get_chapter_full_body(lv1_idx)
        if full_body.strip():
            lang_parts.append(f"【{ch['heading']}】\n{full_body[:500]}")
    language_sample = "\n\n".join(lang_parts)

    # 引用标注抽样：从所有章节找含 [数字] 的句子，每章最多3句，共≤2000字
    citation_sents = []
    total_len = 0
    for ch in chapters:
        ch_sents = [l for l in ch["lines"] if _CITATION_PAT.search(l) and len(l) > 15]
        for sent in ch_sents[:3]:
            citation_sents.append(sent)
            total_len += len(sent)
            if total_len > 2000:
                break
        if total_len > 2000:
            break
    citation_samples = "\n".join(citation_sents) or "（正文中未检测到 [数字] 格式的引用标注）"

    return {
        "title":             title,
        "abstract":          abstract or "（未提取到摘要）",
        "keywords":          keywords or "（未提取到关键词）",
        "outline":           outline,
        "conclusion":        conclusion or "（未提取到结论）",
        "references":        references or "（未提取到参考文献）",
        "chapter_summaries": chapter_summaries or "（未提取到章节内容）",
        "language_sample":   language_sample or "（未提取到正文内容）",
        "citation_samples":  citation_samples,
    }


# ─────────────────────────────────────────────────────────────
# API 路由
# ─────────────────────────────────────────────────────────────

@router.post("/upload", summary="上传论文并提交评价任务")
async def evaluate_paper(file: UploadFile = File(...)):
    """上传论文 → 提交 Celery 评价任务 → 立即返回 task_id"""
    if not settings.BAILIAN_API_KEY and not settings.DEEPSEEK_API_KEY:
        raise HTTPException(503, "AI 评价服务未配置：请在 .env 中配置 BAILIAN_API_KEY 或 DEEPSEEK_API_KEY")

    validate_file(file, max_size_mb=settings.MAX_FILE_SIZE)

    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(400, "仅支持 .docx 格式")

    file_path = await save_upload_file(file, settings.UPLOAD_PATH)
    logger.info(f"[evaluation] 文件已保存: {file_path}")

    paper_struct = None
    fallback_mode = False
    
    try:
        doc          = await asyncio.to_thread(Document, str(file_path))
        paper_struct = await asyncio.to_thread(_extract_structure, doc)
    except Exception as exc:
        error_msg = str(exc)
        logger.error(f"[evaluation] 文档解析失败: {error_msg}")
        
        # 判断是否是文档格式问题
        if "no item named" in error_msg.lower() or "badzip" in error_msg.lower():
            logger.warning(f"[evaluation] 检测到文档格式异常，尝试降级处理...")
            
            try:
                # 尝试降级处理：提取文本 + 简化结构提取
                text_content = await asyncio.to_thread(_extract_text_fallback, str(file_path))
                
                if text_content:
                    paper_struct = await asyncio.to_thread(_extract_structure_from_text, text_content)
                    fallback_mode = True
                    logger.warning(f"[evaluation] 降级处理成功，已提取 {len(text_content)} 字符")
                else:
                    # 文本提取也失败
                    logger.error(f"[evaluation] 降级处理失败：无法提取文本内容")
                    raise HTTPException(
                        400, 
                        "文档格式异常且无法自动修复。建议：在 Word 中打开文档，选择'另存为'，保存为新的 .docx 文件后重试。"
                    )
                    
            except HTTPException:
                raise
            except Exception as fallback_error:
                logger.error(f"[evaluation] 降级处理失败: {fallback_error}")
                raise HTTPException(
                    400, 
                    "文档格式异常且无法自动修复。建议：在 Word 中打开文档，选择'另存为'，保存为新的 .docx 文件后重试。"
                )
        else:
            raise HTTPException(400, f"文档解析失败: {error_msg}")

    title = paper_struct["title"]
    logger.info(
        f"[evaluation] 结构提取完成: title={title[:40]!r} "
        f"refs={paper_struct['references'].count(chr(10))+1}条 "
        f"lang_sample={len(paper_struct['language_sample'])}字"
        f"{' (降级模式)' if fallback_mode else ''}"
    )

    task_id = str(uuid.uuid4())
    celery_result = run_evaluation.apply_async(
        args=[task_id, paper_struct, str(settings.OUTPUT_PATH)],
        task_id=task_id,
    )

    logger.info(f"[evaluation] 任务已入队 task_id={celery_result.id} title={title}")

    response = {
        "task_id":     celery_result.id,
        "status":      "pending",
        "paper_title": title,
        "status_url":  f"/api/v1/evaluation/status/{celery_result.id}",
        "result_url":  f"/api/v1/evaluation/result/{celery_result.id}",
        "message":     "任务已提交，多维度评价正在后台并发执行",
    }
    
    # 如果使用了降级模式，添加警告信息
    if fallback_mode:
        response["warning"] = "文档格式异常，已使用简化模式处理（评价结果可能受影响）"
    
    return response


@router.get("/status/{task_id}", summary="查询评价进度")
def get_status(task_id: str):
    result = celery_app.AsyncResult(task_id)
    state  = result.state

    state_map = {
        "PENDING": ("pending",    0),
        "STARTED": ("processing", 40),
        "RETRY":   ("processing", 20),
        "SUCCESS": ("completed",  100),
        "FAILURE": ("failed",     0),
    }
    status, progress = state_map.get(state, ("pending", 0))

    resp = {"task_id": task_id, "status": status, "progress": progress}
    if state == "FAILURE":
        resp["error"] = str(result.result)
    return resp


@router.get("/result/{task_id}", summary="获取评价结果")
def get_result(task_id: str):
    result = celery_app.AsyncResult(task_id)

    if result.state == "PENDING":
        raise HTTPException(202, "任务排队中，请稍候")
    if result.state in ("STARTED", "RETRY"):
        raise HTTPException(202, "评价进行中，请稍候")
    if result.state == "FAILURE":
        raise HTTPException(503, f"评价失败: {result.result}")
    if result.state != "SUCCESS":
        raise HTTPException(400, f"未知状态: {result.state}")

    return result.result


@router.get("/download/{report_id}", summary="下载评价报告")
def download_report(report_id: str):
    report_path = settings.OUTPUT_PATH / f"{report_id}_report.docx"
    if not report_path.exists():
        raise HTTPException(404, "报告文件不存在或已过期")

    return FileResponse(
        path=str(report_path),
        filename=f"论文评价报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
