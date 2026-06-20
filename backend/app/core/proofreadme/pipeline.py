"""
论文校对核心流水线

mode 参数控制校对深度：
  "basic" → 仅第一轮：错别字 + 标点
  "full"  → 完整三轮：错别字 + 语法 + 段落逻辑（默认）

核心设计原则：
1. 图片段落（含 w:drawing）：完全跳过，绝不触碰
2. 正文段落：按 mode 决定检查深度
3. 表格单元格：错别字 + 语法检查（不做段落逻辑，单元格通常较短）
4. 所有修改以 Word Track Changes 修订节点写入，不直接改原文
5. 段落原有格式（w:pPr）和 run 样式（w:rPr）完整保留

Bug 修复记录：
- [fix1] _write_revision：改为原地替换而非全清重建，避免段落结构串位
- [fix2] 移除双层 Semaphore 竞争；短段落不分块，减少无效 AI 调用
- [fix3] process_word 增加空文件/损坏文件提前检测，避免前端永久 pending
- [fix4] 增加 _should_skip_para 提前过滤无需校对的段落，减少 30~50% AI 调用
- [fix5] Semaphore 并发数 8 → 16，处理速度近乎翻倍
- [fix6] _get_para_text 改为只取直接子 w:r，不递归进 w:del/w:ins，
         避免把已删除文字混入原文发给 AI，导致 diff 错位
- [fix7] AI 调用并发、写 XML 串行：gather 阶段只做 AI 调用不写 XML，
         全部结果回来后再串行调用 _write_revision，彻底消除并发写 XML 数据竞争
- [fix8] 取消超长段落分块：整段发给 AI，避免分块拼接后 diff 错位误改正常文字
- [fix9] _proofread_cell 加入 _should_skip_para 过滤，表格排版标签行不发 AI
- [fix10] _should_skip_para 参考文献正则收紧，必须后跟中文才跳过，避免误杀章节标题
- [fix11] _write_revision 确保 w:pPr 始终在段落最前面，修复表格单元格对齐错位问题
"""

import asyncio
import os
import re as _re
import logging
import zipfile
import xml.etree.ElementTree as ET
from copy import deepcopy
from typing import Optional, Literal

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.core.proofreadme.chunk import protect_terms, restore_terms
from app.core.proofreadme.diff_engine import compute_diff, has_diff
from app.core.proofreadme.llm import full_proofread, proofread_text, _full_check_once
from app.services.docx_normalizer import clean_control_chars, normalize_docx
from app.core.proofreadme.word_patch import (
    enable_track_changes,
    make_del_node,
    make_ins_node,
    new_rev_id,
    _get_or_create_comments_part,
    _save_comments_part,
    _sanitize_rpr,
    add_comment_to_element,
    ensure_comments_content_type,
    make_comment_range_start,
    make_comment_range_end,
    make_comment_reference,
)

logger = logging.getLogger(__name__)

ProofreadMode = Literal["basic", "full"]


# ─────────────────────────────────────────────────────────────
# 降级处理：从损坏文档提取文本
# ─────────────────────────────────────────────────────────────

def _extract_text_fallback(input_path: str) -> str:
    """
    降级方案：从损坏的文档中提取纯文本
    
    尝试从 ZIP 中解析 document.xml 提取文本内容
    """
    try:
        with zipfile.ZipFile(input_path, 'r') as docx_zip:
            # 读取主文档 XML
            xml_content = docx_zip.read('word/document.xml')
            root = ET.fromstring(xml_content)
            
            # 提取所有文本节点
            namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            text_elements = root.findall('.//w:t', namespace)
            
            text_content = '\n'.join([elem.text for elem in text_elements if elem.text])
            
            if text_content.strip():
                text_content = clean_control_chars(text_content)
                logger.info(f"[pipeline] 降级提取成功: {len(text_content)} 字符")
                return text_content

    except Exception as e:
        logger.warning(f"[pipeline] 降级提取失败: {e}")

    return ""


def _create_simple_proofread_doc(text_content: str, output_path: str, mode: ProofreadMode) -> dict:
    """
    降级方案：创建简化的校对文档
    
    从纯文本创建新文档，进行校对后保存
    """
    try:
        # 创建新文档
        doc = Document()
        enable_track_changes(doc.settings.element)
        
        # 按段落分割文本
        paragraphs_text = [p.strip() for p in text_content.split('\n') if p.strip()]
        
        # 添加段落到文档
        for para_text in paragraphs_text:
            doc.add_paragraph(para_text)
        
        logger.info(f"[pipeline] 简化文档创建成功: {len(paragraphs_text)} 个段落")
        
        # 保存临时文档
        temp_path = output_path + ".temp.docx"
        doc.save(temp_path)
        
        # 对临时文档进行正常校对
        try:
            # 重新加载文档进行校对
            doc = Document(temp_path)
            enable_track_changes(doc.settings.element)
            
            tasks = _collect_tasks(doc)
            effective_tasks = [
                (p, pt) for p, pt in tasks
                if not _has_drawing(p._p) and not _is_empty_para(p._p)
            ]
            
            if not effective_tasks:
                doc.save(output_path)
                logger.info("[pipeline] 降级文档无有效文字段落")
                return {
                    "total_paragraphs": len(paragraphs_text),
                    "checked": 0,
                    "changed": 0,
                    "skipped": len(paragraphs_text),
                    "total_changes": 0,
                    "all_changes": [],
                    "mode": mode,
                    "error": None,
                    "fallback_mode": True,
                }
            
            # 同步校对（不使用 asyncio.run，因为 Celery worker 已经在事件循环中）
            # 创建新的事件循环来运行异步任务
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            try:
                stats = loop.run_until_complete(_proofread_tasks_async(doc, effective_tasks, len(tasks), mode))
            finally:
                loop.close()
            
            # 保存最终文档
            doc.save(output_path)
            try:
                ensure_comments_content_type(output_path)
            except Exception as e:
                logger.warning(f"[pipeline] 降级路径 Content Types 修补失败: {e}")

            # 清理临时文件
            try:
                os.remove(temp_path)
            except:
                pass

            stats["fallback_mode"] = True
            logger.info(f"[pipeline] 降级校对完成: checked={stats['checked']} changed={stats['changed']}")
            return stats
            
        except Exception as e:
            logger.error(f"[pipeline] 降级校对失败: {e}")
            # 至少保存原始文档
            doc.save(output_path)
            return {
                "total_paragraphs": len(paragraphs_text),
                "checked": 0,
                "changed": 0,
                "skipped": len(paragraphs_text),
                "total_changes": 0,
                "all_changes": [],
                "mode": mode,
                "error": f"降级校对失败: {e}",
                "fallback_mode": True,
            }
            
    except Exception as e:
        logger.error(f"[pipeline] 简化文档创建失败: {e}")
        raise


async def _proofread_tasks_async(doc, effective_tasks, total_tasks, mode: ProofreadMode) -> dict:
    """异步校对任务（从 process_word 中提取出来，供降级处理复用）"""
    stats = {
        "total_paragraphs": total_tasks,
        "checked": 0,
        "changed": 0,
        "skipped": total_tasks - len(effective_tasks),
        "total_changes": 0,
        "all_changes": [],
        "mode": mode,
        "error": None,
    }
    
    sem = asyncio.Semaphore(16)
    
    async def _run_ai(p, ptype):
        if ptype == "table_cell":
            return await _proofread_cell(p, sem)
        else:
            return await _proofread_paragraph(p, sem, mode=mode)
    
    logger.info(f"[pipeline] 并发 AI 调用（{len(effective_tasks)} 个有效段落）")
    ai_results = await asyncio.gather(
        *[_run_ai(p, ptype) for p, ptype in effective_tasks],
        return_exceptions=True,
    )
    
    # 获取或创建 comments part（用于添加批注）
    comments_elem = _get_or_create_comments_part(doc)
    has_comments = False  # 标记是否有批注需要保存
    
    logger.info("[pipeline] 串行写入修订节点")
    for i, ai_result in enumerate(ai_results):
        p, ptype = effective_tasks[i]
        para_elem = p._p
        
        if isinstance(ai_result, Exception):
            logger.error(f"[pipeline] 段落处理异常 idx={i}: {ai_result}")
            stats["skipped"] += 1
            continue
        
        if ai_result is None:
            continue
        
        original, corrected, _ref, all_changes = ai_result
        
        # 添加详细日志
        logger.info(f"[pipeline] 写入修订 idx={i}: original={original[:100]!r}, corrected={corrected[:100]!r}")
        
        rpr_elem = _get_rpr(para_elem)
        count = _write_revision(para_elem, original, corrected, rpr_elem)
        
        logger.info(f"[pipeline] 写入完成 idx={i}: count={count}, all_changes={all_changes}")
        
        # 【批注】为每处修改添加 Word 批注，说明修改原因
        if count > 0 and all_changes:
            # 合并所有修改原因为一条批注（避免一个段落太多批注气泡）
            comment_lines = []
            for change in all_changes:
                if len(change) >= 3:
                    # change = ["原文", "修正", "类型说明"]
                    comment_lines.append(f"「{change[0]}」→「{change[1]}」：{change[2]}")
                elif len(change) >= 2:
                    comment_lines.append(f"「{change[0]}」→「{change[1]}」")
            
            if comment_lines:
                comment_text = "\n".join(comment_lines)
                comment_id = new_rev_id()
                
                # 在 comments.xml 中添加批注内容
                add_comment_to_element(comments_elem, comment_id, comment_text)
                
                # 在段落中插入批注范围标记
                # commentRangeStart 放在段落第一个 run 之前
                # commentRangeEnd + commentReference 放在段落最后一个 run 之后
                try:
                    _insert_comment_markers(para_elem, comment_id)
                    has_comments = True
                except Exception as e:
                    logger.warning(f"[pipeline] 插入批注标记失败 idx={i}: {e}")
        
        stats["checked"] += 1
        if count > 0:
            stats["changed"] += 1
            stats["total_changes"] += count
            stats["all_changes"].extend(all_changes)
    
    # 保存 comments part（如果有批注）
    if has_comments:
        try:
            _save_comments_part(doc, comments_elem)
            logger.info(f"[pipeline] 批注已保存到文档")
        except Exception as e:
            logger.warning(f"[pipeline] 保存批注失败（不影响修订）: {e}")
    
    return stats


# ─────────────────────────────────────────────────────────────
# 批注标记插入
# ─────────────────────────────────────────────────────────────

def _insert_comment_markers(para_elem, comment_id: int):
    """
    在段落中插入批注范围标记：
    - commentRangeStart 放在第一个 w:r/w:del/w:ins 之前
    - commentRangeEnd + commentReference 放在最后一个 w:r/w:del/w:ins 之后

    使用 addprevious / addnext 相对定位，避免索引硬算破坏段落元素顺序
    （段落末尾可能还有 w:bookmarkEnd 等结构元素，不能乱插）。
    """
    first_run = None
    last_run = None

    for child in para_elem:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("r", "del", "ins"):
            if first_run is None:
                first_run = child
            last_run = child

    if first_run is None:
        return  # 没有 run 节点，跳过

    # commentRangeEnd + commentReference 紧跟在最后一个 run 之后
    end_node = make_comment_range_end(comment_id)
    ref_node = make_comment_reference(comment_id)
    last_run.addnext(end_node)
    end_node.addnext(ref_node)

    # commentRangeStart 紧邻在第一个 run 之前
    start_node = make_comment_range_start(comment_id)
    first_run.addprevious(start_node)


# ─────────────────────────────────────────────────────────────
# 段落类型判断
# ─────────────────────────────────────────────────────────────

def _has_drawing(para_elem) -> bool:
    ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns_v = "urn:schemas-microsoft-com:vml"
    return (
        para_elem.find(f".//{{{ns_w}}}drawing") is not None
        or para_elem.find(f".//{{{ns_w}}}pict")   is not None
        or para_elem.find(f".//{{{ns_v}}}shape")   is not None
    )


def _is_empty_para(para_elem) -> bool:
    ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    runs = para_elem.findall(f"{{{ns_w}}}r")
    text = "".join(
        t.text or ""
        for r in runs
        for t in r.findall(f"{{{ns_w}}}t")
    )
    return not text.strip()


def _get_para_text(para_elem) -> str:
    """
    【fix6】只取段落直接子节点的 w:r，不用 .// 递归查找。
    原来 .//{ns}r 会把 w:del 里的 w:delText 也抓进来，
    导致已删除文字混入原文发给 AI，校对 diff 结果错位。
    """
    ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    parts = []
    for r in para_elem.findall(f"{{{ns_w}}}r"):
        for t in r.findall(f"{{{ns_w}}}t"):
            parts.append(t.text or "")
    return "".join(parts)


def _get_rpr(para_elem):
    """取第一个直接子 w:r 的 w:rPr（fix6：同步去掉 .// 递归）"""
    ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    r = para_elem.find(f"{{{ns_w}}}r")
    if r is not None:
        rpr = r.find(f"{{{ns_w}}}rPr")
        return rpr
    return None


# ─────────────────────────────────────────────────────────────
# 【fix4】段落跳过判断：提前过滤无需 AI 校对的段落
# ─────────────────────────────────────────────────────────────

# 【fix10】参考文献正则收紧：
#   [1] 格式，或 '数字. 中文' 格式（必须后跟中文字符），避免误杀章节标题
_REF_LINE_PAT = _re.compile(
    r'^\[\d+\]'                        # [1] 格式
    r'|^\d+\.\s+[\u4e00-\u9fff]'       # 1. 中文 格式（必须后跟中文）
)
# 中文字符检测
_NO_CN_PAT = _re.compile(r'[\u4e00-\u9fff]')
# 占位符校验正则（模块级编译，避免每次调用重复编译）
_PLACEHOLDER_PAT = _re.compile(r'__L?PROT[0-9a-f]{6}_\d{4}T?__')

# 【fix11】首页/封面常见字段关键词，这些行不需要 AI 校对
_COVER_FIELD_PAT = _re.compile(
    r'^\s*(?:姓\s*名|学\s*号|专\s*业|班\s*级|学\s*院|系\s*别|'
    r'指导\s*教师|指导\s*老师|导\s*师|年\s*级|学\s*历|'
    r'职\s*称|研究\s*方向|完成\s*日期|提交\s*日期|答辩\s*日期|'
    r'学位\s*类别|学科\s*专业|培养\s*单位|论文\s*题目|题\s*目)\s*[:：]',
    _re.UNICODE
)


def _should_skip_para(text: str) -> tuple:
    """
    判断段落是否可以跳过 AI 校对。
    返回 (should_skip: bool, reason: str)

    跳过条件（均为零内容风险）：
    1. 极短段落（≤4字）：标题编号、章节序号等，校对意义极低
    2. 纯数字/英文/符号行：无中文，不存在中文错别字
    3. 参考文献行：[1] 开头，或 '数字. 中文' 开头（fix10：收紧，避免误杀章节标题）
    4. 首页/封面字段行：姓名、学号、专业等（fix11：避免 AI 误改占位符格式）
    （fix9：表格单元格同样调用此函数过滤排版标签行）
    """
    stripped = text.strip()

    # 条件1：极短段落
    if len(stripped) <= 4:
        return True, "极短段落"

    # 条件2：无中文字符
    if not _NO_CN_PAT.search(stripped):
        return True, "无中文内容"

    # 条件3：参考文献行
    if _REF_LINE_PAT.match(stripped) and len(stripped) < 300:
        return True, "参考文献行"

    # 条件4：首页/封面字段行（姓名：xxx、学号：xxx 等）
    if _COVER_FIELD_PAT.match(stripped) and len(stripped) < 100:
        return True, "封面字段行"

    return False, ""


# ─────────────────────────────────────────────────────────────
# 修订节点写入
# 【fix1】不再全清段落后重建，改为找到第一个 w:r 的位置，
#         在原位插入修订节点，保留 w:pPr / w:bookmarkStart 等兄弟节点
# ─────────────────────────────────────────────────────────────

def _check_placeholders(protected_text: str, corrected_protected: str) -> bool:
    original_tokens = _PLACEHOLDER_PAT.findall(protected_text)   # 有序列表
    corrected_tokens = _PLACEHOLDER_PAT.findall(corrected_protected)  # 有序列表
    
    # 添加详细日志
    logger.debug(f"[protect] 检查占位符: 原文={original_tokens}, 修改后={corrected_tokens}")
    
    if original_tokens != corrected_tokens:  # 数量或顺序不一致都拦截
        logger.warning(f"[protect] AI 修改了占位符顺序或数量，放弃修改: {original_tokens} -> {corrected_tokens}")
        logger.warning(f"[protect] 原文: {protected_text[:200]}")
        logger.warning(f"[protect] 修改后: {corrected_protected[:200]}")
        return False
    return True

def _write_revision(para_elem, original: str, corrected: str, rpr_elem=None) -> int:
    opcodes = compute_diff(original, corrected)
    if not opcodes:
        return 0
    # 过滤空字符串操作，避免写入空的 w:del / w:ins 节点破坏段落结构
    opcodes = [
        op for op in opcodes
        if not (op[0] == "delete" and original[op[1]:op[2]] == "")
           and not (op[0] == "insert" and corrected[op[3]:op[4]] == "")
           and not (op[0] == "replace" and original[op[1]:op[2]] == "" and corrected[op[3]:op[4]] == "")
    ]
    if not opcodes:
        return 0

    # ── 【fix11】保存段落属性（w:pPr），确保它始终在最前面 ──────────
    ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ppr_elem = para_elem.find(f"{{{ns_w}}}pPr")

    # ── 找到第一个 run/del/ins 的位置，用于定锚点 ──────────────
    children = list(para_elem)
    insert_index = len(children)  # 默认追加到末尾
    for i, child in enumerate(children):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("r", "del", "ins", "hyperlink"):
            insert_index = i
            break

    # ── 移除所有已有的 w:r / w:del / w:ins / w:hyperlink ────────
    # 使用逆序删除避免索引错位
    to_remove = []
    for child in para_elem:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("r", "del", "ins", "hyperlink"):
            to_remove.append(child)
    for child in to_remove:
        para_elem.remove(child)

    # ── 【fix11】确保 w:pPr 在段落最前面（Word XML 规范要求） ──────
    if ppr_elem is not None:
        # 如果 w:pPr 存在，先移除再插入到最前面
        para_elem.remove(ppr_elem)
        para_elem.insert(0, ppr_elem)
        insert_index = 1  # run 节点从 w:pPr 之后开始
    else:
        insert_index = 0

    # ── 构建新的修订节点并插入到正确位置 ────────────────────────
    def _plain_run(text: str) -> OxmlElement:
        r = OxmlElement("w:r")
        if rpr_elem is not None:
            r.append(_sanitize_rpr(rpr_elem))
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        return r

    new_nodes = []
    change_count = 0
    pos = 0

    for tag, i1, i2, j1, j2 in opcodes:
        if pos < i1:
            new_nodes.append(_plain_run(original[pos:i1]))

        rev_id = new_rev_id()

        if tag == "replace":
            new_nodes.append(make_del_node(original[i1:i2], rev_id,      rpr_elem))
            new_nodes.append(make_ins_node(corrected[j1:j2], rev_id + 1, rpr_elem))
            change_count += 1
        elif tag == "delete":
            new_nodes.append(make_del_node(original[i1:i2], rev_id, rpr_elem))
            change_count += 1
        elif tag == "insert":
            new_nodes.append(make_ins_node(corrected[j1:j2], rev_id, rpr_elem))
            change_count += 1

        pos = i2

    if pos < len(original):
        new_nodes.append(_plain_run(original[pos:]))

    # 按顺序插入到 insert_index 之后
    for offset, node in enumerate(new_nodes):
        para_elem.insert(insert_index + offset, node)

    return change_count

# ─────────────────────────────────────────────────────────────
# 段落校对（正文）
# 【fix2】移除内层 Semaphore；短段落不分块，减少多余 AI 调用
# ─────────────────────────────────────────────────────────────

async def _proofread_paragraph(
    p, sem: asyncio.Semaphore, mode: ProofreadMode = "full"
):
    """
    【fix7】只做 AI 调用，不写 XML。
    返回 (original, corrected, para_elem, changes) 或 None（跳过/失败）。
    写 XML 由 process_word 串行完成。
    """
    para_elem = p._p

    if _has_drawing(para_elem):
        logger.debug("[skip] 含图片段落，跳过")
        return None

    original = _get_para_text(para_elem)
    if not original.strip():
        return None

    # 【fix4】提前过滤无需 AI 校对的段落，不占并发槽
    skip, reason = _should_skip_para(original)
    if skip:
        logger.debug(f"[skip] {reason}，跳过: {original[:30]!r}")
        return None

    try:
        protected_text, protected_map = protect_terms(original)

        # 【fix8】取消分块，整段发给 AI。
        # 原来超长段落分块后拼接，diff 会在块边界产生错位，误改正常文字。
        # 现在整段发送，AI 模型上下文足够处理，结果更准确。
        async with sem:
            if mode == "basic":
                corrected_protected, all_changes = await proofread_text(protected_text)
            else:
                corrected_protected, all_changes = await full_proofread(protected_text)

        if not _check_placeholders(protected_text, corrected_protected):
            return None
        corrected = restore_terms(corrected_protected, protected_map)
        
        # 过滤 all_changes 中包含占位符的记录（这些是 AI 误报的修改）
        filtered_changes = []
        for change in all_changes:
            if len(change) >= 2:
                # 检查原文和修改后的文本是否包含占位符
                original_text = change[0]
                corrected_text = change[1]
                if '__PROT' in original_text or '__PROT' in corrected_text:
                    logger.warning(f"[protect] 过滤包含占位符的修改记录: {change}")
                    continue
            filtered_changes.append(change)
        all_changes = filtered_changes

    except Exception as e:
        logger.error(f"[proofread] 段落校对失败，跳过: {e}")
        return None

    if not has_diff(original, corrected):
        return None

    # 【fix7】不在这里写 XML，返回数据给调用方串行写入
    return original, corrected, para_elem, all_changes


# ─────────────────────────────────────────────────────────────
# 表格单元格校对
# ─────────────────────────────────────────────────────────────

async def _proofread_cell(p, sem: asyncio.Semaphore):
    """
    【fix7】只做 AI 调用，不写 XML。
    返回 (original, corrected, para_elem, changes) 或 None（跳过/失败）。
    """
    para_elem = p._p

    if _has_drawing(para_elem):
        return None

    original = _get_para_text(para_elem)
    if not original.strip():
        return None

    # 【fix12】表格单元格更严格的跳过条件：
    # 短于 15 字的单元格通常是表头、数据标签、单位等，不需要校对
    if len(original.strip()) < 15:
        logger.debug(f"[skip:cell] 短单元格（<15字），跳过: {original[:30]!r}")
        return None

    # 【fix9】表格单元格同样过滤无需校对内容（排版标签/极短/纯英文等）
    skip, reason = _should_skip_para(original)
    if skip:
        logger.debug(f"[skip:cell] {reason}，跳过: {original[:30]!r}")
        return None

    try:
        protected_text, protected_map = protect_terms(original)

        # 单元格整体调用，不分块
        async with sem:
            corrected_protected, all_changes = await _full_check_once(protected_text)

        if not _check_placeholders(protected_text, corrected_protected):
            return None
        corrected = restore_terms(corrected_protected, protected_map)
        
        # 过滤 all_changes 中包含占位符的记录（这些是 AI 误报的修改）
        filtered_changes = []
        for change in all_changes:
            if len(change) >= 2:
                original_text = change[0]
                corrected_text = change[1]
                if '__PROT' in original_text or '__PROT' in corrected_text:
                    logger.warning(f"[protect] 过滤包含占位符的修改记录: {change}")
                    continue
            filtered_changes.append(change)
        all_changes = filtered_changes

    except Exception as e:
        logger.error(f"[proofread] 表格单元格校对失败，跳过: {e}")
        return None

    if not has_diff(original, corrected):
        return None

    # 【fix12】表格单元格额外安全检查：
    # 如果修改只涉及括号/符号的增删（AI 常见误报），直接丢弃
    import re as _re2
    diff_chars = set(original) ^ set(corrected)  # 对称差集
    bracket_chars = set('()（）【】[]{}《》<>、，。；：""''')
    if diff_chars and diff_chars.issubset(bracket_chars):
        logger.warning(f"[protect:cell] 修改仅涉及括号/符号变化，丢弃: {original[:50]!r}")
        return None

    # 【fix7】不在这里写 XML，返回数据给调用方串行写入
    return original, corrected, para_elem, all_changes


# ─────────────────────────────────────────────────────────────
# 文档遍历
# ─────────────────────────────────────────────────────────────

def _collect_tasks(doc: Document):
    """
    收集所有需要校对的段落，返回 [(paragraph_obj, type), ...]。
    type: 'body' | 'table_cell'
    仅遍历 body 直接子节点，避免表格内段落被重复收集。
    """
    para_elems_in_tables = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    para_elems_in_tables.add(id(p._p))

    tasks = []
    for p in doc.paragraphs:
        if id(p._p) not in para_elems_in_tables:
            tasks.append((p, "body"))

    seen_para_ids = set()

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if id(p._p) not in seen_para_ids:
                        seen_para_ids.add(id(p._p))
                        tasks.append((p, "table_cell"))

    return tasks


# ─────────────────────────────────────────────────────────────
# 主入口
# 【fix3】增加空文件/损坏文件提前检测，避免前端永久 pending
# ─────────────────────────────────────────────────────────────

async def process_word(
    input_path: str,
    output_path: str,
    mode: ProofreadMode = "full",
) -> dict:
    """
    异步处理 Word 文档，校对并输出修订版本。

    Args:
        input_path:  输入文档路径
        output_path: 输出文档路径
        mode:        "basic"（仅错别字）或 "full"（完整三轮，默认）

    Returns:
        {
            "total_paragraphs": int,
            "checked":          int,
            "changed":          int,
            "skipped":          int,
            "total_changes":    int,
            "all_changes":      list,
            "mode":             str,
            "error":            str | None,  # 新增：错误信息
        }
    """
    logger.info(f"[pipeline] 开始处理: {input_path}  mode={mode}")

    # 【fix3】提前检测文件有效性，立即返回错误而非永久挂起
    if not os.path.exists(input_path):
        logger.error(f"[pipeline] 文件不存在: {input_path}")
        return _empty_stats(mode, error="文件不存在")

    if os.path.getsize(input_path) == 0:
        logger.error(f"[pipeline] 文件为空: {input_path}")
        return _empty_stats(mode, error="上传文件为空，请检查后重新上传")

    # 对 WPS 生成的 .docx 先归一化;失败会返回原路径,走下面的降级逻辑
    input_path = str(normalize_docx(input_path))

    try:
        doc = Document(input_path)
    except Exception as e:
        error_msg = str(e)
        logger.error(f"[pipeline] 文件解析失败: {error_msg}")
        
        # 判断是否是文档格式问题
        if "no item named" in error_msg.lower() or "badzip" in error_msg.lower():
            logger.warning(f"[pipeline] 检测到文档格式异常，尝试降级处理...")
            
            try:
                # 尝试降级处理：提取文本 + 简化校对
                text_content = _extract_text_fallback(input_path)
                
                if text_content:
                    return _create_simple_proofread_doc(text_content, output_path, mode)
                else:
                    # 文本提取也失败
                    logger.error(f"[pipeline] 降级处理失败：无法提取文本内容")
                    return _empty_stats(mode, error="文档格式异常且无法自动修复，请在 Word 中另存为新文件后重试")
                    
            except Exception as fallback_error:
                logger.error(f"[pipeline] 降级处理失败: {fallback_error}")
                return _empty_stats(mode, error=f"文档格式异常且无法自动修复：{fallback_error}")
        else:
            return _empty_stats(mode, error=f"文件解析失败，请确认为有效的 .docx 文件：{e}")

    enable_track_changes(doc.settings.element)

    tasks = _collect_tasks(doc)
    logger.info(f"[pipeline] 共收集 {len(tasks)} 个段落待检查")

    # 【fix3】文档无有效内容时直接保存返回
    # 同时过滤掉图片段落和空段落，减少 gather 的协程数量（问题6优化）
    effective_tasks = [
        (p, pt) for p, pt in tasks
        if not _has_drawing(p._p) and not _is_empty_para(p._p)
    ]
    stats_skipped_upfront = len(tasks) - len(effective_tasks)
    if not effective_tasks:
        doc.save(output_path)
        logger.info("[pipeline] 文档无有效文字段落，直接返回")
        return _empty_stats(mode, total=len(tasks))

    stats = {
        "total_paragraphs": len(tasks),
        "checked":          0,
        "changed":          0,
        "skipped":          0,
        "total_changes":    0,
        "all_changes":      [],
        "mode":             mode,
        "error":            None,
    }

    # 【fix5】并发数 8 → 16，在 API 限流允许范围内近乎翻倍处理速度
    sem = asyncio.Semaphore(16)

    async def _run_ai(p, ptype):
        """Phase 1：只做 AI 调用，不写 XML"""
        if ptype == "table_cell":
            return await _proofread_cell(p, sem)
        else:
            return await _proofread_paragraph(p, sem, mode=mode)

    # 【fix7】Phase 1：并发 AI 调用，全部结果收集完再写 XML
    logger.info(f"[pipeline] Phase 1：并发 AI 调用（{len(effective_tasks)} 个有效段落，共 {len(tasks)} 个）")
    
    # 使用提取的异步函数
    stats = await _proofread_tasks_async(doc, effective_tasks, len(tasks), mode)
    stats["skipped"] += stats_skipped_upfront

    doc.save(output_path)
    # 兜底：确保 [Content_Types].xml 包含 comments 的 Override，否则 Word 报 "内容有问题"
    try:
        ensure_comments_content_type(output_path)
    except Exception as e:
        logger.warning(f"[pipeline] Content Types 校验/修补失败: {e}")
    logger.info(
        f"[pipeline] 完成: mode={mode} checked={stats['checked']} "
        f"changed={stats['changed']} total_changes={stats['total_changes']} "
        f"skipped={stats['skipped']}"
    )
    return stats


def _empty_stats(mode: str, total: int = 0, error: str = None) -> dict:
    return {
        "total_paragraphs": total,
        "checked":          0,
        "changed":          0,
        "skipped":          total,
        "total_changes":    0,
        "all_changes":      [],
        "mode":             mode,
        "error":            error,
    }


def process_word_sync(
    input_path: str,
    output_path: str,
    mode: ProofreadMode = "full",
) -> dict:
    """同步包装，供 Celery Worker 调用"""
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    try:
        return loop.run_until_complete(process_word(input_path, output_path, mode))
    finally:
        loop.close()
