# backend/app/core/formatter/template_manager.py
"""
企业级模板管理器
支持模板CRUD、缓存、版本控制
"""


from pathlib import Path
from typing import List, Dict, Optional, Union
from dataclasses import dataclass, asdict
from datetime import datetime
import json
import hashlib
import shutil  #文件复制、移动操作
import zipfile
from loguru import logger

from docx import Document
from app.services.docx_normalizer import normalize_docx
from docx.shared import Pt, Inches, RGBColor  #Word 字体尺寸、颜色单位。
from docx.enum.text import WD_ALIGN_PARAGRAPH  #Word 对齐方式枚举。
from docx.oxml.ns import qn  #XML 命名空间工具（高级 Word 操作）


@dataclass
class TemplateMetadata:   #定义模板数据库模型（类似 ORM 表结构）。
    """模板元数据"""
    template_id: str  #模板唯一ID
    name: str
    category: str  # 分类：大学模板、期刊模板、自定义模板
    school_or_journal: str  #学校名或期刊名。
    description: str
    version: str = "1.0"
    author: str = "system"
    created_at: str = None
    updated_at: str = None
    usage_count: int = 0   #使用次数统计（推荐排序用）。
    is_public: bool = True  #是否内置模板（不可删除）。
    tags: List[str] = None  #模板标签（搜索用）
    file_path: str = None  #Word 文件路径、
    preview_image: str = None  #预览图路径
    
    def __post_init__(self):  #__post_init__ = dataclass 构造函数之后自动执行的钩子函数
        if self.created_at is None:
            self.created_at = datetime.now().isoformat()
        if self.updated_at is None:
            self.updated_at = datetime.now().isoformat()
        if self.tags is None:
            self.tags = []


@dataclass
class TemplateConfig:
    """模板配置结构"""

    # 页面设置
    page_width: float  # 单位：cm
    page_height: float
    margin_top: float   #页边距。
    margin_bottom: float
    margin_left: float
    margin_right: float

    # 样式配置
    styles: Dict[str, Dict]  # 样式名 -> 样式配置Word 样式库提取结果。

    # 编号配置
    numbering: Dict[str, str]  # 章节编号类型 -> 格式

    # 页眉页脚
    header: Optional[Dict] = None
    footer: Optional[Dict] = None

    # 其他配置
    line_number: bool = False
    track_changes: bool = False

    # 模板是否包含目录（决定排版时是否覆盖用户文档的原目录样式）
    has_toc: bool = False

    # 默认值（方法内硬编码的兜底值，可在模板中覆盖）
    default_toc_line_spacing: float = 1.5    # 目录默认行距
    default_toc_line_spacing_rule: str = 'MULTIPLE'  # 目录行距规则：MULTIPLE=倍数模式
    default_caption_space_before: float = 6.0  # 图题/表题段前 pt
    default_caption_space_after: float = 6.0   # 图题/表题段后 pt


class TemplateManager:
    """
    企业级模板管理器
    
    功能:
    - 模板CRUD操作
    - 样式提取和应用
    - 模板缓存机制
    - 版本控制
    - 模板验证
    """
    
    def __init__(
        self, 
        builtin_dir: Path,
        user_dir: Path,
        cache_dir: Optional[Path] = None
    ):   #三种模板目录
        self.builtin_dir = Path(builtin_dir)
        self.user_dir = Path(user_dir)
        self.cache_dir = Path(cache_dir) if cache_dir else self.user_dir / ".cache"
        
        # 创建必要目录
        self.builtin_dir.mkdir(parents=True, exist_ok=True)  #自动创建目录（Linux/Windows 兼容
        self.user_dir.mkdir(parents=True, exist_ok=True)
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        
        # 模板索引（内存缓存）
        self._template_index: Dict[str, TemplateMetadata] = {}
        self._config_cache: Dict[str, TemplateConfig] = {}
        # 模板 XML 部件缓存（styles.xml / numbering.xml / theme1.xml）
        self._xml_parts_cache: Dict[str, Dict[str, bytes]] = {}
        
        # 初始化索引
        self._build_index()
        
        logger.info(f"模板管理器初始化完成")
        logger.info(f"内置模板目录: {self.builtin_dir}")
        logger.info(f"用户模板目录: {self.user_dir}")
    
    def _build_index(self):
        """构建模板索引"""
        # 扫描内置模板
        self._scan_directory(self.builtin_dir, is_builtin=True)
        
        # 扫描用户模板
        self._scan_directory(self.user_dir, is_builtin=False)
        
        logger.info(f"索引构建完成，共 {len(self._template_index)} 个模板")
    
    def _scan_directory(self, directory: Path, is_builtin: bool):
        """扫描目录中的模板"""
        if not directory.exists():
            return
        
        for template_file in directory.rglob("*.docx"): #递归遍历目录下所有 .docx 文件（包括子目录）返回的是 Path 对象，不是字符串
            # 跳过临时文件
            if template_file.name.startswith("~"):#过滤临时文件
                continue
            
            try:
                # 读取元数据
                metadata_file = template_file.with_suffix(".json") #把文件扩展名改为 .json，返回新的Path
                
                if metadata_file.exists():
                    # 从JSON读取
                    with open(metadata_file, 'r', encoding='utf-8') as f:
                        metadata_dict = json.load(f) #从文件对象 f 中读取 JSON 内容，并解析成 Python 字典（dict）
                        metadata = TemplateMetadata(**metadata_dict)
                else:
                    # 自动生成元数据
                    metadata = self._generate_metadata(template_file, is_builtin)
                
                metadata.file_path = str(template_file)
                self._template_index[metadata.template_id] = metadata
                
                logger.debug(f"索引模板: {metadata.name}")
            
            except Exception as e:
                logger.error(f"索引模板失败 {template_file}: {e}")
    
    def _generate_metadata(self, template_file: Path, is_builtin: bool) -> TemplateMetadata:
        """自动生成模板元数据"""
        # 从文件名提取信息
        name = template_file.stem
        
        # 生成template_id
        template_id = hashlib.md5(
            str(template_file).encode()
        ).hexdigest()[:16]
        
        # 判断分类
        category = "builtin" if is_builtin else "custom"
        if "university" in str(template_file).lower() or "大学" in name:
            category = "universities"
        elif "journal" in str(template_file).lower() or "期刊" in name:
            category = "journals"
        
        return TemplateMetadata(
            template_id=template_id,
            name=name,
            category=category,
            school_or_journal=name,
            description=f"自动生成的模板：{name}",
            is_public=is_builtin
        )
    
    def get_template(self, template_id: str) -> Optional[TemplateMetadata]:
        """获取模板元数据。若未命中，重新扫描用户目录（兼容跨进程：Web 写入/Worker 读取）"""
        meta = self._template_index.get(template_id)
        if meta is None:
            # Web 进程新上传的模板，Celery Worker 的内存索引还不知道 → 重扫盘
            logger.info(f"模板未命中内存索引，重新扫描用户目录: {template_id}")
            self._scan_directory(self.user_dir, is_builtin=False)
            meta = self._template_index.get(template_id)
        return meta
    
    def list_templates(
        self,
        category: Optional[str] = None,
        search: Optional[str] = None,
        limit: int = 100
    ) -> List[TemplateMetadata]:
        """
        列出模板
        
        Args:
            category: 分类过滤
            search: 搜索关键词
            limit: 返回数量限制
        """
        templates = list(self._template_index.values())
        
        # 分类过滤
        if category:
            templates = [t for t in templates if t.category == category]
        
        # 搜索过滤
        if search:
            search_lower = search.lower()
            templates = [
                t for t in templates
                if search_lower in t.name.lower() or
                   search_lower in t.school_or_journal.lower() or
                   search_lower in t.description.lower()
            ]
        
        # 排序（按使用次数降序）
        templates.sort(key=lambda x: x.usage_count, reverse=True)
        
        return templates[:limit]
    
    def load_document(self, template_id: str) -> Document:
        """加载模板文档"""
        metadata = self.get_template(template_id)

        if not metadata:
            raise ValueError(f"模板不存在: {template_id}")

        try:
            metadata.file_path = str(normalize_docx(metadata.file_path))
            doc = Document(metadata.file_path)

            # 更新使用次数（后台线程写磁盘，不阻塞格式化主流程）
            metadata.usage_count += 1
            import threading
            threading.Thread(
                target=self._save_metadata,
                args=(metadata,),
                daemon=True
            ).start()

            logger.info(f"模板已加载: {metadata.name}")
            return doc

        except Exception as e:
            logger.error(f"加载模板失败: {e}")
            raise
    
    # 样式注入所需的核心 XML 部件
    _STYLE_INJECT_PARTS = (
        "word/styles.xml",
        "word/numbering.xml",
        "word/theme/theme1.xml",
    )

    def get_template_xml_parts(self, template_id: str) -> Dict[str, bytes]:
        """
        从模板 .docx ZIP 里抽出核心 XML 部件的字节内容。

        用于"样式注入"流程：把这些部件写入用户文档 ZIP，让 Word 用模板的
        样式定义渲染。返回 dict：{'word/styles.xml': b'...', ...}
        带内存缓存避免重复解析。
        """
        if template_id in self._xml_parts_cache:
            return self._xml_parts_cache[template_id]

        metadata = self.get_template(template_id)  # 内部已带重扫盘兜底
        if metadata is None:
            raise ValueError(f"模板不存在: {template_id}")

        template_path = Path(metadata.file_path)
        if not template_path.exists():
            raise FileNotFoundError(f"模板文件丢失: {template_path}")

        parts: Dict[str, bytes] = {}
        with zipfile.ZipFile(template_path, "r") as z:
            names = set(z.namelist())
            for t in self._STYLE_INJECT_PARTS:
                if t in names:
                    parts[t] = z.read(t)

        self._xml_parts_cache[template_id] = parts
        logger.info(f"模板 XML 部件已抽取 {template_id}: {list(parts.keys())}")
        return parts

    def extract_config(self, template_id: str) -> TemplateConfig:
        """
        提取模板配置
        带缓存机制
        """
        # 检查缓存
        if template_id in self._config_cache:
            logger.debug(f"从缓存读取配置: {template_id}")
            return self._config_cache[template_id]
        
        # 加载文档
        doc = self.load_document(template_id)
        
        # 提取配置
        config = TemplateConfig(
            page_width=self._extract_page_width(doc),
            page_height=self._extract_page_height(doc),
            margin_top=self._extract_margin(doc, 'top'),
            margin_bottom=self._extract_margin(doc, 'bottom'),
            margin_left=self._extract_margin(doc, 'left'),
            margin_right=self._extract_margin(doc, 'right'),
            styles=self._extract_styles(doc),
            numbering=self._extract_numbering(doc),
            header=self._extract_header(doc),
            footer=self._extract_footer(doc),
            has_toc=self._detect_has_toc(doc),
        )
        
        # 缓存配置
        self._config_cache[template_id] = config
        
        logger.info(f"配置提取完成: {template_id}")
        return config
    
    def _extract_page_width(self, doc: Document) -> float:
        """提取页面宽度（cm）"""
        section = doc.sections[0]
        return section.page_width.cm
    
    def _extract_page_height(self, doc: Document) -> float:
        """提取页面高度（cm）"""
        section = doc.sections[0]
        return section.page_height.cm
    
    def _extract_margin(self, doc: Document, position: str) -> float:
        """提取页边距（cm）"""
        section = doc.sections[0]
        margin_map = {
            'top': section.top_margin,
            'bottom': section.bottom_margin,
            'left': section.left_margin,
            'right': section.right_margin
        }
        return margin_map[position].cm
    
    def _resolve_font_xml_attr(self, style, child_tag: str, attr_name: str):
        """
        沿样式继承链查找 rPr 下某个子元素的某个属性的字符串值。
        python-docx 的 style.font.name 只返回 w:ascii（西文字体），
        无法拿到 w:eastAsia（中文字体）和 w:color（颜色），所以直接读 XML。

        child_tag / attr_name 使用不带命名空间的本地名，例如:
          - ('rFonts', 'eastAsia') → 中文字体
          - ('rFonts', 'ascii')    → 西文字体
          - ('color',  'val')      → 字体颜色（十六进制 或 'auto'）
        """
        visited = set()
        current = style
        while current is not None:
            if id(current) in visited:
                break
            visited.add(id(current))
            try:
                rPr = current.element.find(qn('w:rPr'))
                if rPr is not None:
                    child = rPr.find(qn(f'w:{child_tag}'))
                    if child is not None:
                        val = child.get(qn(f'w:{attr_name}'))
                        if val is not None and val != '':
                            return val
            except Exception:
                pass
            current = getattr(current, 'base_style', None)
        return None

    def _resolve_style_attr(self, style, attr: str, pf_attr: str = None, is_font: bool = False):
        """
        沿样式继承链查找属性值，直到找到非 None 值或到达根样式。
        attr: font 属性名 或 paragraph_format 属性名
        pf_attr: 若 is_font=False，paragraph_format 上的属性名（默认同 attr）
        """
        visited = set()
        current = style
        while current is not None:
            if id(current) in visited:
                break
            visited.add(id(current))
            try:
                if is_font:
                    val = getattr(current.font, attr, None)
                else:
                    key = pf_attr or attr
                    val = getattr(current.paragraph_format, key, None)
                if val is not None:
                    return val
            except Exception:
                pass
            current = getattr(current, 'base_style', None)
        return None

    def _extract_styles(self, doc: Document) -> Dict[str, Dict]:
        """提取段落样式，沿继承链解析 None 值"""
        styles = {}

        for style in doc.styles:
            # 只提取段落样式
            if style.type != 1:  # 1 = WD_STYLE_TYPE.PARAGRAPH
                continue

            try:
                style_config = {
                    "name": style.name,
                    "font_name": None,          # 兼容旧字段（优先西文，其次中文）
                    "font_name_cn": None,       # 中文字体（eastAsia）
                    "font_name_en": None,       # 西文字体（ascii）
                    "font_size": None,
                    "bold": False,
                    "italic": False,
                    "underline": False,
                    "color": None,              # 'auto' / 'FF0000' / None
                    "alignment": None,
                    "line_spacing": None,
                    "space_before": None,
                    "space_after": None,
                    "first_line_indent": None,
                    "left_indent": None,
                    "right_indent": None
                }

                # 字体属性（直接读 XML，区分中/西文）
                font_en = self._resolve_font_xml_attr(style, 'rFonts', 'ascii')
                font_cn = self._resolve_font_xml_attr(style, 'rFonts', 'eastAsia')
                # 回退：python-docx 的 font.name（通常等价 ascii）
                font_fallback = self._resolve_style_attr(style, 'name', is_font=True)

                if font_en:
                    style_config["font_name_en"] = font_en
                if font_cn:
                    style_config["font_name_cn"] = font_cn
                style_config["font_name"] = font_en or font_fallback or font_cn

                # 字体颜色（可能是 'auto' 或十六进制）
                color_val = self._resolve_font_xml_attr(style, 'color', 'val')
                if color_val:
                    style_config["color"] = color_val

                font_size = self._resolve_style_attr(style, 'size', is_font=True)
                if font_size:
                    style_config["font_size"] = font_size.pt

                bold = self._resolve_style_attr(style, 'bold', is_font=True)
                style_config["bold"] = bool(bold)

                italic = self._resolve_style_attr(style, 'italic', is_font=True)
                style_config["italic"] = bool(italic)

                underline = self._resolve_style_attr(style, 'underline', is_font=True)
                style_config["underline"] = bool(underline)

                # 段落格式（沿继承链查找）
                alignment = self._resolve_style_attr(style, 'alignment')
                if alignment is not None:
                    style_config["alignment"] = str(alignment)

                line_spacing = self._resolve_style_attr(style, 'line_spacing')
                if line_spacing is not None:
                    style_config["line_spacing"] = line_spacing

                # 行距规则（关键：倍数 vs 固定值）
                line_spacing_rule = self._resolve_style_attr(style, 'line_spacing_rule')
                if line_spacing_rule is not None:
                    # 转换为字符串便于序列化和比对
                    style_config["line_spacing_rule"] = str(line_spacing_rule)

                space_before = self._resolve_style_attr(style, 'space_before')
                if space_before is not None:
                    style_config["space_before"] = space_before.pt

                space_after = self._resolve_style_attr(style, 'space_after')
                if space_after is not None:
                    style_config["space_after"] = space_after.pt

                first_line_indent = self._resolve_style_attr(style, 'first_line_indent')
                if first_line_indent is not None:
                    style_config["first_line_indent"] = first_line_indent.cm

                left_indent = self._resolve_style_attr(style, 'left_indent')
                if left_indent is not None:
                    style_config["left_indent"] = left_indent.cm

                right_indent = self._resolve_style_attr(style, 'right_indent')
                if right_indent is not None:
                    style_config["right_indent"] = right_indent.cm

                styles[style.name] = style_config

            except Exception as e:
                logger.warning(f"提取样式失败 {style.name}: {e}")

        return styles
    
    def _detect_has_toc(self, doc: Document) -> bool:
        """
        判断模板里是否出现目录。用于决定排版时是否覆盖用户文档的原目录样式：
          - 模板有目录 → 用模板 TOC 样式覆盖
          - 模板无目录 → 保留用户原目录格式，不动
        识别条件（命中任一即视为"有"）：
          1. 段落样式名包含 TOC / 目录
          2. 段落文本等于"目录"（兼容"目 录"、"目　录"）
          3. 文档 XML 里存在 TOC 域（w:fldChar + TOC 指令）
        """
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ''
            if 'TOC' in style_name.upper() or '目录' in style_name:
                return True
            text = para.text.strip().replace('　', '').replace(' ', '')
            if text == '目录':
                return True
        try:
            body = doc.element.body
            for instr in body.iter(qn('w:instrText')):
                if instr.text and 'TOC' in instr.text.upper():
                    return True
        except Exception:
            pass
        return False

    def _extract_numbering(self, doc: Document) -> Dict[str, str]:
        """
        提取编号配置

        通过扫描文档中有标题样式的段落，推断编号格式
        """
        result = {
            "chapter": "第{n}章",      # 默认值
            "section1": "{n}",
            "section2": "{n}.{m}",
            "section3": "{n}.{m}.{k}"
        }

        # 样式识别：检查段落是否为标题
        chapter_patterns = [
            re.compile(r"^第([一二三四五六七八九十百千]+)章"),  # "第一章"
            re.compile(r"^第(\d+)章"),  # "第1章"
            re.compile(r"^Chapter\s+(\d+)", re.IGNORECASE),  # "Chapter 1"
            re.compile(r"^([A-Z])\s+"),  # "A " 或 "A."
        ]

        section1_patterns = [
            re.compile(r"^(\d+)\.?\s+"),  # "1. " 或 "1 "
            re.compile(r"^([一二三四五六七八九十]+)[、]"),  # "一、"
        ]

        section2_patterns = [
            re.compile(r"^(\d+)\.(\d+)\.?\s+"),  # "1.1. " 或 "1.1 "
            re.compile(r"^(\d+)\.(\d+)\s+"),
        ]

        # 扫描段落，查找标题样式
        heading_styles = {"Heading 1", "Heading 2", "Heading 3", "标题 1", "标题 2", "标题 3"}

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text or len(text) < 2:
                continue

            style_name = para.style.name if para.style else ""

            # 章标题识别
            if style_name in {"Heading 1", "标题 1"} or style_name.startswith("Heading 1"):
                for pattern in chapter_patterns:
                    match = pattern.match(text)
                    if match:
                        # 推断格式
                        if "第" in text and "章" in text:
                            if any(c in match.group(1) for c in "一二三四五六七八九十百千"):
                                result["chapter"] = "第{n}章"
                            else:
                                result["chapter"] = "第{n}章"
                        elif "Chapter" in text.lower():
                            result["chapter"] = "Chapter {n}"
                        break

            # 一级节标题识别
            elif style_name in {"Heading 2", "标题 2"} or style_name.startswith("Heading 2"):
                for pattern in section1_patterns:
                    match = pattern.match(text)
                    if match:
                        if "、" in text:
                            result["section1"] = "{n}、"
                        elif "." in text:
                            result["section1"] = "{n}."
                        else:
                            result["section1"] = "{n}"
                        break

            # 二级节标题识别
            elif style_name in {"Heading 3", "标题 3"} or style_name.startswith("Heading 3"):
                for pattern in section2_patterns:
                    match = pattern.match(text)
                    if match:
                        if "." in text:
                            result["section2"] = "{n}.{m}"
                        break

        return result
    
    def _extract_header(self, doc: Document) -> Optional[Dict]:
        """提取页眉配置：仅当模板页眉有实际文本内容时才返回"""
        try:
            section = doc.sections[0]
            header = section.header
            text = ''.join(p.text for p in header.paragraphs).strip()
            if not text:
                return None
            return {"text": text, "align": "center"}
        except Exception:
            pass
        return None

    def _extract_footer(self, doc: Document) -> Optional[Dict]:
        """提取页脚配置：检测实际页码域，防止无页码模板强制插页码"""
        try:
            section = doc.sections[0]
            footer = section.footer
            text = ''.join(p.text for p in footer.paragraphs).strip()

            # 检查 XML 中是否存在 PAGE 域
            has_page_number = False
            for para in footer.paragraphs:
                for elem in para._element.iter():
                    if elem.tag == qn('w:instrText') and elem.text:
                        if 'PAGE' in elem.text.upper():
                            has_page_number = True
                            break

            if not text and not has_page_number:
                return None
            return {"text": text, "align": "center", "page_number": has_page_number}
        except Exception:
            pass
        return None
    
    def save_template(
        self,
        file_path: Union[str, Path],
        metadata: TemplateMetadata
    ) -> str:
        """
        保存用户上传的模板

        Returns:
            template_id
        """
        file_path = Path(file_path)

        # 生成保存路径
        save_dir = self.user_dir / metadata.category
        save_dir.mkdir(parents=True, exist_ok=True)

        # 保存文件
        template_file = save_dir / f"{metadata.template_id}.docx"
        shutil.copy(file_path, template_file)

        # 更新元数据，存储相对路径（相对于模板库根目录）
        templates_root = self.user_dir.parent
        try:
            relative_path = template_file.relative_to(templates_root)
            metadata.file_path = str(relative_path)
        except ValueError:
            # 如果无法计算相对路径，回落到绝对路径
            metadata.file_path = str(template_file)

        metadata.created_at = datetime.now().isoformat()

        # 保存元数据
        self._save_metadata(metadata)

        # 更新索引（内存中用绝对路径）
        metadata.file_path = str(template_file)
        self._template_index[metadata.template_id] = metadata

        logger.info(f"模板已保存: {metadata.name}")

        return metadata.template_id
    
    def _save_metadata(self, metadata: TemplateMetadata):
        """保存元数据到JSON文件"""
        if not metadata.file_path:
            return
        
        metadata_file = Path(metadata.file_path).with_suffix(".json")
        
        with open(metadata_file, 'w', encoding='utf-8') as f:
            json.dump(asdict(metadata), f, ensure_ascii=False, indent=2)
    
    def delete_template(self, template_id: str) -> bool:
        """删除模板"""
        metadata = self.get_template(template_id)
        
        if not metadata:
            return False
        
        # 只能删除用户模板
        if metadata.is_public:
            raise PermissionError("不能删除内置模板")
        
        try:
            # 删除文件
            Path(metadata.file_path).unlink()
            
            # 删除元数据文件
            metadata_file = Path(metadata.file_path).with_suffix(".json")
            if metadata_file.exists():
                metadata_file.unlink()
            
            # 从索引移除
            del self._template_index[template_id]
            
            # 清除缓存
            if template_id in self._config_cache:
                del self._config_cache[template_id]
            
            logger.info(f"模板已删除: {metadata.name}")
            return True
        
        except Exception as e:
            logger.error(f"删除模板失败: {e}")
            return False
    
    def validate_template(self, file_path: Union[str, Path]) -> Dict:
        """
        验证模板文件
        
        Returns:
            验证结果 {valid: bool, errors: List[str], warnings: List[str]}
        """
        errors = []
        warnings = []
        
        try:
            file_path = str(normalize_docx(file_path))
            doc = Document(file_path)

            # 检查是否有样式
            if len(doc.styles) == 0:
                warnings.append("模板中没有自定义样式")
            
            # 检查是否有内容
            if len(doc.paragraphs) == 0:
                warnings.append("模板中没有示例内容")
            
            # 检查页面设置
            section = doc.sections[0]
            if section.page_width.cm < 10 or section.page_height.cm < 10:
                errors.append("页面尺寸异常")
            
        except Exception as e:
            errors.append(f"无法打开文件: {str(e)}")
        
        return {
            "valid": len(errors) == 0,
            "errors": errors,
            "warnings": warnings
        }
    
    def extract_config_from_path(self, file_path: Union[str, Path]) -> TemplateConfig:
        """
        从文件路径直接提取模板配置（不需要先入库）。
        用于模板上传预览场景。
        """
        file_path = str(normalize_docx(file_path))
        doc = Document(file_path)
        return TemplateConfig(
            page_width=self._extract_page_width(doc),
            page_height=self._extract_page_height(doc),
            margin_top=self._extract_margin(doc, 'top'),
            margin_bottom=self._extract_margin(doc, 'bottom'),
            margin_left=self._extract_margin(doc, 'left'),
            margin_right=self._extract_margin(doc, 'right'),
            styles=self._extract_styles(doc),
            numbering=self._extract_numbering(doc),
            header=self._extract_header(doc),
            footer=self._extract_footer(doc),
            has_toc=self._detect_has_toc(doc),
        )

    def generate_style_summary(self, config: TemplateConfig) -> Dict[str, str]:
        """
        生成人类可读的样式摘要，用于前端预览确认。
        返回 { "正文": "宋体 小四，1.5倍行距", "一级标题": "黑体 三号，加粗，居中", ... }
        """
        summary = {}

        # ── pt → 中文字号映射 ──
        pt_to_chinese = {
            42: "初号", 36: "小初", 26: "一号", 24: "小一",
            22: "二号", 18: "小二", 16: "三号", 15: "小三",
            14: "四号", 12: "小四", 10.5: "五号", 9: "小五",
            7.5: "六号", 6.5: "小六", 5.5: "七号", 5: "八号",
        }

        def format_font_size(pt_val):
            """将 pt 值转为友好显示：优先中文字号，否则取整"""
            if pt_val is None:
                return None
            # 四舍五入到1位小数再匹配
            rounded = round(pt_val, 1)
            if rounded in pt_to_chinese:
                return f"{pt_to_chinese[rounded]}（{int(rounded) if rounded == int(rounded) else rounded}pt）"
            # 没有对应中文字号，显示整数 pt
            int_val = int(round(pt_val))
            return f"{int_val}pt"

        def format_line_spacing(ls):
            """行距格式化，最多保留1位小数"""
            if ls is None:
                return None
            if isinstance(ls, (int, float)):
                rounded = round(ls, 1)
                if rounded == int(rounded):
                    return f"{int(rounded)}倍行距"
                return f"{rounded}倍行距"
            return None

        # 同一逻辑标签的多个候选样式名（中英文 Word / WPS 不同命名）
        # 我们从每组里挑"非空字段最多"的那个作为展示来源，
        # 避免 Word 内置的 Heading 1（未定制，仅默认值）盖过模板实际使用的 标题 1。
        label_groups = [
            ("正文",     ["正文",   "Normal"]),
            ("一级标题", ["标题 1", "标题1", "Heading 1"]),
            ("二级标题", ["标题 2", "标题2", "Heading 2"]),
            ("三级标题", ["标题 3", "标题3", "Heading 3"]),
            ("四级标题", ["标题 4", "标题4", "Heading 4"]),
            ("论文标题", ["论文标题", "Title", "标题"]),
        ]

        def _score(style_dict):
            """非空关键字段计数，用来挑更像是"手工定制过"的样式。"""
            score = 0
            for key in ("font_name_cn", "font_name_en", "font_name",
                        "font_size", "color", "line_spacing",
                        "space_before", "space_after", "first_line_indent"):
                if style_dict.get(key):
                    score += 1
            if style_dict.get("bold"): score += 1
            if style_dict.get("italic"): score += 1
            return score

        def _pick_best(names):
            best = None
            best_score = -1
            for n in names:
                s = config.styles.get(n)
                if not s:
                    continue
                sc = _score(s)
                if sc > best_score:
                    best_score = sc
                    best = s
            return best

        def _format_color(val: str) -> Optional[str]:
            if not val:
                return None
            if val.lower() == "auto":
                return "颜色自动"
            # 只显示前 6 位 hex
            hex_val = val.lstrip('#')[:6].upper()
            return f"颜色#{hex_val}"

        for label, names in label_groups:
            s = _pick_best(names)
            if s is None:
                continue
            parts = []

            # 字体：区分中文/西文
            cn = s.get("font_name_cn")
            en = s.get("font_name_en") or s.get("font_name")
            if cn and en and cn != en:
                parts.append(f"中文{cn}/西文{en}")
            elif cn:
                parts.append(f"中文{cn}")
            elif en:
                parts.append(en)

            # 字号
            if s.get("font_size"):
                size_str = format_font_size(s["font_size"])
                if size_str:
                    parts.append(size_str)

            # 加粗/斜体
            if s.get("bold"):
                parts.append("加粗")
            if s.get("italic"):
                parts.append("斜体")

            # 颜色
            color_str = _format_color(s.get("color"))
            if color_str:
                parts.append(color_str)

            # 行距
            ls_str = format_line_spacing(s.get("line_spacing"))
            if ls_str:
                parts.append(ls_str)

            # 对齐
            if s.get("alignment"):
                align_str = str(s["alignment"]).upper()
                if "CENTER" in align_str:
                    parts.append("居中")
                elif "RIGHT" in align_str:
                    parts.append("右对齐")
                elif "JUSTIFY" in align_str or "BOTH" in align_str:
                    parts.append("两端对齐")
                elif "LEFT" in align_str:
                    parts.append("左对齐")

            # 首行缩进
            if s.get("first_line_indent"):
                parts.append(f"首行缩进{s['first_line_indent']:.1f}cm")

            summary[label] = "，".join(parts) if parts else "（未检测到具体样式设置）"

        return summary

    def get_statistics(self) -> Dict:
        """获取统计信息"""
        templates = list(self._template_index.values())
        
        return {
            "total": len(templates),
            "categories": {
                "universities": len([t for t in templates if t.category == "universities"]),
                "journals": len([t for t in templates if t.category == "journals"]),
                "custom": len([t for t in templates if t.category == "custom"])
            },
            "most_used": sorted(templates, key=lambda x: x.usage_count, reverse=True)[:5],
            "cache_size": len(self._config_cache)
        }
