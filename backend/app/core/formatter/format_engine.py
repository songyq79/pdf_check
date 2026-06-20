# backend/app/core/formatter/format_engine.py
"""
格式化引擎 - 整合所有组件
"""

from docx import Document
from pathlib import Path
from typing import Callable, Dict, Optional
from loguru import logger
from datetime import datetime
import os
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from lxml import etree

from .template_manager import TemplateManager
from .structure_analyzer import StructureAnalyzer
from .style_applicator import StyleApplicator
from app.services.docx_normalizer import clean_control_chars, normalize_docx


class FormatEngine:
    """
    企业级格式化引擎
    
    整合所有格式化组件，提供统一接口
    """
    
    def __init__(self, template_dir: Path, use_ai: bool = False):
        self.template_manager = TemplateManager(
            builtin_dir=template_dir / "builtin",
            user_dir=template_dir / "user"
        )
        
        self.analyzer = StructureAnalyzer(use_ai=use_ai)
        
        logger.info(f"格式化引擎初始化完成 (AI: {use_ai})")
    
    # ZIP 层样式注入：把模板的 styles.xml / theme.xml 覆盖进用户文档，
    # numbering.xml 走合并策略（保留用户 numId 避免列表编号错位）。
    _INJECT_REPLACE_PARTS = ('word/styles.xml', 'word/theme/theme1.xml')
    _INJECT_MERGE_PARTS = ('word/numbering.xml',)
    _W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    def _merge_numbering(self, user_xml: bytes, template_xml: bytes) -> bytes:
        """
        合并 numbering.xml：
        - 用户 XML 为主（保留用户文档里已使用的 numId / abstractNumId 定义）
        - 模板中用户没有的 abstractNum / num 节点追加进来
        - ID 冲突时保留用户版本，避免已有编号列表错位
        """
        try:
            ns = {'w': self._W_NS}
            w = f'{{{self._W_NS}}}'
            user_root = etree.fromstring(user_xml)
            tpl_root = etree.fromstring(template_xml)

            user_num_ids = {
                el.get(f'{w}numId') for el in user_root.findall('w:num', ns)
            }
            user_abs_ids = {
                el.get(f'{w}abstractNumId')
                for el in user_root.findall('w:abstractNum', ns)
            }

            # 追加模板独有的 abstractNum（必须在 num 之前）
            for abs_num in tpl_root.findall('w:abstractNum', ns):
                if abs_num.get(f'{w}abstractNumId') not in user_abs_ids:
                    user_root.append(abs_num)

            # 追加模板独有的 num
            for num in tpl_root.findall('w:num', ns):
                if num.get(f'{w}numId') not in user_num_ids:
                    user_root.append(num)

            return etree.tostring(
                user_root,
                xml_declaration=True,
                encoding='UTF-8',
                standalone=True,
            )
        except Exception as e:
            logger.warning(f"  numbering.xml 合并失败，保留用户原文件: {e}")
            return user_xml

    def _inject_template_parts(
        self,
        input_path: str,
        template_parts: Dict[str, bytes],
    ) -> str:
        """
        把模板的 styles.xml / theme.xml / numbering.xml 注入用户文档，
        返回注入后的临时 docx 路径（调用方负责加载后清理）。
        """
        if not template_parts:
            logger.info("  模板无 XML 部件可注入，跳过")
            return input_path

        tmp_fd, tmp_out = tempfile.mkstemp(suffix='.docx', prefix='injected_')
        os.close(tmp_fd)

        replaced, merged, skipped_orphan = [], [], []
        with zipfile.ZipFile(input_path, 'r') as zin:
            user_names = set(zin.namelist())
            with zipfile.ZipFile(tmp_out, 'w', zipfile.ZIP_DEFLATED) as zout:
                for name in zin.namelist():
                    if name in self._INJECT_REPLACE_PARTS and name in template_parts:
                        zout.writestr(name, template_parts[name])
                        replaced.append(name)
                    elif name in self._INJECT_MERGE_PARTS and name in template_parts:
                        merged_bytes = self._merge_numbering(
                            zin.read(name), template_parts[name]
                        )
                        zout.writestr(name, merged_bytes)
                        merged.append(name)
                    else:
                        zout.writestr(name, zin.read(name))

            # 用户文档没有的部件不能直接追加（[Content_Types].xml / rels 没同步 →
            # Word 会判定文件损坏）。记录下来便于观察，后续 Phase 3 可做完整注入。
            for name in template_parts:
                if name not in user_names:
                    skipped_orphan.append(name)

        if skipped_orphan:
            logger.info(
                f"  样式注入完成 replace={replaced} merge={merged} "
                f"skip_orphan={skipped_orphan}（用户文档未包含该部件，跳过以免损坏）"
            )
        else:
            logger.info(
                f"  样式注入完成 replace={replaced} merge={merged}"
            )
        return tmp_out

    def _extract_text_fallback(self, input_path: str) -> str:
        """
        降级方案：从损坏的文档中提取纯文本
        
        尝试多种方法提取文本内容：
        1. 直接解析 ZIP 中的 document.xml
        2. 使用 zipfile 提取所有文本节点
        """
        try:
            # 方法1：解析 document.xml
            with zipfile.ZipFile(input_path, 'r') as docx_zip:
                # 读取主文档 XML
                xml_content = docx_zip.read('word/document.xml')
                root = ET.fromstring(xml_content)
                
                # 提取所有文本节点
                namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                text_elements = root.findall('.//w:t', namespace)
                
                text_content = '\n'.join([elem.text for elem in text_elements if elem.text])
                
                if text_content.strip():
                    text_content = clean_control_chars(text_content)
                    logger.info(f"  降级提取成功: {len(text_content)} 字符")
                    return text_content

        except Exception as e:
            logger.warning(f"  降级提取失败: {e}")

        return ""
    
    def _create_simple_document(self, text_content: str, template_id: str, output_path: str) -> Dict:
        """
        降级方案：创建简化的格式化文档
        
        只保留文本内容，应用基本的模板样式
        """
        try:
            # 创建新文档
            doc = Document()
            
            # 加载模板配置
            template_config = self.template_manager.extract_config(template_id)
            template_meta = self.template_manager.get_template(template_id)
            
            # 按段落分割文本
            paragraphs = [p.strip() for p in text_content.split('\n') if p.strip()]
            
            # 添加段落到文档
            for para_text in paragraphs:
                doc.add_paragraph(para_text)
            
            # 应用基本样式（如果模板支持）
            try:
                applicator = StyleApplicator(template_config)
                simple_sections = [{'index': i, 'type': 'body'} for i in range(len(paragraphs))]
                doc = applicator.apply(doc, simple_sections)
            except Exception as e:
                logger.warning(f"  样式应用失败（使用默认样式）: {e}")
            
            # 保存文档
            doc.save(output_path)
            
            logger.info(f"  简化文档生成成功: {len(paragraphs)} 个段落")
            
            return {
                'success': True,
                'fallback_mode': True,
                'paragraphs': len(paragraphs),
                'template': {
                    'id': template_id,
                    'name': template_meta.name
                }
            }
            
        except Exception as e:
            logger.error(f"  简化文档生成失败: {e}")
            raise
    
    def format_document(
        self,
        input_path: str,
        output_path: str,
        template_id: str,
        options: Optional[Dict] = None,
        progress_cb: Optional[Callable[[int, str], None]] = None,
    ) -> Dict:
        """
        格式化文档（主入口）
        
        Args:
            input_path: 输入文档路径
            output_path: 输出文档路径
            template_id: 模板ID
            options: 可选配置
        
        Returns:
            格式化结果
        """
        start_time = datetime.now()
        
        logger.info("=" * 60)
        logger.info(f"开始格式化")
        logger.info(f"输入: {input_path}")
        logger.info(f"模板: {template_id}")
        logger.info("=" * 60)
        
        def _report(progress: int, stage: str):
            if progress_cb:
                try:
                    progress_cb(progress, stage)
                except Exception:
                    pass

        injected_path = None
        normalized_path = None
        try:
            _report(5, "加载模板")
            # 1. 归一化 + 提取模板 XML 部件
            logger.info("[1/6] 加载模板与预处理...")
            _orig = input_path
            input_path = str(normalize_docx(input_path))
            normalized_path = input_path if input_path != _orig else None
            template_config = self.template_manager.extract_config(template_id)
            template_meta = self.template_manager.get_template(template_id)
            try:
                template_parts = self.template_manager.get_template_xml_parts(template_id)
            except Exception as e:
                logger.warning(f"  模板 XML 部件提取失败，跳过样式注入: {e}")
                template_parts = {}
            logger.info(f"  模板: {template_meta.name}")

            _report(20, "注入模板样式")
            # 2. 样式注入（ZIP 层替换 styles.xml / theme.xml，合并 numbering.xml）
            logger.info("[2/6] 注入模板样式...")
            doc = None
            try:
                injected_path = self._inject_template_parts(input_path, template_parts)
                work_path = injected_path if injected_path != input_path else input_path
                # 注入后立即加载校验：能被 python-docx 正常打开才算成功
                doc = Document(work_path)
                _ = len(doc.paragraphs)  # 触发 lazy load
            except Exception as e:
                logger.warning(f"  样式注入失败，降级用原文档 + 模板候选样式: {e}")
                if injected_path and injected_path != input_path:
                    try:
                        Path(injected_path).unlink(missing_ok=True)
                    except Exception:
                        pass
                injected_path = None
                doc = None

            _report(35, "加载文档")
            # 3. 加载文档（若注入路径未加载成功，回退到原文档）
            logger.info("[3/6] 加载文档...")
            if doc is None:
                doc = Document(input_path)
            para_count = len(doc.paragraphs)
            logger.info(f"  文档加载成功: {para_count} 个段落")

            _report(55, "识别结构")
            # 4. 识别结构
            logger.info("[4/6] 识别结构...")
            analysis_result = self.analyzer.analyze(doc)

            logger.info(
                f"  识别完成: {len(analysis_result['sections'])} 个章节, "
                f"质量: {analysis_result['quality']:.2%}"
            )

            _report(75, "应用样式")
            # 5. 应用样式（清 direct formatting + 打 pStyle + 多 section 页面设置）
            logger.info("[5/6] 应用样式...")
            applicator = StyleApplicator(template_config)
            doc = applicator.apply(doc, analysis_result['sections'])
            
            logger.info(
                f"  应用完成: 成功 {applicator.applied_count}, "
                f"失败 {applicator.error_count}"
            )

            _report(95, "保存文档")
            # 6. 保存文档
            logger.info("[6/6] 保存文档...")
            doc.save(output_path)
            _report(100, "完成")
            
            # 计算耗时
            time_elapsed = (datetime.now() - start_time).total_seconds()
            
            logger.info("=" * 60)
            logger.info(f"✅ 格式化完成！耗时: {time_elapsed:.2f}秒")
            logger.info(f"输出: {output_path}")
            logger.info("=" * 60)
            
            return {
                'success': True,
                'structure': analysis_result,
                'template': {
                    'id': template_id,
                    'name': template_meta.name
                },
                'stats': {
                    'paragraphs': para_count,
                    'sections': len(analysis_result['sections']),
                    'applied': applicator.applied_count,
                    'errors': applicator.error_count,
                    'quality': analysis_result['quality']
                },
                'time_elapsed': time_elapsed,
                'output_path': output_path
            }
        
        except Exception as e:
            error_msg = str(e)
            logger.error(f"❌ 格式化失败: {error_msg}")
            
            # 判断是否是文档格式问题
            if "no item named" in error_msg.lower() or "badzip" in error_msg.lower():
                logger.warning("⚠️ 检测到文档格式异常，尝试降级处理...")
                
                try:
                    # 尝试降级处理：提取文本 + 简化格式化
                    text_content = self._extract_text_fallback(input_path)
                    
                    if text_content:
                        result = self._create_simple_document(text_content, template_id, output_path)
                        time_elapsed = (datetime.now() - start_time).total_seconds()
                        
                        logger.warning("=" * 60)
                        logger.warning(f"⚠️ 降级处理完成！耗时: {time_elapsed:.2f}秒")
                        logger.warning(f"输出: {output_path}")
                        logger.warning("=" * 60)
                        
                        return {
                            'success': True,
                            'warning': '文档格式异常，已使用简化模式处理（可能丢失部分格式、图片等内容）',
                            'fallback_mode': True,
                            'stats': {
                                'paragraphs': result.get('paragraphs', 0),
                                'sections': 0,
                                'applied': 0,
                                'errors': 0,
                                'quality': 0.0
                            },
                            'template': result.get('template', {}),
                            'time_elapsed': time_elapsed,
                            'output_path': output_path
                        }
                    else:
                        # 文本提取也失败
                        logger.error("❌ 降级处理失败：无法提取文本内容")
                        
                except Exception as fallback_error:
                    logger.error(f"❌ 降级处理失败: {fallback_error}")
                
                # 降级处理失败，返回友好错误
                return {
                    'success': False,
                    'error': '文档格式异常且无法自动修复',
                    'detail': '文档内部结构损坏或格式不正确',
                    'suggestions': [
                        '请在 Word 中打开文档，选择"另存为"，保存为新的 .docx 文件',
                        '确保文档能在 Word 中正常打开',
                        '尝试删除文档中的特殊对象（如嵌入的 OLE 对象、损坏的图片等）'
                    ],
                    'time_elapsed': (datetime.now() - start_time).total_seconds()
                }
            
            # 其他类型的错误，直接返回
            return {
                'success': False,
                'error': error_msg,
                'time_elapsed': (datetime.now() - start_time).total_seconds()
            }
        finally:
            if normalized_path:
                try:
                    Path(normalized_path).unlink(missing_ok=True)
                except Exception:
                    pass
            if injected_path and injected_path != input_path:
                try:
                    Path(injected_path).unlink(missing_ok=True)
                except Exception:
                    pass
    
    def preview_structure(self, input_path: str) -> Dict:
        """预览文档结构（不格式化）"""
        _orig = input_path
        input_path = str(normalize_docx(input_path))
        normalized_path = input_path if input_path != _orig else None
        try:
            doc = Document(input_path)
            return self.analyzer.analyze(doc)
        finally:
            if normalized_path:
                try:
                    Path(normalized_path).unlink(missing_ok=True)
                except Exception:
                    pass
    
    def validate_template(self, template_file: str) -> Dict:
        """验证模板文件"""
        return self.template_manager.validate_template(template_file)
    
    def get_template_list(self, category: str = None) -> list:
        """获取模板列表"""
        return [
            {
                'id': t.template_id,
                'name': t.name,
                'category': t.category,
                'school': t.school_or_journal,
                'usage_count': t.usage_count
            }
            for t in self.template_manager.list_templates(category=category)
        ]
