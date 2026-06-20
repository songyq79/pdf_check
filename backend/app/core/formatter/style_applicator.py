# backend/app/core/formatter/style_applicator.py
"""
企业级样式应用器
支持复杂的样式映射和批量应用
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree
from typing import Dict, List, Optional
from loguru import logger

from .structure_analyzer import SectionInfo, SectionType
from .template_manager import TemplateConfig


class StyleApplicator:
    """
    样式应用器
    
    功能：
    - 页面设置应用
    - 字体样式应用  
    - 段落格式应用
    - 编号应用
    - 页眉页脚应用
    """
    
    def __init__(self, config: TemplateConfig):
        self.config = config
        self.applied_count = 0
        self.error_count = 0
    
    def apply(self, doc: Document, sections: List[Dict]) -> Document:
        """
        应用样式到文档
        
        Args:
            doc: Word文档对象
            sections: 识别的章节列表
        """
        logger.info("开始应用样式...")
        
        # 1. 页面设置
        self._apply_page_setup(doc)
        
        # 2. 批量应用段落样式
        for section in sections:
            try:
                para = doc.paragraphs[section['index']]
                self._apply_paragraph_style(para, section)
                self.applied_count += 1
            except Exception as e:
                logger.error(f"应用样式失败 [{section['index']}]: {e}")
                self.error_count += 1

        # 2b. 目录段落行距单独处理（TOC 段落用 Word 内置 TOC 样式名识别）
        self._apply_toc_spacing(doc, sections)
        
        # 3. 应用编号
        self._apply_numbering(doc, sections)
        
        # 4. 页眉页脚
        if self.config.header or self.config.footer:
            self._apply_header_footer(doc)

        # 5. 正文引用标注转上标
        self._apply_citation_superscript(doc, sections)

        # 6. 图题/表题/公式居中对齐（caption 样式兜底）
        self._apply_caption_format(doc, sections)

        # 7. 参考文献国标格式整理
        self._apply_reference_format(doc, sections)
        
        logger.info(f"样式应用完成: 成功 {self.applied_count}, 失败 {self.error_count}")
        
        return doc
    
    def _apply_page_setup(self, doc: Document):
        """应用页面设置（遍历所有 section，处理多节文档如封面/目录/正文分节符）"""
        n = 0
        for section in doc.sections:
            section.page_width = Cm(self.config.page_width)
            section.page_height = Cm(self.config.page_height)
            section.top_margin = Cm(self.config.margin_top)
            section.bottom_margin = Cm(self.config.margin_bottom)
            section.left_margin = Cm(self.config.margin_left)
            section.right_margin = Cm(self.config.margin_right)
            n += 1

        logger.info(
            f"页面设置({n} sections): {self.config.page_width}×{self.config.page_height}cm, "
            f"边距 T{self.config.margin_top}/B{self.config.margin_bottom}/"
            f"L{self.config.margin_left}/R{self.config.margin_right}cm"
        )

    # 章节类型 → Word 样式名候选列表（按优先级）
    # 样式注入后，doc.styles 里是模板的样式，所以候选名应覆盖模板可能使用的命名。
    # 同时兼容中英文 Word / WPS 的不同命名习惯。
    # 重要：中文模板里 "标题 1" 通常是被模板定制过的样式（黑体、特定字号、颜色自动），
    #       而 "Heading 1" 往往只是 Word 内置的默认样式（Calibri/宋体、默认字号）。
    #       因此中文样式名必须排在英文样式名前面，否则导出文档会套用未定制的英文样式。
    _STYLE_ALIASES = {
        'title':         ['论文标题', '标题', 'Title', 'title'],
        'chapter':       ['标题 1', '标题1', '一级标题', '章标题',
                          'Heading 1', 'heading1', 'Heading1'],
        'section_1':     ['标题 2', '标题2', '二级标题',
                          'Heading 2', 'heading2', 'Heading2'],
        'section_2':     ['标题 3', '标题3', '三级标题',
                          'Heading 3', 'heading3', 'Heading3'],
        'section_3':     ['标题 4', '标题4', '四级标题',
                          'Heading 4', 'heading4', 'Heading4'],
        'abstract_cn':   ['中文摘要', '摘要', 'Abstract', 'abstract'],
        'abstract_en':   ['英文摘要', 'English Abstract', 'Abstract', 'abstract'],
        'keywords_cn':   ['中文关键词', '关键词', 'Keywords', '正文', 'Normal'],
        'keywords_en':   ['英文关键词', 'Keywords', '正文', 'Normal'],
        'body':          ['正文', 'Normal', '正文首行缩进', 'body', 'Body Text'],
        'references':    ['参考文献', 'References', 'Bibliography', 'references'],
        'toc':           ['目录 1', '目录', 'TOC 1', 'toc'],
        'figure':        ['图题', '题注', 'Caption', 'caption'],
        'table':         ['表题', '题注', 'Caption', 'caption'],
        'formula':       ['公式', '题注', 'Caption', 'caption'],
        'appendix':      ['附录', '标题 1', 'Heading 1', '正文', 'Normal'],
        'acknowledgement': ['致谢', '标题 1', 'Heading 1', '正文', 'Normal'],
    }

    # run 级 direct formatting 里要清掉的子节点名
    # （这些直接格式化会覆盖 pStyle 样式定义，必须清除才能让样式注入生效）
    _CLEAR_TAGS = {
        'sz', 'szCs', 'rFonts', 'b', 'bCs', 'i', 'iCs', 'u', 'color',
    }

    def _resolve_style_name(self, doc: Document, section_type: str) -> Optional[str]:
        """按候选列表查文档里实际存在的样式名，命中第一个就返回"""
        candidates = self._STYLE_ALIASES.get(section_type, ['Normal', '正文'])
        available = {s.name for s in doc.styles}
        for name in candidates:
            if name in available:
                return name
        return None

    def _clear_direct_formatting(self, paragraph):
        """
        清除段落所有 run 的 direct formatting（字体/字号/加粗/斜体/下划线/颜色）。

        为什么必须清：Word 的渲染优先级是 run direct formatting > pStyle 样式定义。
        用户段落里的 run 如果有直接设置的 font.size/font.name 等，会盖过我们
        注入的模板样式定义。清完再打 pStyle，样式才会真正生效。

        保留：superscript/subscript 等特殊属性（引用上标需要）。
        """
        for run in paragraph.runs:
            rPr = run._element.find(qn('w:rPr'))
            if rPr is None:
                continue
            for child in list(rPr):
                if etree.QName(child).localname in self._CLEAR_TAGS:
                    rPr.remove(child)

    def _apply_paragraph_style(self, paragraph, section: Dict):
        """
        打 pStyle 标签应用样式：
          1. 清除 run 的 direct formatting
          2. 把段落的 pStyle 指向目标样式
        Word 渲染时会自动从模板样式定义取字体/字号/行距/对齐等。
        """
        section_type = section['type']

        # 模板没有目录时，保留用户文档的原目录样式（标题和条目都不动）
        if section_type == 'toc' and not getattr(self.config, 'has_toc', False):
            return

        doc = paragraph.part.document
        style_name = self._resolve_style_name(doc, section_type)

        if style_name is None:
            logger.debug(
                f"无匹配样式 section_type={section_type}, "
                f"跳过段落 idx={section.get('index')}"
            )
            return

        try:
            self._clear_direct_formatting(paragraph)
            paragraph.style = doc.styles[style_name]
        except Exception as e:
            logger.warning(
                f"打 pStyle 失败 idx={section.get('index')} "
                f"type={section_type} style={style_name}: {e}"
            )
            raise
    
    # 已有编号的检测模式（防止对已编号段落二次添加）
    _NUMBERED_CHAPTER = re.compile(r'^(第[一二三四五六七八九十百千\d]+章|Chapter\s+\d+)', re.IGNORECASE)
    _NUMBERED_SECTION1 = re.compile(r'^(\d+[\.、\s]|[一二三四五六七八九十]+[、])')
    _NUMBERED_SECTION2 = re.compile(r'^\d+\.\d+')

    def _apply_numbering(self, doc: Document, sections: List[Dict]):
        """应用编号（保留原 run 的字体格式，不用 para.text= 整体覆盖）"""
        logger.info("应用编号...")

        chapter_num = 0
        section_nums = {}

        def _prepend_numbering(para, prefix: str):
            """把 prefix 拼到段落第一个 run 前面，保留所有 run 的样式"""
            if not prefix:
                return
            runs = list(para.runs)
            if runs:
                runs[0].text = f"{prefix} {runs[0].text}"
            else:
                para.add_run(prefix)

        for section in sections:
            stype = section['type']
            para = doc.paragraphs[section['index']]
            text = para.text.strip()

            if stype == 'chapter':
                chapter_num += 1
                section_nums = {1: chapter_num, 2: 0, 3: 0}

                # 只在段落没有自带编号（pattern识别）且文本未已含章节编号时才前插
                if not section.get('numbering') and not self._NUMBERED_CHAPTER.match(text):
                    fmt = self.config.numbering.get('chapter', '第{n}章')
                    _prepend_numbering(para, fmt.format(n=chapter_num))

            elif stype == 'section_1':
                section_nums[2] = section_nums.get(2, 0) + 1
                section_nums[3] = 0

                if not section.get('numbering') and not self._NUMBERED_SECTION1.match(text):
                    fmt = self.config.numbering.get('section1', '{n}')
                    _prepend_numbering(para, fmt.format(n=section_nums[2]))

            elif stype == 'section_2':
                section_nums[3] = section_nums.get(3, 0) + 1

                if not section.get('numbering') and not self._NUMBERED_SECTION2.match(text):
                    fmt = self.config.numbering.get('section2', '{n}.{m}')
                    _prepend_numbering(para, fmt.format(n=section_nums[2], m=section_nums[3]))
    
    def _apply_header_footer(self, doc: Document):
        """应用页眉页脚（遍历所有 section，支持多节文档）"""
        for section in doc.sections:
            # 页眉
            if self.config.header:
                header = section.header
                para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                para.text = self.config.header.get('text', '')
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 页脚
            if self.config.footer:
                footer = section.footer
                para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

                if self.config.footer.get('page_number'):
                    self._add_page_number(para)
                else:
                    para.text = self.config.footer.get('text', '')

                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_page_number(self, paragraph):
        """添加页码"""
        run = paragraph.add_run()
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def _apply_citation_superscript(self, doc: Document, sections: List[Dict]):
        """
        将正文中的引用标注 [1]、[1,2]、[1-3] 转为上标。
        保留段落内的其他格式（斜体、加粗、颜色等）。
        正确处理同一 run 内含多个引用标注的情况。
        """
        from copy import deepcopy
        from collections import defaultdict

        # 参考文献起始段落索引：从该段起（含该段）全部跳过，
        # 避免将参考文献条目开头的 [1] 误转为上标
        ref_start = min((s['index'] for s in sections if s.get('type') == 'references'), default=None)
        citation_pattern = re.compile(r'\[\d+(?:[,，\-–]\d+)*\]')

        def _ensure_rPr(run_elem):
            rPr = run_elem.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                run_elem.insert(0, rPr)
            return rPr

        def _set_superscript(rPr_elem):
            for va in rPr_elem.findall(qn('w:vertAlign')):
                rPr_elem.remove(va)
            va = OxmlElement('w:vertAlign')
            va.set(qn('w:val'), 'superscript')
            rPr_elem.append(va)

        def _clear_superscript(rPr_elem):
            for va in rPr_elem.findall(qn('w:vertAlign')):
                rPr_elem.remove(va)

        converted = 0
        for idx, para in enumerate(doc.paragraphs):
            # 跳过参考文献节及其后所有段落（含条目中的 [1] [2] 序号）
            if ref_start is not None and idx >= ref_start:
                continue
            if not para.text or not para.runs:
                continue
            if not citation_pattern.search(para.text):
                continue

            # 构建 run 区间映射（基于 para.text 的原始位置）
            run_ranges = {}
            pos = 0
            for ri, run in enumerate(para.runs):
                run_ranges[ri] = (pos, pos + len(run.text))
                pos += len(run.text)

            # 将所有引用按所属 run 分组
            citations_by_run = defaultdict(list)
            for m in citation_pattern.finditer(para.text):
                cs, ce = m.start(), m.end()
                for ri, (rs, re_) in run_ranges.items():
                    if rs <= cs and ce <= re_:
                        citations_by_run[ri].append((cs - rs, ce - rs, m.group(0)))
                        break
                # 跨 run 的引用（极少见）：跳过

            para_element = para._element
            # 在修改 DOM 之前快照 runs 列表，避免插入新元素后索引漂移
            all_runs = list(para.runs)

            for ri in sorted(citations_by_run.keys()):
                run = all_runs[ri]
                original_text = run.text
                # 保存原始 XML 供后续新 run 复制格式
                original_elem = deepcopy(run._element)

                cites = sorted(citations_by_run[ri], key=lambda x: x[0])

                # 整个 run 只有一个引用且占满 run
                if len(cites) == 1 and cites[0][0] == 0 and cites[0][1] == len(original_text):
                    _set_superscript(_ensure_rPr(run._element))
                    continue

                # 将 run 文本分割为 [(text, is_superscript)] 段
                segments = []
                last = 0
                for (c_start, c_end, c_text) in cites:
                    if last < c_start:
                        segments.append((original_text[last:c_start], False))
                    segments.append((c_text, True))
                    last = c_end
                if last < len(original_text):
                    segments.append((original_text[last:], False))

                # 修改原 run 为第一段
                first_text, first_super = segments[0]
                run.text = first_text
                if first_super:
                    _set_superscript(_ensure_rPr(run._element))
                else:
                    _clear_superscript(_ensure_rPr(run._element))

                # 在原 run 之后依次插入后续段
                insert_pos = list(para_element).index(run._element) + 1
                for i, (seg_text, seg_super) in enumerate(segments[1:]):
                    new_elem = deepcopy(original_elem)
                    w_t = new_elem.find(qn('w:t'))
                    if w_t is None:
                        w_t = OxmlElement('w:t')
                        new_elem.append(w_t)
                    w_t.text = seg_text
                    rPr = _ensure_rPr(new_elem)
                    if seg_super:
                        _set_superscript(rPr)
                    else:
                        _clear_superscript(rPr)
                    para_element.insert(insert_pos + i, new_elem)

            converted += 1

        if converted:
            logger.info(f"引用标注上标转换: {converted} 个段落")

    def _apply_caption_format(self, doc: Document, sections: List[Dict]):
        """
        对图题、表题、公式段落应用格式。
        优先从模板 caption 样式读取配置，没有则用默认值：
        - 居中对齐
        - 字体同正文（从 body 样式继承）
        - 段前 6pt，段后 6pt
        """
        caption_types = {'figure', 'table', 'formula'}
        caption_indices = {s['index'] for s in sections if s.get('type') in caption_types}

        if not caption_indices:
            return

        # 从模板配置读 caption 样式，没有则用默认
        caption_cfg = self.config.styles.get('caption', self.config.styles.get('Caption', {}))
        body_cfg = self.config.styles.get('body', self.config.styles.get('Normal', {}))

        font_name = caption_cfg.get('font_name') or body_cfg.get('font_name')
        font_size = caption_cfg.get('font_size') or body_cfg.get('font_size')
        line_spacing = caption_cfg.get('line_spacing') or body_cfg.get('line_spacing')
        line_spacing_rule = caption_cfg.get('line_spacing_rule') or body_cfg.get('line_spacing_rule')
        space_before = caption_cfg.get('space_before', self.config.default_caption_space_before)
        space_after  = caption_cfg.get('space_after',  self.config.default_caption_space_after)

        applied = 0
        for idx in caption_indices:
            try:
                para = doc.paragraphs[idx]
                pf = para.paragraph_format
                pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pf.space_before = Pt(space_before)
                pf.space_after  = Pt(space_after)
                if line_spacing:
                    # 应用行距规则，确保非标准值（如 1.3）被正确保存
                    if line_spacing_rule:
                        try:
                            rule_map = {
                                'SINGLE': WD_LINE_SPACING.SINGLE,
                                'ONE_POINT_FIVE': WD_LINE_SPACING.ONE_POINT_FIVE,
                                'DOUBLE': WD_LINE_SPACING.DOUBLE,
                                'AT_LEAST': WD_LINE_SPACING.AT_LEAST,
                                'EXACTLY': WD_LINE_SPACING.EXACTLY,
                                'MULTIPLE': WD_LINE_SPACING.MULTIPLE,
                            }
                            rule_str = str(line_spacing_rule).split('.')[-1].split()[0]
                            pf.line_spacing_rule = rule_map.get(rule_str, WD_LINE_SPACING.MULTIPLE)
                        except Exception as e:
                            logger.warning(f"图表行距规则设置失败: {e}")
                            if line_spacing not in (1.0, 1.5, 2.0):
                                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    else:
                        # 根据值自动判断
                        if line_spacing not in (1.0, 1.5, 2.0):
                            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

                    pf.line_spacing = line_spacing

                for run in para.runs:
                    if font_name:
                        run.font.name = font_name
                        # 安全设置中文字体（rPr/rFonts 可能不存在）
                        rPr = run._element.find(qn('w:rPr'))
                        if rPr is None:
                            rPr = OxmlElement('w:rPr')
                            run._element.insert(0, rPr)
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is None:
                            rFonts = OxmlElement('w:rFonts')
                            rPr.insert(0, rFonts)
                        rFonts.set(qn('w:eastAsia'), font_name)
                    if font_size:
                        run.font.size = Pt(font_size)

                applied += 1
            except Exception as e:
                logger.warning(f"图表格式应用失败 [{idx}]: {e}")

        if applied:
            logger.info(f"图题/表题/公式格式已应用: {applied} 个段落")

    def _apply_toc_spacing(self, doc: Document, sections: List[Dict]):
        """
        对目录段落应用行距。
        优先从模板配置的 toc 样式读取 line_spacing，
        没有则默认 1.5 倍行距，段前段后各 0pt。
        同时兼容两种来源：
          1. structure_analyzer 识别出的 toc 类型段落
          2. 段落样式名包含 TOC / 目录 的 Word 内置目录样式

        如果模板本身没有目录（has_toc=False），则保留用户文档的原目录样式不动。
        """
        if not getattr(self.config, 'has_toc', False):
            logger.info("模板无目录，跳过目录格式覆盖，保留用户文档原目录样式")
            return

        toc_style_config = self.config.styles.get('toc', {})
        line_spacing = toc_style_config.get('line_spacing') or self.config.default_toc_line_spacing
        line_spacing_rule = toc_style_config.get('line_spacing_rule') or self.config.default_toc_line_spacing_rule
        space_before = toc_style_config.get('space_before', 0)
        space_after  = toc_style_config.get('space_after',  0)

        toc_indices = {s['index'] for s in sections if s.get('type') == 'toc'}

        applied = 0
        for idx, para in enumerate(doc.paragraphs):
            style_name = para.style.name if para.style else ''
            is_toc = (
                idx in toc_indices
                or 'TOC' in style_name.upper()
                or '目录' in style_name
            )
            if not is_toc:
                continue

            pf = para.paragraph_format
            # 设置行距规则，确保倍数模式下行距被正确解释
            if line_spacing_rule:
                try:
                    rule_map = {
                        'SINGLE': WD_LINE_SPACING.SINGLE,
                        'ONE_POINT_FIVE': WD_LINE_SPACING.ONE_POINT_FIVE,
                        'DOUBLE': WD_LINE_SPACING.DOUBLE,
                        'AT_LEAST': WD_LINE_SPACING.AT_LEAST,
                        'EXACTLY': WD_LINE_SPACING.EXACTLY,
                        'MULTIPLE': WD_LINE_SPACING.MULTIPLE,
                    }
                    rule_str = str(line_spacing_rule).split('.')[-1].split()[0]
                    pf.line_spacing_rule = rule_map.get(rule_str, WD_LINE_SPACING.MULTIPLE)
                except Exception as e:
                    logger.warning(f"目录行距规则设置失败: {e}，根据值自动判断")
                    if line_spacing in (1.0, 1.5, 2.0):
                        pf.line_spacing = line_spacing  # Word 自动优化
                    else:
                        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                        pf.line_spacing = line_spacing
            else:
                # 根据值自动判断
                if line_spacing in (1.0, 1.5, 2.0):
                    pf.line_spacing = line_spacing  # Word 自动优化
                else:
                    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    pf.line_spacing = line_spacing
            if space_before is not None:
                pf.space_before = Pt(space_before)
            if space_after is not None:
                pf.space_after = Pt(space_after)
            applied += 1

        if applied:
            logger.info(f"目录行距已应用: {applied} 个段落，行距={line_spacing}，规则={line_spacing_rule}")

    def _apply_reference_format(self, doc: Document, sections: List[Dict]):
        """
        按 GB/T 7714 国标格式整理参考文献段落。
        策略：用正则识别各字段，重新拼装为标准格式。
        识别不了的条目保持原样不动。

        支持的文献类型：
          [J] 期刊  [M] 专著  [C] 会议  [D] 学位论文
          [R] 报告  [N] 报纸  [EB/OL] 网络资源
        """
        ref_indices = [s['index'] for s in sections if s.get('type') == 'references']
        if not ref_indices:
            return

        seq_pat = re.compile(r'^\s*\[(\d+)\]\s*')

        journal_pat = re.compile(
            r'^(?P<authors>.+?)\.\s*'
            r'(?P<title>.+?)'
            r'(?:\[J\])?\s*[\.。]\s*'
            r'(?P<journal>[^,，]+?)\s*[,，]\s*'
            r'(?P<year>\d{4})\s*[,，]\s*'
            r'(?P<vol_issue>[^:：]+?)\s*[:：]\s*'
            r'(?P<pages>[\d\-]+)\s*[\.。]?\s*$'
        )

        book_pat = re.compile(
            r'^(?P<authors>.+?)\.\s*'
            r'(?P<title>.+?)'
            r'(?:\[M\])?\s*[\.。]\s*'
            r'(?P<place>[^:：]+?)\s*[:：]\s*'
            r'(?P<publisher>[^,，]+?)\s*[,，]\s*'
            r'(?P<year>\d{4})\s*[\.。]?\s*$'
        )

        thesis_pat = re.compile(
            r'^(?P<authors>.+?)\.\s*'
            r'(?P<title>.+?)'
            r'(?:\[D\])?\s*[\.。]\s*'
            r'(?P<place>[^:：]+?)\s*[:：]\s*'
            r'(?P<school>[^,，]+?)\s*[,，]\s*'
            r'(?P<year>\d{4})\s*[\.。]?\s*$'
        )

        ebol_pat = re.compile(
            r'^(?P<authors>.+?)\.\s*'
            r'(?P<title>.+?)'
            r'(?:\[EB/OL\])?\s*[\.。]\s*'
            r'(?P<url>https?://\S+)\s*[,，]?\s*'
            r'(?P<year>\d{4})\s*[\.。]?\s*$'
        )

        def _reformat(raw: str) -> str:
            m_seq = seq_pat.match(raw)
            seq = ''
            body = raw
            if m_seq:
                seq = f'[{m_seq.group(1)}] '
                body = raw[m_seq.end():]

            m = journal_pat.match(body)
            if m:
                return (f"{seq}{m.group('authors')}. "
                        f"{m.group('title')}[J]. "
                        f"{m.group('journal')}, "
                        f"{m.group('year')}, "
                        f"{m.group('vol_issue')}: "
                        f"{m.group('pages')}.")

            m = book_pat.match(body)
            if m:
                return (f"{seq}{m.group('authors')}. "
                        f"{m.group('title')}[M]. "
                        f"{m.group('place')}: "
                        f"{m.group('publisher')}, "
                        f"{m.group('year')}.")

            m = thesis_pat.match(body)
            if m:
                return (f"{seq}{m.group('authors')}. "
                        f"{m.group('title')}[D]. "
                        f"{m.group('place')}: "
                        f"{m.group('school')}, "
                        f"{m.group('year')}.")

            m = ebol_pat.match(body)
            if m:
                return (f"{seq}{m.group('authors')}. "
                        f"{m.group('title')}[EB/OL]. "
                        f"{m.group('url')}, "
                        f"{m.group('year')}.")

            return raw

        reformatted = 0
        for idx in ref_indices:
            try:
                para = doc.paragraphs[idx]
                original = para.text.strip()
                if not original:
                    continue
                new_text = _reformat(original)
                if new_text != original:
                    runs = para.runs
                    if runs:
                        # 保留首个 run 的格式，修改其文本，移除多余 run
                        runs[0].text = new_text
                        for run in runs[1:]:
                            run._element.getparent().remove(run._element)
                    else:
                        para.add_run(new_text)
                    reformatted += 1
            except Exception as e:
                logger.warning(f"参考文献格式化失败 [{idx}]: {e}")

        if reformatted:
            logger.info(f"参考文献国标格式整理: {reformatted} 条")
