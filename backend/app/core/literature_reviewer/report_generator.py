"""
文献综述 Word 报告生成器。
输入为 reviewer.generate_review() 的结果 dict。
参考文献用 gbt7714 格式化（中文 GB/T 7714；英文条目自动 et al）。
"""
from __future__ import annotations

import platform
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from loguru import logger

from app.core.literature_reviewer.gbt7714 import format_references


def _set_document_style(doc: Document):
    system = platform.system()
    if system == "Windows":
        cn_font, fallback = "微软雅黑", "Microsoft YaHei"
    elif system == "Darwin":
        cn_font, fallback = "PingFang SC", "Arial Unicode MS"
    else:
        cn_font, fallback = "Noto Sans CJK SC", "Arial Unicode MS"
    doc.styles["Normal"].font.name = fallback
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), cn_font)
    doc.styles["Normal"].font.size = Pt(10.5)


def _add_title(doc: Document, title: str):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 108, 73)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()


def generate_review_report(result: dict, output_path: Path) -> Path:
    """根据综述结果 dict 生成 Word，返回文件路径。"""
    try:
        doc = Document()
        _set_document_style(doc)

        meta = result.get("meta", {})
        draft = result.get("draft", {})
        categorization = result.get("categorization", {})
        papers = result.get("papers", [])

        title = meta.get("topic") or "文献综述"
        _add_title(doc, f"{title} · 文献综述初稿")

        # 元信息
        p = doc.add_paragraph()
        p.add_run("学科：").font.bold = True
        p.add_run(meta.get("discipline") or "未指定")
        p = doc.add_paragraph()
        p.add_run("文献数量：").font.bold = True
        p.add_run(f"输入 {meta.get('papers_identified', 0)} 篇 / 共 {meta.get('papers_total', 0)} 篇")
        p = doc.add_paragraph()
        p.add_run("生成时间：").font.bold = True
        p.add_run(datetime.now().strftime("%Y年%m月%d日 %H:%M"))
        doc.add_paragraph("_" * 80)

        # 研究现状总览
        if draft.get("overview"):
            doc.add_heading("一、研究现状总览", level=1)
            doc.add_paragraph(draft["overview"])

        # 各小节
        sections = draft.get("sections", [])
        if sections:
            doc.add_heading("二、主要研究方向", level=1)
            for i, sec in enumerate(sections, 1):
                doc.add_heading(f"{i}. {sec.get('title','')}", level=2)
                if sec.get("content"):
                    doc.add_paragraph(sec["content"])

        # 研究空白
        gaps = categorization.get("research_gaps", [])
        if gaps:
            doc.add_heading("三、研究空白与未来方向", level=1)
            for g in gaps:
                doc.add_paragraph(str(g), style="List Bullet")

        # 结论
        if draft.get("conclusion"):
            doc.add_heading("四、结论", level=1)
            doc.add_paragraph(draft["conclusion"])

        # 参考文献（GB/T 7714）
        doc.add_heading("参考文献", level=1)
        if papers:
            refs = format_references(papers)
            for line in refs.split("\n"):
                doc.add_paragraph(line)
        else:
            doc.add_paragraph("（无）")

        # 页脚
        doc.add_paragraph()
        doc.add_paragraph("_" * 80)
        fp = doc.add_paragraph()
        fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        frun = fp.add_run("本综述初稿由 VRonly 自动生成，需作者二次修改完善后方可使用")
        frun.font.size = Pt(9)
        frun.font.color.rgb = RGBColor(128, 128, 128)
        frun.italic = True

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        logger.info(f"[lit.report] 文献综述报告生成: {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"[lit.report] 报告生成失败: {e}")
        raise
