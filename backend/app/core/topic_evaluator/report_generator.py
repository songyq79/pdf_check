"""
选题评估 Word 报告生成器。
镜像 evaluator/report_generator.py 的字体/着色/工厂函数风格，
输入为 evaluator.evaluate_topic() 返回的结果 dict。
"""
from __future__ import annotations

import platform
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from loguru import logger

from app.core.topic_evaluator.prompts import get_type_name


def _set_document_style(doc: Document):
    system = platform.system()
    if system == "Windows":
        cn_font, fallback = "微软雅黑", "Microsoft YaHei"
    elif system == "Darwin":
        cn_font, fallback = "PingFang SC", "Arial Unicode MS"
    else:
        cn_font, fallback = "Noto Sans CJK SC", "Arial Unicode MS"
    doc.styles["Normal"].font.name = fallback
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), cn_font)
    doc.styles["Normal"].font.size = Pt(10.5)


def _score_color(score10: float) -> RGBColor:
    """分数为 1-10 制。"""
    if score10 >= 8:
        return RGBColor(0, 128, 0)
    if score10 >= 6:
        return RGBColor(0, 100, 200)
    if score10 >= 4:
        return RGBColor(255, 165, 0)
    return RGBColor(255, 0, 0)


def _add_title(doc: Document, title: str):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 108, 73)  # VRonly 主题绿
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()


def _add_score_block(doc: Document, label: str, score10: int, analysis: str):
    p = doc.add_paragraph()
    name_run = p.add_run(f"{label}：")
    name_run.font.bold = True
    name_run.font.size = Pt(13)
    score_run = p.add_run(f"{score10} / 10")
    score_run.font.bold = True
    score_run.font.size = Pt(13)
    score_run.font.color.rgb = _score_color(score10)
    if analysis:
        doc.add_paragraph(analysis)


def _add_bullets(doc: Document, heading: str, items: list):
    if not items:
        return
    p = doc.add_paragraph()
    p.add_run(heading).font.bold = True
    for it in items:
        doc.add_paragraph(str(it), style="List Bullet")


def generate_topic_report(result: dict, output_path: Path) -> Path:
    """根据评估结果 dict 生成 Word 报告，返回文件路径。"""
    try:
        doc = Document()
        _set_document_style(doc)

        _add_title(doc, "选题评估报告")

        topic = result.get("topic", {})
        scores = result.get("scores", {})

        # 基本信息
        for label, key in [("研究问题", "question"), ("研究方向", "description"),
                           ("学科", "discipline"), ("学位阶段", "degree_level")]:
            val = topic.get(key)
            if val:
                p = doc.add_paragraph()
                p.add_run(f"{label}：").font.bold = True
                p.add_run(str(val))
        p = doc.add_paragraph()
        p.add_run("论文类别：").font.bold = True
        p.add_run(get_type_name(topic.get("paper_type", "humanities")))
        p = doc.add_paragraph()
        p.add_run("评估时间：").font.bold = True
        p.add_run(datetime.now().strftime("%Y年%m月%d日 %H:%M"))
        doc.add_paragraph("_" * 80)

        # 综合结论
        doc.add_heading("一、综合结论", level=1)
        overall = scores.get("overall", 0)
        cp = doc.add_paragraph()
        cp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = cp.add_run(f"{overall}")
        run.font.size = Pt(40)
        run.font.bold = True
        run.font.color.rgb = _score_color(overall)
        cp.add_run(" / 10").font.size = Pt(14)
        vp = doc.add_paragraph()
        vp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        vrun = vp.add_run(f"评估结论：{scores.get('verdict', '')}")
        vrun.font.size = Pt(13)
        vrun.font.bold = True

        # 三维评分
        doc.add_heading("二、三维评分", level=1)
        analysis = result.get("analysis", {})
        _add_score_block(doc, "创新性", scores.get("innovation", 5), analysis.get("innovation", ""))
        _add_score_block(doc, "可行性", scores.get("feasibility", 5), analysis.get("feasibility", ""))
        _add_score_block(doc, "重要性", scores.get("importance", 5), analysis.get("importance", ""))

        # 亮点 / 挑战 / 建议
        doc.add_heading("三、分析与建议", level=1)
        _add_bullets(doc, "核心创新点：", result.get("key_novelties"))
        _add_bullets(doc, "技术难点 / 挑战：", result.get("technical_challenges"))
        _add_bullets(doc, "改进建议：", result.get("improvement_suggestions"))

        # 相关文献
        papers = result.get("related_papers") or []
        doc.add_heading("四、相关文献", level=1)
        if papers:
            for i, p in enumerate(papers, 1):
                authors = ", ".join(p.get("authors", [])[:3])
                year = p.get("year") or ""
                line = f"[{i}] {authors}. {p.get('title', '')}. {year}".strip()
                doc.add_paragraph(line, style="List Number")
        else:
            doc.add_paragraph("未检索到直接相关文献（评估基于学科常识，建议人工补充检索）。")

        # 页脚
        doc.add_paragraph()
        doc.add_paragraph("_" * 80)
        fp = doc.add_paragraph()
        fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        frun = fp.add_run("本报告由 VRonly 选题评估系统自动生成，仅供参考")
        frun.font.size = Pt(9)
        frun.font.color.rgb = RGBColor(128, 128, 128)
        frun.italic = True

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        logger.info(f"[topic.report] 选题评估报告生成: {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"[topic.report] 报告生成失败: {e}")
        raise
