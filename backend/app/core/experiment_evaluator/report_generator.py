"""
实验设计评审 Word 报告生成器。
输入为 evaluator.review_experiment() 的结果 dict。
"""
from __future__ import annotations

import platform
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from loguru import logger

_SEVERITY_COLOR = {"high": RGBColor(255, 0, 0), "medium": RGBColor(255, 140, 0), "low": RGBColor(0, 100, 200)}
_SEVERITY_LABEL = {"high": "高", "medium": "中", "low": "低"}


def _set_document_style(doc: Document):
    system = platform.system()
    if system == "Windows":
        cn_font, fallback = "微软雅黑", "Microsoft YaHei"
    elif system == "Darwin":
        cn_font, fallback = "PingFang SC", "Arial Unicode MS"
    else:
        cn_font, fallback = "Noto Sans CJK SC", "Arial Unicode MS"
    doc.styles["Normal"].font.name = fallback
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), cn_font)
    doc.styles["Normal"].font.size = Pt(10.5)


def _score_color(s10: float) -> RGBColor:
    if s10 >= 8:
        return RGBColor(0, 128, 0)
    if s10 >= 6:
        return RGBColor(0, 100, 200)
    if s10 >= 4:
        return RGBColor(255, 165, 0)
    return RGBColor(255, 0, 0)


def _add_bullets(doc: Document, heading: str, items: list):
    if not items:
        return
    p = doc.add_paragraph()
    p.add_run(heading).font.bold = True
    for it in items:
        doc.add_paragraph(str(it), style="List Bullet")


def generate_experiment_report(result: dict, output_path: Path) -> Path:
    try:
        doc = Document()
        _set_document_style(doc)

        # 标题
        p = doc.add_paragraph()
        run = p.add_run("实验设计评审报告")
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 108, 73)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()

        scores = result.get("scores", {})
        if result.get("discipline"):
            pp = doc.add_paragraph()
            pp.add_run("学科：").font.bold = True
            pp.add_run(result["discipline"])
        pp = doc.add_paragraph()
        pp.add_run("评审时间：").font.bold = True
        pp.add_run(datetime.now().strftime("%Y年%m月%d日 %H:%M"))
        doc.add_paragraph("_" * 80)

        # 综合结论
        doc.add_heading("一、综合结论", level=1)
        cp = doc.add_paragraph()
        cp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r = cp.add_run(f"{scores.get('overall', 0)}")
        r.font.size = Pt(40)
        r.font.bold = True
        r.font.color.rgb = _score_color(scores.get("overall", 0))
        cp.add_run(" / 10").font.size = Pt(14)
        vp = doc.add_paragraph()
        vp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        vr = vp.add_run(f"评审结论：{scores.get('verdict', '')}")
        vr.font.size = Pt(13)
        vr.font.bold = True

        # 两维评分
        doc.add_heading("二、评分", level=1)
        analysis = result.get("analysis", {})
        for label, key in [("科学性", "scientific_validity"), ("完整性", "completeness")]:
            sp = doc.add_paragraph()
            sp.add_run(f"{label}：").font.bold = True
            sr = sp.add_run(f"{scores.get(key, 5)} / 10")
            sr.font.bold = True
            sr.font.color.rgb = _score_color(scores.get(key, 5))
            if analysis.get(key):
                doc.add_paragraph(analysis[key])

        # 结构性错误
        flaws = result.get("detected_flaws", [])
        doc.add_heading("三、结构性错误排查", level=1)
        if flaws:
            for f in flaws:
                fp = doc.add_paragraph(style="List Bullet")
                fr = fp.add_run(str(f))
                fr.font.color.rgb = RGBColor(200, 0, 0)
        else:
            doc.add_paragraph("✓ 未发现清单内的结构性错误（伪重复/混杂/批次/别名等）。")

        # 风险
        risks = result.get("risks", [])
        doc.add_heading("四、风险识别", level=1)
        if risks:
            for rk in risks:
                sev = (rk.get("severity") or "medium").lower()
                rp = doc.add_paragraph(style="List Bullet")
                tag = rp.add_run(f"[{_SEVERITY_LABEL.get(sev, '中')}] ")
                tag.font.bold = True
                tag.font.color.rgb = _SEVERITY_COLOR.get(sev, _SEVERITY_COLOR["medium"])
                rp.add_run(f"{rk.get('type', '')}：{rk.get('description', '')}")
        else:
            doc.add_paragraph("未识别出显著风险。")

        # 成本 + 建议
        doc.add_heading("五、成本估算与改进建议", level=1)
        if result.get("cost_estimate"):
            cp2 = doc.add_paragraph()
            cp2.add_run("成本/时间估算：").font.bold = True
            cp2.add_run(result["cost_estimate"])
        _add_bullets(doc, "方法论改进建议：", result.get("methodology_suggestions"))

        # 页脚
        doc.add_paragraph()
        doc.add_paragraph("_" * 80)
        fp = doc.add_paragraph()
        fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        fr = fp.add_run("本报告由 VRonly 实验设计评审系统自动生成，仅供参考")
        fr.font.size = Pt(9)
        fr.font.color.rgb = RGBColor(128, 128, 128)
        fr.italic = True

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        logger.info(f"[exp.report] 实验评审报告生成: {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"[exp.report] 报告生成失败: {e}")
        raise
