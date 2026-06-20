"""
评价报告生成器 - 生成Word格式的评价报告
使用python-docx库生成格式良好的评价报告
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from typing import Dict, Optional, List, Tuple
from datetime import datetime
import logging

from app.schemas.evaluation import EvaluationResponse, EvaluationResult
from app.core.evaluator.chart_generator import generate_radar_chart
from app.core.evaluator.prompts import DIMENSION_WEIGHTS, DIMENSION_NAMES, get_type_name

logger = logging.getLogger(__name__)


class ReportGenerator:
    """评价报告生成器"""

    def __init__(self):
        """初始化报告生成器"""
        pass

    def generate_report(
        self,
        evaluation: EvaluationResponse,
        output_path: Path,
        include_chart: bool = True,
        paper_type: str = "humanities",
        integrity_result: Optional[Dict] = None,
    ) -> Path:
        """
        生成评价报告

        :param evaluation: 评价结果数据
        :param output_path: 输出文件路径
        :param include_chart: 是否包含雷达图
        :param paper_type: 论文类别
        :param integrity_result: 学术诚信检查结果
        :return: 生成的报告文件路径
        """
        try:
            # 创建文档
            doc = Document()

            # 设置文档默认字体
            self._set_document_style(doc)

            # 添加标题
            self._add_title(doc, "论文智能评价报告")

            # 添加基本信息
            self._add_basic_info(doc, evaluation, paper_type)

            # 添加综合评分
            self._add_overall_score(doc, evaluation.overall_score)

            # 添加执行摘要
            self._add_executive_summary(doc, evaluation, paper_type)

            # 添加雷达图（如果需要）
            if include_chart:
                chart_path = self._generate_and_add_chart(doc, evaluation, output_path)
                self._add_dimension_score_table(doc, evaluation)
            else:
                # 即使没有图表，也添加标题和汇总表
                doc.add_heading("二、评价维度分析", level=1)
                self._add_dimension_score_table(doc, evaluation)

            # 添加关键问题聚合分析
            self._add_issue_aggregation(doc, evaluation)

            # 添加分维度评价
            self._add_dimension_evaluations(doc, evaluation)

            # 添加学术诚信评估（如果有）
            if integrity_result is not None:
                self._add_integrity_section(doc, integrity_result)

            # 添加页脚
            self._add_footer(doc)

            # 保存文档
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)

            logger.info(f"评价报告生成成功: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"报告生成失败: {str(e)}")
            raise ReportGenerationError(f"报告生成失败: {str(e)}")

    def _set_document_style(self, doc: Document):
        """设置文档样式"""
        import platform
        system = platform.system()
        if system == 'Windows':
            cn_font = '微软雅黑'
            fallback = 'Microsoft YaHei'
        elif system == 'Darwin':
            cn_font = 'PingFang SC'
            fallback = 'Arial Unicode MS'
        else:
            cn_font = 'Noto Sans CJK SC'
            fallback = 'Arial Unicode MS'
        doc.styles['Normal'].font.name = fallback
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), cn_font)
        doc.styles['Normal'].font.size = Pt(10.5)

    def _add_title(self, doc: Document, title: str):
        """添加文档标题"""
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 添加空行
        doc.add_paragraph()

    def _add_basic_info(self, doc: Document, evaluation: EvaluationResponse, paper_type: str = "humanities"):
        """添加基本信息"""
        # 论文标题
        p = doc.add_paragraph()
        p.add_run("论文标题：").font.bold = True
        p.add_run(evaluation.paper_title)

        # 论文类别
        p = doc.add_paragraph()
        p.add_run("论文类别：").font.bold = True
        p.add_run(get_type_name(paper_type))

        # 评价标准
        p = doc.add_paragraph()
        p.add_run("评价标准：").font.bold = True
        p.add_run("山东省本科毕业论文(设计)抽检评议要素(试行)")

        # 评价时间
        p = doc.add_paragraph()
        p.add_run("评价时间：").font.bold = True
        p.add_run(evaluation.evaluated_at.strftime("%Y年%m月%d日 %H:%M:%S"))

        # 添加分隔线
        doc.add_paragraph("_" * 80)

    def _add_overall_score(self, doc: Document, score: float):
        """添加综合评分"""
        doc.add_heading("一、综合评分", level=1)

        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        score_run = p.add_run(f"{score:.1f}")
        score_run.font.size = Pt(48)
        score_run.font.bold = True
        score_run.font.color.rgb = self._get_score_color(score)

        p.add_run(" / 100").font.size = Pt(16)

        # 评分等级
        grade = self._get_score_grade(score)
        grade_p = doc.add_paragraph()
        grade_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        grade_run = grade_p.add_run(f"评价等级：{grade}")
        grade_run.font.size = Pt(14)
        grade_run.font.bold = True

        doc.add_paragraph()

    def _generate_and_add_chart(
        self,
        doc: Document,
        evaluation: EvaluationResponse,
        output_path: Path
    ) -> Optional[Path]:
        """生成并添加雷达图"""
        try:
            # 准备雷达图数据
            dimensions_data = {
                result.dimension_name: result.score
                for result in evaluation.dimensions.values()
            }

            # 生成雷达图
            chart_filename = output_path.stem + "_radar.png"
            chart_path = output_path.parent / chart_filename

            generate_radar_chart(
                dimensions=dimensions_data,
                output_path=chart_path,
                title=f"《{evaluation.paper_title}》评价雷达图"
            )

            # 添加图表标题
            doc.add_heading("二、评价维度分析", level=1)

            # 插入图片
            doc.add_picture(str(chart_path), width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            doc.add_paragraph()

            return chart_path

        except Exception as e:
            logger.warning(f"雷达图生成失败，跳过图表: {str(e)}")
            return None

    def _add_dimension_evaluations(self, doc: Document, evaluation: EvaluationResponse):
        """添加各维度详细评价"""
        doc.add_heading("四、分维度评价", level=1)

        # 维度顺序（V2.0：5个维度）
        dimension_order = [
            "topic_significance",
            "writing_arrangement",
            "logic_construction",
            "professional_ability",
            "academic_standard",
        ]

        for idx, dimension_key in enumerate(dimension_order, 1):
            if dimension_key not in evaluation.dimensions:
                continue

            result = evaluation.dimensions[dimension_key]
            weight = DIMENSION_WEIGHTS.get(dimension_key, 0)

            # 维度标题（含权重）
            heading = f"{idx}. {result.dimension_name}（权重{weight}分）"
            doc.add_heading(heading, level=2)

            # 得分
            score_p = doc.add_paragraph()
            score_p.add_run("得分：").font.bold = True
            score_run = score_p.add_run(f"{result.score}")
            score_run.font.bold = True
            score_run.font.size = Pt(14)
            score_run.font.color.rgb = self._get_score_color(result.score)
            score_p.add_run(" / 100")

            # 优点
            if result.strengths:
                p = doc.add_paragraph()
                p.add_run("优点：").font.bold = True

                for strength in result.strengths:
                    doc.add_paragraph(strength, style='List Bullet')

            # 不足
            if result.weaknesses:
                p = doc.add_paragraph()
                p.add_run("不足之处：").font.bold = True

                for weakness in result.weaknesses:
                    doc.add_paragraph(weakness, style='List Bullet')

            # 改进建议
            if result.suggestions:
                p = doc.add_paragraph()
                p.add_run("改进建议：").font.bold = True

                for suggestion in result.suggestions:
                    doc.add_paragraph(suggestion, style='List Bullet')

            # 添加分隔线（最后一个维度除外）
            if idx < len(dimension_order):
                doc.add_paragraph()

    def _add_footer(self, doc: Document):
        """添加页脚"""
        doc.add_paragraph()
        doc.add_paragraph("_" * 80)

        footer_p = doc.add_paragraph()
        footer_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        footer_run = footer_p.add_run("本报告由论文智能评价系统自动生成")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        footer_run.italic = True

    def _get_score_color(self, score: float) -> RGBColor:
        """根据分数获取颜色（四档制）"""
        if score >= 90:
            return RGBColor(0, 128, 0)    # 绿色 - 优秀
        elif score >= 75:
            return RGBColor(0, 100, 200)   # 蓝色 - 良好
        elif score >= 60:
            return RGBColor(255, 165, 0)   # 橙色 - 合格
        else:
            return RGBColor(255, 0, 0)     # 红色 - 不合格

    def _get_score_grade(self, score: float) -> str:
        """根据分数获取等级（山东省四档制）"""
        if score >= 90:
            return "优秀"
        elif score >= 75:
            return "良好"
        elif score >= 60:
            return "合格"
        else:
            return "不合格"

    def _set_cell_bg_color(self, cell, hex_color: str):
        """设置表格单元格背景色（需操作 XML）"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def _add_executive_summary(self, doc: Document, evaluation: EvaluationResponse, paper_type: str = "humanities"):
        """添加执行摘要：最强/最弱维度 + 改进优先级"""
        doc.add_heading("总体评价概述", level=2)

        # 找最强和最弱维度
        best_dim_key = None
        best_score = -1
        worst_dim_key = None
        worst_score = 101

        for dim_key, result in evaluation.dimensions.items():
            if result.score > best_score:
                best_score = result.score
                best_dim_key = dim_key
            if result.score < worst_score:
                worst_score = result.score
                worst_dim_key = dim_key

        # 最强维度
        if best_dim_key:
            p = doc.add_paragraph()
            p.add_run("最强维度：")
            best_run = p.add_run(f"{evaluation.dimensions[best_dim_key].dimension_name}（得分：{best_score}）")
            best_run.font.color.rgb = RGBColor(0, 128, 0)
            best_run.font.bold = True

        # 最弱维度
        if worst_dim_key:
            p = doc.add_paragraph()
            p.add_run("需改进维度：")
            worst_run = p.add_run(f"{evaluation.dimensions[worst_dim_key].dimension_name}（得分：{worst_score}）")
            worst_run.font.color.rgb = RGBColor(255, 140, 0)
            worst_run.font.bold = True

        # 论文类别说明
        type_descriptions = {
            "humanities": "（人文社科类）评价重点关注理论运用合理性与论证逻辑严密性。",
            "science_engineering": "（理工农医类）评价重点关注研究方法科学性与数据分析严谨性。",
            "arts": "（艺术类）评价重点关注理论研究与艺术创作实践的结合程度。",
        }
        doc.add_paragraph(type_descriptions.get(paper_type, ""))

        # 计算改进优先级（按失分贡献）
        dimension_order = [
            "topic_significance",
            "writing_arrangement",
            "logic_construction",
            "professional_ability",
            "academic_standard",
        ]

        priority_list = []
        for dim_key in dimension_order:
            if dim_key in evaluation.dimensions:
                weight = DIMENSION_WEIGHTS.get(dim_key, 0)
                score = evaluation.dimensions[dim_key].score
                loss_score = weight * (100 - score) / 100
                priority_list.append((loss_score, evaluation.dimensions[dim_key].dimension_name, dim_key))

        priority_list.sort(key=lambda x: x[0], reverse=True)

        # 显示改进优先级 TOP3
        p = doc.add_paragraph()
        p.add_run("改进优先级：").font.bold = True

        for idx, (loss, dim_name, _) in enumerate(priority_list[:3], 1):
            doc.add_paragraph(
                f"{dim_name}（失分贡献：{loss:.1f}分）",
                style='List Number'
            )

        doc.add_paragraph()

    def _add_dimension_score_table(self, doc: Document, evaluation: EvaluationResponse):
        """添加维度分数汇总表"""
        # 添加小标题
        p = doc.add_paragraph()
        p_run = p.add_run("维度得分汇总")
        p_run.font.size = Pt(14)
        p_run.font.bold = True

        # 创建表格：5个维度 + 1个汇总行
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        # 表头
        header_cells = table.rows[0].cells
        header_texts = ["维度", "权重", "得分", "实际分值"]
        for idx, text in enumerate(header_texts):
            cell = header_cells[idx]
            cell.text = text
            # 设置表头背景色（浅蓝灰）
            self._set_cell_bg_color(cell, "D9E1F2")
            # 表头加粗
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # 维度顺序
        dimension_order = [
            "topic_significance",
            "writing_arrangement",
            "logic_construction",
            "professional_ability",
            "academic_standard",
        ]

        # 添加各维度行
        total_score = 0
        total_weight = 0
        for dim_key in dimension_order:
            if dim_key not in evaluation.dimensions:
                continue

            result = evaluation.dimensions[dim_key]
            weight = DIMENSION_WEIGHTS.get(dim_key, 0)
            actual_score = round(weight * result.score / 100, 1)

            row = table.add_row()
            row.cells[0].text = result.dimension_name
            row.cells[1].text = str(weight)
            row.cells[2].text = str(result.score)
            row.cells[3].text = str(actual_score)

            # 得分列着色
            score_para = row.cells[2].paragraphs[0]
            for run in score_para.runs:
                run.font.color.rgb = self._get_score_color(result.score)
                run.font.bold = True

            total_weight += weight
            total_score += actual_score

        # 添加汇总行
        summary_row = table.add_row()
        summary_row.cells[0].text = "综合评分"
        summary_row.cells[1].text = str(total_weight)
        summary_row.cells[2].text = str(round(evaluation.overall_score, 1))
        summary_row.cells[3].text = str(round(evaluation.overall_score, 1))

        # 汇总行加粗
        for cell in summary_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # 设置列宽
        table.columns[0].width = Inches(2.0)
        table.columns[1].width = Inches(0.8)
        table.columns[2].width = Inches(0.8)
        table.columns[3].width = Inches(1.2)

        # 表格居中对齐
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    if idx == 0:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    else:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph()

    def _add_issue_aggregation(self, doc: Document, evaluation: EvaluationResponse):
        """添加关键问题聚合分析：去重并按主题词归并"""
        doc.add_heading("三、关键问题聚合分析", level=1)

        # 收集所有问题
        all_weaknesses = []
        for dim_key, result in evaluation.dimensions.items():
            for w in result.weaknesses:
                all_weaknesses.append((dim_key, w))

        if not all_weaknesses:
            doc.add_paragraph("各维度评价条目未发现明显问题。")
            doc.add_paragraph()
            return

        # 关键词分组定义
        topic_keywords = {
            "文献综述与背景": ["文献", "综述", "参考文献", "文献调研", "引用", "背景"],
            "研究方法与实证": ["方法", "实证", "研究设计", "数据", "分析方法", "问卷"],
            "逻辑论证与衔接": ["逻辑", "论证", "论据", "衔接", "过渡", "层次", "递进"],
            "语言表达与格式": ["表达", "语言", "文字", "措辞", "行文", "术语", "格式", "标注", "规范"],
            "创新性与深度": ["创新", "新颖", "贡献", "突破", "独特", "深度"],
        }

        # 按主题词匹配
        theme_hits = {}  # theme -> [(dim_key, weakness_text), ...]
        for topic, keywords in topic_keywords.items():
            for dim_key, w_text in all_weaknesses:
                if any(kw in w_text for kw in keywords):
                    if topic not in theme_hits:
                        theme_hits[topic] = []
                    theme_hits[topic].append((dim_key, w_text))

        # 按涉及维度数排序，过滤只有1条命中的主题
        sorted_themes = [
            (t, hits) for t, hits in sorted(
                theme_hits.items(),
                key=lambda x: len(set(dim for dim, _ in x[1])),
                reverse=True
            )
            if len(set(dim for dim, _ in hits)) >= 2 and len(hits) >= 2
        ]

        if not sorted_themes:
            doc.add_paragraph("各维度评价条目未发现高频共性问题。")
            doc.add_paragraph()
            return

        # 展示前4个主题
        for topic, hits in sorted_themes[:4]:
            affected_dims = set(dim for dim, _ in hits)
            p = doc.add_paragraph()
            p_run = p.add_run(f"【{topic}】")
            p_run.font.bold = True
            p_run.font.color.rgb = RGBColor(200, 0, 0)
            p.add_run(f" 涉及 {len(affected_dims)} 个维度")

            # 去重并展示问题
            seen = set()
            for dim_key, w_text in hits:
                if w_text not in seen:
                    dim_name = evaluation.dimensions[dim_key].dimension_name
                    doc.add_paragraph(
                        f"[{dim_name}] {w_text}",
                        style='List Bullet'
                    )
                    seen.add(w_text)

            doc.add_paragraph()

    def _add_integrity_section(self, doc: Document, integrity_result: Dict):
        """添加学术诚信评估部分"""
        doc.add_heading("五、学术诚信评估", level=1)

        has_risk = integrity_result.get("has_risk", False)
        risk_items = integrity_result.get("risk_items", [])
        note = integrity_result.get("note", "")

        if not has_risk:
            # 无风险提示
            p = doc.add_paragraph()
            p_run = p.add_run("✓ 未发现明显学术诚信风险")
            p_run.font.color.rgb = RGBColor(0, 128, 0)
            p_run.font.bold = True
        else:
            # 有风险提示
            if risk_items:
                p = doc.add_paragraph()
                p_run = p.add_run("⚠ 发现以下潜在风险，建议人工复核：")
                p_run.font.color.rgb = RGBColor(200, 60, 0)
                p_run.font.bold = True

                for item in risk_items:
                    doc.add_paragraph(item, style='List Bullet')
            else:
                p = doc.add_paragraph()
                p_run = p.add_run("⚠ 风险条目未能解析，建议人工审查")
                p_run.font.color.rgb = RGBColor(200, 60, 0)

        # 评估说明
        if note:
            p = doc.add_paragraph()
            p_run = p.add_run(f"评估说明：{note}")
            p_run.font.color.rgb = RGBColor(100, 100, 100)
            p_run.italic = True

        # 免责声明
        p = doc.add_paragraph()
        p_run = p.add_run("注：本诚信评估基于文本特征的初步分析，不构成最终学术不端认定依据。")
        p_run.font.size = Pt(9)
        p_run.font.color.rgb = RGBColor(128, 128, 128)
        p_run.italic = True

        doc.add_paragraph()


class ReportGenerationError(Exception):
    """报告生成错误"""
    pass


# 工厂函数
def generate_report(
    evaluation: EvaluationResponse,
    output_path: Path,
    include_chart: bool = True,
    paper_type: str = "humanities",
    integrity_result: Optional[Dict] = None,
) -> Path:
    """
    生成评价报告的便捷函数

    :param evaluation: 评价结果
    :param output_path: 输出路径
    :param include_chart: 是否包含图表
    :param paper_type: 论文类别
    :param integrity_result: 学术诚信检查结果
    :return: 报告文件路径
    """
    generator = ReportGenerator()
    return generator.generate_report(evaluation, output_path, include_chart, paper_type, integrity_result)